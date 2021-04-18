from docx import *
from shutil import copyfile
from docx.shared import Pt


def Bio_Word(word,year,month,day):
    f = open('bio latest17.docx', 'rb')
    f.read()

    document = Document(f)
    for i in document.tables:
        for k in i.rows:
            for j in k.cells:
                for n in j.paragraphs:

                    if n.text == 'أسـم المريض :       المحترم':
                        n.text = f'أسـم المريض :{word}       المحترم'
                        for run in n.runs:
                            run.font.size = Pt(11)
                            run.font.name = 'Monotype Koufi'
                    if n.text == 'حضرة الدكتور   :       المحترم':
                        n.text = f'حضرة الدكتور   : {word}      المحترم'
                        for run in n.runs:
                            run.font.size = Pt(11)
                            run.font.name = 'Monotype Koufi'
                    if n.text == 'Random  blood sugar :':
                        n.text = f'Random  blood sugar : {word}'
                        for run in n.runs:
                            run.bold = True
                            run.font.size = Pt(11)
                            run.font.name = 'Tahoma'
                    if n.text == 'Blood Urea               :':
                        n.text = f'Blood Urea               : {word}'
                        for run in n.runs:
                            run.bold = True
                            run.font.size = Pt(11)
                            run.font.name = 'Tahoma'
                    if n.text == 'S. Creatinin               :':
                        n.text = f'S. Creatinin               : {word}'
                        for run in n.runs:
                            run.bold = True
                            run.font.size = Pt(11)
                            run.font.name = 'Tahoma'
                    if n.text == 'S. Uric acid                  :':
                        n.text = f'S. Uric acid                  : {word}'
                        for run in n.runs:
                            run.bold = True
                            run.font.size = Pt(11)
                            run.font.name = 'Tahoma'
                    if n.text == 'S. Cholesterol            :':
                        n.text = f'S. Cholesterol            : {word}'
                        for run in n.runs:
                            run.bold = True
                            run.font.size = Pt(11)
                            run.font.name = 'Tahoma'
                    if n.text == 'S. Triglycerid             :':
                        n.text = f'S. Triglycerid             : {word}'
                        for run in n.runs:
                            run.bold = True
                            run.font.size = Pt(11)
                            run.font.name = 'Tahoma'
                    if n.text == 'Total serum Bilirubin:':
                        n.text = f'Total serum Bilirubin: {word}'
                        for run in n.runs:
                            run.bold = True
                            run.font.size = Pt(11)
                            run.font.name = 'Tahoma'
                    if n.text == 'S.Calcium :':
                        n.text = f'S.Calcium : {word}'
                        for run in n.runs:
                            run.bold = True
                            run.font.size = Pt(11)
                            run.font.name = 'Tahoma'
                    if n.text == 'Vitamin D              :':
                        n.text = f'Vitamin D              : {word}'
                        for run in n.runs:
                            run.bold = True
                            run.font.size = Pt(11)
                            run.font.name = 'Tahoma'

                    if n.text == 'Date:    /     / 20':
                        n.text = f'Date:   {day} /  {month} / {year}'
                        for run in n.runs:
                            run.bold = True
                            run.font.size = Pt(12)
                            run.font.name = 'Tahoma'
                    print('2' + n.text + '1')
    document.save('bio latest17.docx')
    f.close()
