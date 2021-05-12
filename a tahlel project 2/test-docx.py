
# from shutil import copyfile
#
# # def Bio_Word(word):
from docx.shared import Pt, Inches

# word_type = ['hemo', 'bio']
# all_analyst = ['r', 'h', 'random', 'd', 'g', 'll', 'llii', 'iio9', 'kkj', 'iiip', 'gggyu7', 'jjjk', 'ppp', 'new1',
#                'new2', 'new3', 'new4', 'new5', 'new6', 'new7', 'new8', 'new9', 'new11', 'new22', 'new33', 'new44',
#                'new55', 'new66', 'new77', 'new88', 'new99']
# all_result = ['r1', '2h', '5k', '5d', '5g', '5j', '5op', '5lo', ';5l', 'y5u', 'f5hk', 'llojuo', 'iiui', 'anew1',
#               'anew2', 'anew3', 'anew4', 'anew5', 'anew6', 'anew7', 'anew8', 'anew9', 'anew11', 'qnew22', 'anew33',
#               'anew44', 'anew55', 'anew66', 'anew77', 'anew88', 'anew99']
# list_of_hemo_and_bio_analysts2 = []
# list_of_hemo_and_bio_results2 = []
# bio_hemo_all_analysts = ['r', 'h', 'random', 'd', 'g', 'll', 'llii', 'iio9', 'kkj', 'iiip', 'gggyu7', 'jjjk', 'ppp',
#                          'new1', 'new2', 'new3', 'new4', 'new5', 'new6', 'new7', 'new8', 'new9', 'new11', 'new22',
#                          'new33', 'new44', 'new55', 'new66', 'new77', 'new88', 'new99']
# for index, itr in enumerate(all_analyst):
#     if itr in bio_hemo_all_analysts and itr not in list_of_hemo_and_bio_analysts2:
#         list_of_hemo_and_bio_analysts2.append(itr)
#         list_of_hemo_and_bio_results2.append(all_result[index])
# number_of_docx_rows = 23
# if len(bio_hemo_all_analysts) < number_of_docx_rows:
#     f = open('test-mydocx.docx', 'rb')
#     f.read()
#     document = Document(f)
# else:
#     print('two files open')
#     f = open('test-mydocx.docx', 'rb')
#     f.read()
#     document = Document(f)
#     f2 = open('test-mydocx2.docx', 'rb')
#     document2 = Document(f2)
# row_num = False
# for i in document.tables:
#     for k in i.rows:
#         for j in k.cells:
#             for n in j.paragraphs:
#                 for row in range(0, len(list_of_hemo_and_bio_analysts2)):
#                     if 'bio' in word_type and 'hemo' in word_type:
#                         if n.text == str(row + 1):
#                             print(all_analyst[row], 'not')
#                             n.text = str(list_of_hemo_and_bio_analysts2[row]) + '  :'
#                             run = n.runs
#                             font = run[0].font
#                             font.bold = True
#                             font.size = Pt(11)
#                             font.name = 'Tahoma'
#                             n2 = n.add_run(str(list_of_hemo_and_bio_results2[row]))
#                             run = n2
#                             font = run.font
#                             font.bold = False
#                             font.size = Pt(11)
#                             font.name = 'Tahoma'
#                         if n.text == str(row + 1) + 'unit':
#                             # n.text = analyst_and_result['unit']
#                             n.text = 'myunit'
#                             run1 = n
#                             font1 = run1.runs[0].font
#                             font1.bold = True
#                             font1.size = Pt(12)
#                             font1.name = 'Times New Roman'
#                         if n.text == str(row + 1) + 'defult':
#                             # n.text = analyst_and_result['normal']
#                             n.text = 'mydef'
#                             run1 = n
#                             font1 = run1.runs[0].font
#                             font1.bold = True
#                             font1.size = Pt(12)
#                             font1.name = 'Times New Roman'
#                         if row > 23:
#                             row_num = True
#
#                     else:
#                         print('hahahah you will dont showing hahahaha')
#                         if 'bio' in word_type or 'hemo' in word_type:
#                             if n.text == str(row + 1):
#                                 print(all_analyst[row], 'not')
#                                 n.text = str(list_of_hemo_and_bio_analysts2[row]) + '  :'
#                                 run = n.runs
#                                 font = run[0].font
#                                 font.bold = True
#                                 font.size = Pt(11)
#                                 font.name = 'Tahoma'
#                                 n2 = n.add_run(str(list_of_hemo_and_bio_results2[row]))
#                                 run = n2
#                                 font = run.font
#                                 font.bold = False
#                                 font.size = Pt(11)
#                                 font.name = 'Tahoma'
#                             if n.text == str(row + 1) + 'unit':
#                                 # n.text = analyst_and_result['unit']
#                                 n.text = 'myunit'
#                                 run1 = n
#                                 font1 = run1.runs[0].font
#                                 font1.bold = True
#                                 font1.size = Pt(12)
#                                 font1.name = 'Times New Roman'
#                             if n.text == str(row + 1) + 'defult':
#                                 # n.text = analyst_and_result['normal']
#                                 n.text = 'mydef'
#                                 run1 = n
#                                 font1 = run1.runs[0].font
#                                 font1.bold = True
#                                 font1.size = Pt(12)
#                                 font1.name = 'Times New Roman'
#                 for mynum in range(1, 25):
#                     if n.text == str(mynum) or n.text == str(mynum) + 'unit' or n.text == str(
#                             mynum) + 'defult' or n.text == '':
#                         n.text = ''
#                         n.clear()
# if row_num:
#     try:
#         for i2 in document2.tables:
#             for k2 in i2.rows:
#                 for j2 in k2.cells:
#                     for n3 in j2.paragraphs:
#                         for row in range(23, len(list_of_hemo_and_bio_analysts2)):
#                             if n3.text == str(row - 22):
#                                 # print(all_analyst[row],'not')
#                                 n3.text = str(list_of_hemo_and_bio_analysts2[row]) + '  :'
#                                 run = n3.runs
#                                 font = run[0].font
#                                 font.bold = True
#                                 font.size = Pt(11)
#                                 font.name = 'Tahoma'
#                                 n2 = n3.add_run(str(list_of_hemo_and_bio_results2[row]))
#                                 run = n2
#                                 font = run.font
#                                 font.bold = False
#                                 font.size = Pt(11)
#                                 font.name = 'Tahoma'
#                             if n3.text == str(row - 22) + 'unit':
#                                 # n.text = analyst_and_result['unit']
#                                 n3.text = 'myunit'
#                                 run1 = n3
#                                 font1 = run1.runs[0].font
#                                 font1.bold = True
#                                 font1.size = Pt(12)
#                                 font1.name = 'Times New Roman'
#                             if n3.text == str(row - 22) + 'defult':
#                                 # n.text = analyst_and_result['normal']
#                                 n3.text = 'mydef'
#                                 run1 = n3
#                                 font1 = run1.runs[0].font
#                                 font1.bold = True
#                                 font1.size = Pt(12)
#                                 font1.name = 'Times New Roman'
#                         for mynum in range(1, 20):
#                             if n3.text == str(mynum) or n3.text == str(mynum) + 'unit' or n3.text == str(
#                                     mynum) + 'defult' or n3.text == '':
#                                 n3.text = ''
#                                 n3.clear()
#     except Exception as e:
#         print(e)
# row_num = False
# document.save('gg.docx')
# f.close()
# document2.save('gg2.docx')
# f2.close()
# # print(row,'row')
# # print(col,'col')
# # print(table,'table')
#
# # from docx import Document
# #
# # from docx.text.paragraph import Paragraph
# # from docx.oxml.xmlchemy import OxmlElement
# # from docx.enum.text import WD_TAB_ALIGNMENT
# #
# # def insert_paragraph_after(paragraph, text=None, style=None):
# #     """Insert a new paragraph after the given paragraph."""
# #     new_p = OxmlElement("w:p")
# #     paragraph._p.addnext(new_p)
# #     new_para = Paragraph(new_p, paragraph._parent)
# #
# #     if text:
# #         new_para.add_run(text)
# #
# #     new_para.style.font.name = 'Tahoma'
# #     new_para.style.font.size = Pt(9)
# #
# #     return new_para
# #
# # def main():
# #     # Create a minimal document
# #     f = open('احمد سعد.docx', 'rb')
# #     f.read()
# #     document = Document(f)
# #     # p1 = document.add_paragraph("First Paragraph.")
# #     # p2 = document.add_paragraph("Second Paragraph.")
# #     # Insert a paragraph wedged between p1 and p2
# #     for p in document.paragraphs:
# #         if p.text=='علي:':
# #             print('start')
# #             n2=n.add_run()
# #             new.font.bold=False
# #             new.font.size=Pt(10)
# #         # print('2'+p.text+'1')
# #         # if p.text=='علي  قاسم':
# #         #     print('d')
# #         #     insert_paragraph_after(p, "Paragraph One And A Half.")
# #
# #
# #     # Test if the function succeeded
# #     document.save('gg.docx')
# # main()
# word_types = ['bio', ' bio', 'bio', 'hemo', 'hemo','GSE','GSE''GSE']
# bio_analysts=0
# hemo_analysts=0
# GSE_analysts=0
# GUE_analysts=0
# SFA_analysts=0
# HRMON_analysts=0
# for iiii in word_types:
#     if iiii =='bio':
#         bio_analysts +=1
#     if iiii =='hemo':
#         hemo_analysts +=1
#     if iiii =='GSE':
#         GSE_analysts +=1
# if bio_analysts != 0:
#     if row ==0+1:
#         if n.text == str(row):
#             n.text=str(word_types[0])
# if hemo_analysts !=0:
#     if row ==(bio_analysts+1):
#         if n.text == str(row):
#             n.text = str(word_types[bio_analysts])
# if GSE_analysts !=0:
#     if row ==(hemo_analysts+bio_analysts+1):
#         if n.text==str(row):
#             n.text =str(word_types[hemo_analysts+bio_analysts])
# if GUE_analysts !=0:
#     if row ==(hemo_analysts+bio_analysts+GSE_analysts+1):
#         if n.text==str(row):
#             n.text =str(word_types[hemo_analysts+bio_analysts+GSE_analysts])
#
# if SFA_analysts !=0:
#     if row ==(hemo_analysts+bio_analysts+GSE_analysts+GUE_analysts+1):
#         if n.text==str(row):
#             n.text =str(word_types[hemo_analysts+bio_analysts+GSE_analysts+GUE_analysts])
#
# if HRMON_analysts !=0:
#     if row ==(hemo_analysts+bio_analysts+GSE_analysts+GUE_analysts+SFA_analysts+1):
#         if n.text==str(row):
#             n.text =str(word_types[hemo_analysts+bio_analysts+GSE_analysts+GUE_analysts+SFA_analysts]
