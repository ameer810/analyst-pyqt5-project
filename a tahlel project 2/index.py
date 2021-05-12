import datetime
import os
import sys
import MySQLdb
from PyQt5.QtCore import *
from PyQt5.QtGui import *
from PyQt5.QtWidgets import *
from PyQt5.uic import loadUiType
from docx import *
from docx.shared import Pt
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from win32com import client
import searchDialog
import searchDialog2
import multyclass
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from threading import Thread
from mymain import Ui_MainWindow as main_wind
show_clients_check = True
FORM_CLASS, _ = loadUiType("design.ui")
show_all_sales_in_clients_page = False
user_id = ''
analyst_name_for_update_before = ''
analyst_name_for_update = ''
client_id_glob = 0
chick_if_add_new = False
if_print = False
analysts_name_glo = []
clients_name_glo = []
clients_name_glo_clients_page = []
from_start = False
first_text = ''
first_text2 = ''
first_text3 = ''
addTrue = True
first_text_category1 = ''
first_text_category2 = ''
mylist = []
select_by_date = False
sd = None
echo_mode_num = 0
echo_mode_num2 = 0
Add_all_analysts_items_list = []
edit_employee_check = True
Edit_employee = False
search_info_by_date = False
all_doctors = []
Edit_Doctor = True
Delete_Doctor = True
default_analysts = []
default_analysts2 = []


class mainapp(QMainWindow, FORM_CLASS):
    def __init__(self, parent=None):
        global from_start
        super(mainapp, self).__init__(parent)
        QMainWindow.__init__(self)
        self.setupUi(self)
        self.tabWidget.tabBar().setVisible(False)
        self.DB()
        self.Show_default_statics()
        self.Delete_Files()
        self.add_clients_to_combo()
        from_start = True
        self.Show_All_The_Analysts()
        self.handel_buttons()
        self.Show_All_The_Sales()
        self.Show_all_analysts_in_combo()
        self.Show_all_buys()
        self.History()
        self.groupBox.setEnabled(False)
        self.Show_paths()
        self.add_Analyst_to_list()
        self.Auto_complete_combo()
        self.add_client_to_list()
        self.Auto_complete_combo2()
        self.add_client_to_list4()
        self.Auto_complete_combo4()
        self.Show_Word_Doc_Data()
        self.Add_Doctor_Data()
        self.dateEdit_6.setDate(datetime.date.today())
        self.dateEdit_5.setDate(datetime.date.today())
        self.dateEdit.setDate(datetime.date.today())
        # os.system("TASKKILL /F /IM WINWORD.exe")
        # os.system('start WINWORD.exe')
        self.tableWidget_5.setColumnHidden(6, True)

    def DB(self):
        self.db = MySQLdb.connect(host='localhost', user='root', password='12345', db='tahlel2', charset="utf8",
                                  use_unicode=True, port=3306)
        self.cur = self.db.cursor()

    def Show_Word_Doc_Data(self):
        self.cur.execute(''' SELECT * FROM word ''')
        data = self.cur.fetchone()
        self.lineEdit_16.setText(data[9])
        self.lineEdit_18.setText(data[10])
        self.lineEdit_20.setText(data[11])
        self.lineEdit_19.setText(data[12])
        self.lineEdit_23.setText(data[2])
        self.lineEdit_22.setText(data[3])
        self.lineEdit_37.setText(data[4])
        self.lineEdit_41.setText(data[5])
        self.lineEdit_35.setText(data[6])
        self.lineEdit_42.setText(data[7])
        self.lineEdit_43.setText(data[1])
        self.lineEdit_34.setText(data[8])
        self.lineEdit_45.setText(data[13])
        self.lineEdit_44.setText(data[14])
        self.lineEdit_46.setText(data[15])
        self.lineEdit_47.setText(data[16])

    def Update_Word_Doc_Data(self):
        var1 = self.lineEdit_16.text()
        var2 = self.lineEdit_18.text()
        var3 = self.lineEdit_20.text()
        var4 = self.lineEdit_19.text()
        var5 = self.lineEdit_23.text()
        var6 = self.lineEdit_22.text()
        var7 = self.lineEdit_37.text()
        var8 = self.lineEdit_41.text()
        var9 = self.lineEdit_35.text()
        var10 = self.lineEdit_42.text()
        var11 = self.lineEdit_43.text()
        var12 = self.lineEdit_34.text()
        var13 = self.lineEdit_45.text()
        var14 = self.lineEdit_44.text()
        var15 = self.lineEdit_46.text()
        var16 = self.lineEdit_47.text()
        self.cur.execute(
            ''' UPDATE word SET  shop_name=%s,phone1=%s,phone2=%s,employee_name1=%s,employee_name2=%s,employee1_shahada=%s,employee2_shahada=%s,gps=%s,client_name=%s,client_lqb=%s,doctor_name=%s,doctor_lqb=%s,client_name2=%s,client_lqb2=%s,doctor_name2=%s,doctor_lqb2=%s WHERE id=1''',
            (var11, var5, var6, var7, var8, var9, var10, var12, var1, var2, var3, var4, var13, var14, var15, var16,))
        self.db.commit()
        QMessageBox.information(self, 'info', 'تم تحديث المعلومات بنجاح')
        self.Show_Word_Doc_Data()
        self.Add_Data_To_history(4, 7)
        self.History()

    def all_per(self):
        if self.checkBox.isChecked() == True:
            self.checkBox_58.setCheckState(True)
            self.checkBox_57.setCheckState(True)
            self.checkBox_56.setCheckState(True)
            self.checkBox_9.setCheckState(True)
            self.checkBox_55.setCheckState(True)
            self.checkBox_44.setCheckState(True)
            self.checkBox_40.setCheckState(True)
            self.checkBox_13.setCheckState(True)
            self.checkBox_8.setCheckState(True)
            self.checkBox_15.setCheckState(True)
            self.checkBox_9.setCheckState(True)
            self.checkBox_55.setCheckState(True)
            self.checkBox_44.setCheckState(True)
            self.checkBox_40.setCheckState(True)
            self.checkBox_13.setCheckState(True)
            self.checkBox_16.setCheckState(True)
            # self.checkBox_3.setCheckState(False)
            self.checkBox_14.setCheckState(True)
            self.checkBox_7.setCheckState(True)
            self.checkBox_11.setCheckState(True)
            self.checkBox_12.setCheckState(True)
            self.checkBox_46.setCheckState(True)
            self.checkBox_49.setCheckState(True)
            self.checkBox_50.setCheckState(True)
            self.checkBox_47.setCheckState(True)
            self.checkBox_48.setCheckState(True)
            self.checkBox_51.setCheckState(True)
            self.checkBox_52.setCheckState(True)
            self.checkBox_53.setCheckState(True)
            self.checkBox_54.setCheckState(True)
            self.checkBox_44.setCheckState(True)
            self.checkBox_45.setCheckState(True)
            self.checkBox_42.setCheckState(True)
            self.checkBox_43.setCheckState(True)
            self.checkBox_41.setCheckState(True)
        else:
            self.False_checkState()

    def handel_buttons(self):
        global addTrue
        self.pushButton_15.clicked.connect(self.Light_Blue_Theme)
        self.pushButton_9.clicked.connect(self.Dark_Orange_Theme)
        self.pushButton_13.clicked.connect(self.Dark_Blue_Theme)
        self.pushButton_14.clicked.connect(self.Dark_Theme)
        self.pushButton_11.clicked.connect(self.Dark_Gray_Theme)
        self.pushButton.clicked.connect(self.Open_Sales_Page)
        # self.pushButton_6.clicked.connect(self.Open_Login_Page)
        self.pushButton_21.clicked.connect(self.Open_Login_Page)
        self.pushButton_4.clicked.connect(self.Open_Settings_Page)
        self.pushButton_12.clicked.connect(self.Open_Login_Page)
        self.pushButton_3.clicked.connect(self.Open_History_Page)
        self.pushButton_5.clicked.connect(self.Open_clients_Page)
        self.pushButton_2.clicked.connect(self.Open_Analyse_Page)
        self.pushButton_8.clicked.connect(self.Open_ResetPassword_Page)
        self.pushButton_17.clicked.connect(lambda arg='no': self.Sales_Page(for_loop2=arg))
        self.pushButton_30.clicked.connect(self.get_client_id)
        # if not addTrue:
        self.comboBox_16.currentIndexChanged.connect(self.Show_Type_of_result_category)
        self.comboBox_2.currentIndexChanged.connect(self.Show_permissions)
        # QListView.currentChanged()
        # self.pushButton_31.clicked.connect(self.Chick_analyst_category)
        self.pushButton_32.clicked.connect(self.get_total_price)
        # self.pushButton_33.clicked.connect(self.Show_analyst_in_Edit_Or_Delete)
        self.comboBox_21.currentTextChanged.connect(self.Show_analyst_in_Edit_Or_Delete)
        self.comboBox_28.currentTextChanged.connect(self.Show_Doctor_Data)
        self.comboBox_3.currentTextChanged.connect(self.Show_employee_data)
        self.pushButton_29.clicked.connect(self.Clients_Page)
        self.pushButton_16.clicked.connect(self.Add_Buys)
        self.pushButton_20.clicked.connect(self.Show_All_The_Analysts)
        self.pushButton_27.clicked.connect(self.Edit_Analyst)
        self.pushButton_7.clicked.connect(self.Log_In_Chieck)
        self.pushButton_28.clicked.connect(self.Delete_Analyst)
        self.pushButton_22.clicked.connect(self.Delete_All_History_Data)
        self.pushButton_18.clicked.connect(self.Print_Sale_Data)
        self.pushButton_36.clicked.connect(self.Search_In_All_Sales)
        self.comboBox_20.currentIndexChanged.connect(self.Search_In_History)
        self.comboBox_24.currentIndexChanged.connect(self.Search_In_History)
        # self.pushButton_21.clicked.connect(self.Search_In_History)
        self.pushButton_35.clicked.connect(self.clear_data_in_sales)
        self.pushButton_10.clicked.connect(self.Reset_password)
        self.pushButton_26.clicked.connect(self.Add_Analyst)
        self.pushButton_34.clicked.connect(self.Preview)
        # self.pushButton_23.clicked.connect(self.Print_empty_papers)
        self.pushButton_24.clicked.connect(self.Add_Path)
        self.lineEdit_27.textChanged.connect(self.Show_All_Clients)
        self.pushButton_39.clicked.connect(self.ADD_one_REsult_to_ANalyst_results)
        self.pushButton_40.clicked.connect(self.ADD_one_REsult_to_ANalyst_results_IN_EDITMODE)
        self.pushButton_41.clicked.connect(self.Remove_one_REsult_to_ANalyst_results_IN_EDITMODE)
        self.pushButton_50.clicked.connect(self.Remove_one_REsult_to_ANalyst_results)
        self.pushButton_43.clicked.connect(self.Show_search_Widget)
        self.pushButton_44.clicked.connect(self.Show_multy_Dialog)
        self.comboBox_16.view().pressed.connect(self.Set_Chick_State2)
        self.pushButton_46.clicked.connect(self.ADD_one_category_to_analyst_categoryes)
        self.pushButton_53.clicked.connect(self.Add_permissions)
        self.pushButton_25.clicked.connect(self.Add_permissions)
        self.pushButton_42.clicked.connect(self.Show_search_Widget2)
        self.pushButton_53.clicked.connect(self.Add_employee)
        self.pushButton_25.clicked.connect(self.Add_employee)
        self.pushButton_23.clicked.connect(self.Show_statics)
        self.pushButton_45.clicked.connect(self.ADD_one_category_to_analyst_categoryes_IN_EDITMODE)
        self.pushButton_31.clicked.connect(self.Update_Word_Doc_Data)
        self.pushButton_33.pressed.connect(self.echo_mode)
        self.pushButton_33.released.connect(self.echo_mode)
        self.pushButton_38.pressed.connect(self.echo_mode2)
        self.pushButton_38.released.connect(self.echo_mode2)
        self.checkBox.clicked.connect(self.all_per)
        self.Add_all_employee_to_comboBox()
        self.comboBox_4.currentTextChanged.connect(self.Show_client_data_by_current_index)
        self.my_def2()
        self.add_all_subCategory_toList()
        self.pushButton_47.clicked.connect(self.Add_Doctor)
        self.pushButton_48.clicked.connect(self.Update_doctor)
        self.pushButton_49.clicked.connect(self.Delete_doctor)
        self.add_doctors_to_list()
        self.tabWidget_4.currentChanged.connect(self.Show_default_statics)
        self.pushButton_51.clicked.connect(self.Remove_category)
        self.pushButton_52.clicked.connect(self.Remove_category_in_edit)
    def add_doctors_to_list(self):
        global all_doctors
        self.cur.execute(''' SELECT name FROM doctor ''')
        data = self.cur.fetchall()
        all_doctors.clear()
        for i in data:
            if i[0] not in all_doctors:
                all_doctors.append(i[0])

    def Add_Doctor(self):
        try:
            genus = self.comboBox_18.currentIndex()
            rgenus = ''
            if genus:
                rgenus = 'male'
            else:
                rgenus = 'female'
            self.cur.execute(''' INSERT INTO doctor (name,genus) VALUES(%s,%s)''', (self.lineEdit_5.text(), rgenus,))
            self.db.commit()
            QMessageBox.information(self, 'info', 'تم اضافة الطبيب بنجاح')
            self.Add_Doctor_Data()
            self.add_doctors_to_list()
            self.Add_Data_To_history(3, 6)
            self.History()
        except Exception as e:
            print(e, '95kerorr')
            QMessageBox.information(self, '', 'هذا الطبيب موجود بالفعل')

    def Update_doctor(self):
        global all_doctors
        if self.comboBox_28.currentText() in all_doctors:
            try:
                genus = self.comboBox_27.currentIndex()
                rgenus = ''
                if genus:
                    rgenus = 'male'
                else:
                    rgenus = 'female'
                self.cur.execute(''' UPDATE doctor SET name=%s,genus=%s WHERE name=%s''',
                                 ((self.lineEdit_6.text(), rgenus, self.comboBox_28.currentText(),)))
                self.db.commit()
                QMessageBox.information(self, 'info', 'تم تحديث بيانات الطبيب بنجاح')
                self.Add_Doctor_Data()
                self.add_doctors_to_list()
                self.Add_Data_To_history(4, 6)
                self.History()
            except:
                QMessageBox.information(self, '', 'لا يمكن تكرار اسم الطبيب , يرجى تغيير اسم الطبيب')
        else:
            QMessageBox.information(self, '', 'هذا الطبيب غير موجود')

    def Delete_doctor(self):
        if self.comboBox_28.currentText() in all_doctors:
            genus = self.comboBox_27.currentIndex()
            rgenus = ''
            if genus:
                rgenus = 'male'
            else:
                rgenus = 'female'
            warning = QMessageBox.warning(self, '', f'هل انت متأكد من مسح الطبيب {self.comboBox_28.currentText()}؟',
                                          QMessageBox.Yes | QMessageBox.No)
            if warning == QMessageBox.Yes:
                self.cur.execute(''' DELETE FROM doctor  WHERE name=%s''', ((self.comboBox_28.currentText(),)))
                self.db.commit()
                QMessageBox.information(self, 'info', 'تم حذف الطبيب بنجاح')
                self.Add_Doctor_Data()
                self.add_doctors_to_list()
                self.Add_Data_To_history(5, 6)
                self.History()
        else:
            QMessageBox.information(self, '', 'هذا الطبيب غير موجود')

    def Show_Doctor_Data(self):
        global Delete_Doctor
        global Edit_Doctor
        self.cur.execute(''' select name,genus from doctor where name=%s ''', (self.comboBox_28.currentText(),))
        data = self.cur.fetchone()
        if data:
            genus = data[1]
            if genus == 'male':
                self.comboBox_27.setCurrentIndex(1)
            else:
                self.comboBox_27.setCurrentIndex(0)
            self.lineEdit_6.setText(data[0])
            if Edit_Doctor:
                self.pushButton_48.setEnabled(True)
            else:
                self.pushButton_48.setEnabled(False)
            if Delete_Doctor:
                self.pushButton_49.setEnabled(True)
            else:
                self.pushButton_49.setEnabled(False)
        else:
            self.pushButton_48.setEnabled(False)
            self.pushButton_49.setEnabled(False)
            self.comboBox_27.setCurrentIndex(0)
            self.lineEdit_6.setText('')

    def mr(self):
        print('finaly')

    def Show_client_data_by_current_index(self):
        self.cur.execute(
            '''SELECT client_name,doctor_name,client_age,genus,notes,client_id FROM addnewitem WHERE client_name = %s''',
            (self.comboBox_4.currentText(),))
        analyst_data = self.cur.fetchone()
        if analyst_data:
            self.spinBox.setValue(int(analyst_data[5]))
            self.comboBox_15.setCurrentText(str(analyst_data[1]))
            self.comboBox_15.setEnabled(False)
            self.comboBox_14.setEnabled(False)
            self.comboBox_4.setCurrentText(analyst_data[0])
            self.spinBox_7.setValue(int(analyst_data[2]))
            self.spinBox_7.setEnabled(False)
            if analyst_data[3] == 'ذكر':
                self.comboBox_14.setCurrentIndex(1)
                self.comboBox_14.setEnabled(False)
            if analyst_data[3] == 'انثى':
                self.comboBox_14.setCurrentIndex(0)
                self.comboBox_14.setEnabled(False)
            self.textEdit.setPlainText(str(analyst_data[4]))
            self.textEdit.setEnabled(False)
        else:
            self.spinBox.setValue(0)
            self.comboBox_15.setCurrentText('')
            self.comboBox_15.setEnabled(True)
            self.comboBox_14.setEnabled(True)
            self.spinBox_7.setValue(20)
            self.spinBox_7.setEnabled(True)
            self.comboBox_14.setCurrentIndex(0)
            self.comboBox_14.setEnabled(True)
            self.textEdit.setPlainText('')
            self.textEdit.setEnabled(True)
        name = self.comboBox_4.currentText()
        if name != '' or name != None:
            self.pushButton_44.setEnabled(True)
            self.pushButton_17.setEnabled(True)
            self.comboBox_16.setEnabled(True)
        if name == '' or name == None:
            self.pushButton_44.setEnabled(False)
            self.pushButton_17.setEnabled(False)
            self.comboBox_16.setEnabled(False)

    def echo_mode(self):
        global echo_mode_num
        if echo_mode_num % 2 == 0:
            self.lineEdit_17.setEchoMode(QLineEdit.Normal)
        else:
            self.lineEdit_17.setEchoMode(QLineEdit.Password)
        echo_mode_num += 1

    def echo_mode2(self):
        global echo_mode_num2
        if echo_mode_num2 % 2 == 0:
            self.lineEdit_2.setEchoMode(QLineEdit.Normal)
        else:
            self.lineEdit_2.setEchoMode(QLineEdit.Password)
        echo_mode_num2 += 1

    def add_all_subCategory_toList(self):
        global mylist
        mylist.clear()
        self.cur.execute(''' SELECT sub_category FROM addanalyst ''')
        data = self.cur.fetchall()
        for i in data:
            if i[0] not in mylist:
                mylist.append(i[0])

    def Set_Chick_State2(self, item):
        nitem = self.comboBox_16.model().itemFromIndex(item)
        ritem = self.comboBox_16.model().itemFromIndex(item)
        self.comboBox_16.setCurrentText(ritem.text())
        self.Chick_analyst_category(item)
        # if nitem.checkState() == Qt.Checked:
        #     nitem.setCheckState(Qt.Unchecked)
        # else:
        #     # print('ggggggg')
        #     nitem.setCheckState(Qt.Checked)

    def Set_Chick_State(self, item=None):
        try:
            if item.checkState() == Qt.Checked:
                item.setCheckState(Qt.Unchecked)
            else:
                item.setCheckState(Qt.Checked)

        except Exception as e:
            print(e, '95poerorr')

    def setCheckSateForAllItems(self, table=None, ListWidget=None):
        # print(table,item)
        # MyObject = self.Dialog.findChild(QListWidget, str(table))
        # print(MyObject)

        if table.text() == 'الكل':
            if table.checkState() == Qt.Checked:
                for row in range(0, ListWidget.count()):
                    row_item = ListWidget.item(row)
                    row_item.setCheckState(Qt.Checked)
            else:
                for row in range(0, ListWidget.count()):
                    row_item = ListWidget.item(row)
                    row_item.setCheckState(Qt.Unchecked)

    def Add_all_analysts_items(self):
        global Add_all_analysts_items_list
        Add_all_analysts_items_list.clear()
        for row in range(0, self.tableWidget_5.rowCount() - 1):
            Add_all_analysts_items_list.append(str(self.tableWidget_5.item(row, 1).text()))

    def Show_multy_Dialog(self):
        self.Add_all_analysts_items()
        global Add_all_analysts_items_list
        self.Dialog = multyclass.MultyDialog()
        self.cur.execute(''' SELECT name,sub_category FROM addanalyst''')
        data = self.cur.fetchall()
        my_list = []
        for index1, i in enumerate(data):
            if data[index1][1] not in my_list:
                my_list.append(data[index1][1])
        all_listWidget = []
        for index, ii in enumerate(my_list):
            tab = QWidget()
            self.Dialog.pushButton.setText('اضافة')
            self.Dialog.tabWidget.addTab(tab, str(ii))
            self.Dialog.my_listWidget = QListWidget(tab)
            self.Dialog.my_listWidget.setObjectName("listWidget_+" + str(index + 1))
            all_listWidget.append(str("listWidget_+" + str(index + 1)))
            self.Dialog.my_listWidget.setGeometry(QRect(0, 10, 641, 491))
            font = QFont()
            font.setPointSize(11)
            self.Dialog.my_listWidget.setFont(font)
            item = QListWidgetItem()
            item.setFlags(item.flags() | Qt.ItemIsUserCheckable)
            item.setCheckState(Qt.Checked)
        for index2 in range(0, self.Dialog.tabWidget.count()):
            if index2 < self.Dialog.tabWidget.count():
                listWidget_object = self.Dialog.findChild(QListWidget, f"listWidget_+{index2 + 1}")
                listWidget_object.itemClicked.connect(self.Set_Chick_State)
                item2 = QListWidgetItem()
                item2.setFlags(item2.flags() | Qt.ItemIsUserCheckable)
                item2.setCheckState(Qt.Unchecked)
                item2.setText('الكل')
                listWidget_object.itemClicked.connect((lambda arg=listWidget_object,
                                                              arg2=listWidget_object: self.setCheckSateForAllItems(
                    table=arg, ListWidget=arg2)))
                listWidget_object.addItem(item2)
                for iy in range(0, len(data)):
                    if self.Dialog.tabWidget.tabText(index2) == str(data[iy][1]):
                        item = QListWidgetItem()
                        item.setFlags(item.flags() | Qt.ItemIsUserCheckable)
                        item.setCheckState(Qt.Unchecked)
                        item.setText(str(data[iy][0]))
                        listWidget_object.addItem(item)
                        if item.text() in Add_all_analysts_items_list:
                            item.setFlags(Qt.NoItemFlags)
                        else:
                            item.setFlags(Qt.ItemIsEnabled | Qt.ItemIsSelectable | Qt.ItemIsUserCheckable)
            else:
                # print('else')
                break
        # print(self.Dialog.tabWidget.count(), 'for sorry')
        self.Dialog.pushButton.clicked.connect(self.Handel_multy_Dialog)
        self.Dialog.exec_()

    def Handel_multy_Dialog(self):
        for i in range(0, self.Dialog.tabWidget.count()):
            # print('ghg')
            listWidget_object = self.Dialog.findChild(QListWidget, f"listWidget_+{i + 1}")
            for row in range(0, listWidget_object.count()):
                if listWidget_object.item(row).checkState() == Qt.Checked:
                    self.comboBox_16.setCurrentText(str(listWidget_object.item(row).text()))
                    self.Sales_Page()
                    # print('1')
        self.Dialog.close()
        self.Show_All_one_client_analyst()
        self.Add_buttons_combo_spin_to_tableWidget()
        self.tableWidget_5.scrollToBottom()
        self.get_total_price()
        self.my_def2()
        self.Add_Data_To_history(3, 1)
        self.History()
        self.Add_all_analysts_items()

    def Show_search_Widget(self):
        # print('yesf')
        self.searchWidget = searchDialog.Dialog()
        self.searchWidget.spinBox.setValue(self.spinBox.value())
        self.searchWidget.show()
        self.searchWidget.dateEdit_6.setDate(datetime.date.today())
        self.searchWidget.dateEdit_5.setDate(datetime.date.today())
        self.searchWidget.pushButton_20.clicked.connect(self.Search_by_date)

    def Show_search_Widget2(self):
        global search_info_by_date
        search_info_by_date = True
        # print('yesf')
        self.searchWidget2 = searchDialog2.Dialog()
        self.searchWidget2.spinBox.setValue(self.spinBox_2.value())
        self.searchWidget2.show()
        self.searchWidget2.dateEdit_6.setDate(datetime.date.today())
        self.searchWidget2.dateEdit_5.setDate(datetime.date.today())
        self.searchWidget2.pushButton_20.clicked.connect(self.Clients_Page)

    def Search_by_date(self):
        # print('clicked')
        global select_by_date
        select_by_date = True
        self.Show_All_one_client_analyst()
        self.searchWidget.close()

    def add_clients_to_combo(self):
        clients = []
        self.cur.execute(''' SELECT client_name FROM addclient''')
        data = self.cur.fetchall()
        for row in data:
            if row[0] not in clients:
                clients.append(row[0])
        self.comboBox_4.clear()
        self.comboBox_4.addItem('')
        self.comboBox_4.addItems(clients)

    def ADD_one_category_to_analyst_categoryes_IN_EDITMODE(self):
        global first_text_category1
        first_text_category1 = self.comboBox_26.currentText()
        all_items = []
        for row in range(0, self.comboBox_26.count()):
            self.comboBox_30.setCurrentIndex(row)
            all_items.append(str(self.comboBox_26.currentText()))
        if str(first_text_category1) not in all_items and first_text_category1 != '':
            self.comboBox_26.addItem(str(first_text_category1))
            self.comboBox_26.setCurrentText('')
        else:
            self.comboBox_26.setCurrentText('')
            QMessageBox.information(self, '', 'هذا العنصر موجود بالفعل')

    def ADD_one_category_to_analyst_categoryes(self):
        global first_text_category2
        first_text_category2 = self.comboBox_23.currentText()
        all_items = []
        for row in range(0, self.comboBox_23.count()):
            self.comboBox_23.setCurrentIndex(row)
            all_items.append(str(self.comboBox_23.currentText()))
        if str(first_text_category2) not in all_items and first_text_category2 != '':
            self.comboBox_23.addItem(str(first_text_category2))
            self.comboBox_23.setCurrentText('')
        else:
            self.comboBox_23.setCurrentText('')
            QMessageBox.information(self, '', 'هذا العنصر موجود بالفعل')

    def all_categorys(self):
        self.cur.execute(''' SELECT sub_category FROM addanalyst ''')
        data = self.cur.fetchall()
        mylist = []
        for i in data:
            if i[0] not in mylist:
                mylist.append(i[0])
        print(mylist)
        return mylist

    def Remove_category(self):
        if self.comboBox_23.currentText() in self.all_categorys():
            warning = QMessageBox.warning(self, '',
                                          f'سوف يتم مسح الصنف {self.comboBox_23.currentText()} وجميع التحاليل الخاصة به هل انت متأكد؟',
                                          QMessageBox.Yes | QMessageBox.No)
            if warning == QMessageBox.Yes:
                self.cur.execute(''' DELETE FROM addanalyst WHERE sub_category=%s ''',
                                 (self.comboBox_23.currentText(),))
                self.db.commit()
                index = self.comboBox_23.findText(self.comboBox_23.currentText())
                self.comboBox_23.removeItem(index)
                QMessageBox.information(self, '',
                                        f'تم مسح الصنف {self.comboBox_23.currentText()} وجميع التحاليل الخاصة به')
                print('end')
        else:
            QMessageBox.information(self, '', 'هذا الصنف غير موجود')

    def Remove_category_in_edit(self):
        if self.comboBox_26.currentText() in self.all_categorys():
            warning = QMessageBox.warning(self, '',
                                          f'سوف يتم مسح الصنف {self.comboBox_26.currentText()} وجميع التحاليل الخاصة به هل انت متأكد؟',
                                          QMessageBox.Yes | QMessageBox.No)
            if warning == QMessageBox.Yes:
                self.cur.execute(''' DELETE FROM addanalyst WHERE sub_category=%s ''',
                                 (self.comboBox_26.currentText(),))
                self.db.commit()
                index = self.comboBox_26.findText(self.comboBox_26.currentText())
                self.comboBox_26.removeItem(index)
                QMessageBox.information(self, '',
                                        f'تم مسح الصنف {self.comboBox_26.currentText()} وجميع التحاليل الخاصة به')
        else:
            QMessageBox.information(self, '', 'هذا الصنف غير موجود')

    def ADD_one_REsult_to_ANalyst_results(self):
        global first_text
        first_text = self.comboBox_30.currentText()
        all_items = []
        for row in range(0, self.comboBox_30.count()):
            self.comboBox_30.setCurrentIndex(row)
            all_items.append(str(self.comboBox_30.currentText()))
        if str(first_text) not in all_items and first_text != '':
            self.comboBox_30.addItem(str(first_text))
            self.comboBox_30.setCurrentText('')
        else:
            self.comboBox_30.setCurrentText('')
            QMessageBox.information(self, '', 'هذا العنصر موجود بالفعل')
        # print(first_text,'here my boy')

    def ADD_one_REsult_to_ANalyst_results_IN_EDITMODE(self):
        global first_text2
        first_text2 = self.comboBox_31.currentText()
        all_items = []
        for row in range(0, self.comboBox_31.count()):
            self.comboBox_31.setCurrentIndex(row)
            all_items.append(str(self.comboBox_31.currentText()))
        if str(first_text2) not in all_items and first_text2 != '':
            self.comboBox_31.addItem(str(first_text2))
            self.comboBox_31.setCurrentText('')
        else:
            self.comboBox_31.setCurrentText('')
            QMessageBox.information(self, '', 'هذا العنصر موجود بالفعل')

    def Remove_one_REsult_to_ANalyst_results_IN_EDITMODE(self):
        self.comboBox_31.removeItem(self.comboBox_31.currentIndex())
        self.comboBox_31.setCurrentText('')

    def Remove_one_REsult_to_ANalyst_results(self):
        self.comboBox_30.removeItem(self.comboBox_30.currentIndex())
        self.comboBox_30.setCurrentText('')

    def my_def2(self):
        self.cur.execute(''' SELECT * FROM addclient group by client_name,date(date) ''')
        data = self.cur.fetchall()
        all_clients = []
        all_clients_id = []
        for index in range(0, len(data)):
            all_clients.append(data[index][1])
            all_clients_id.append(data[index][0])
        self.tableWidget_2.setSortingEnabled(False)
        self.tableWidget_2.setRowCount(0)
        self.tableWidget_2.insertRow(0)
        for row, form in enumerate(data):
            for col, item in enumerate(form):
                if col == 0:
                    self.tableWidget_2.setItem(row, col,
                                               QTableWidgetItem(str(all_clients_id[all_clients.index(data[row][1])])))
                elif col == 3:
                    self.tableWidget_2.setItem(row, col, QTableWidgetItem(str(data[row][4])))
                elif col == 4:
                    self.tableWidget_2.setItem(row, col, QTableWidgetItem(str(data[row][5])[:10]))
                else:
                    self.tableWidget_2.setItem(row, col, QTableWidgetItem(str(item)))
                col += 1
            row_pos = self.tableWidget_2.rowCount()
            self.tableWidget_2.insertRow(row_pos)
        self.tableWidget_2.setSortingEnabled(True)
    def add_Analyst_to_list(self):
        global analysts_name_glo
        self.cur.execute(''' SELECT name FROM addanalyst ''')
        data = self.cur.fetchall()
        analysts_name_glo.clear()
        for i in data:
            if i[0] not in analysts_name_glo:
                analysts_name_glo.append(i[0])

    def Auto_Complete(self, model):
        model.setStringList(analysts_name_glo)

    def please(self):
        print('run2')
        self.comboBox_16.setCurrentText('pleas1')
        self.comboBox_21.setCurrentText('pleas1')

    def Auto_complete_combo(self):
        combo = self.comboBox_16, self.comboBox_21
        completer = QCompleter()
        for i in combo:
            i.setCompleter(completer)
            model = QStringListModel()
            completer.setModel(model)
            completer.setCaseSensitivity(Qt.CaseInsensitive)
            self.Auto_Complete(model)
        completer.activated.connect(self.please)

    def add_client_to_list(self):
        global clients_name_glo
        clients_name_glo.clear()
        self.cur.execute(''' SELECT client_name FROM addclient  ''')
        data = self.cur.fetchall()
        for i in data:
            if i[0] not in clients_name_glo:
                clients_name_glo.append(i[0])

    def Auto_Complete2(self, model):
        model.setStringList(clients_name_glo)

    def Auto_complete_combo2(self):
        combo = self.lineEdit_25, self.lineEdit_27
        completer = QCompleter()
        for i in combo:
            i.setCompleter(completer)
            model = QStringListModel()
            completer.setModel(model)
            completer.setCaseSensitivity(Qt.CaseInsensitive)
            self.Auto_Complete2(model)

    def add_client_to_list4(self):
        self.cur.execute(''' SELECT client_name FROM addclient''')
        data = self.cur.fetchall()
        for i in data:
            if i[0] not in clients_name_glo_clients_page:
                clients_name_glo_clients_page.append(i[0])

    def Auto_Complete4(self, model):
        model.setStringList(clients_name_glo_clients_page)

    def Auto_complete_combo4(self):
        combo = self.lineEdit_27
        completer = QCompleter()
        combo.setCompleter(completer)
        model = QStringListModel()
        completer.setModel(model)
        completer.setCaseSensitivity(Qt.CaseInsensitive)
        self.Auto_Complete4(model)

    def Delete_Files(self):
        self.cur.execute(''' SELECT * FROM paths WHERE id=1 ''')
        mydata = self.cur.fetchone()
        word_files = mydata[1]
        save_word_files = mydata[2]
        try:
            if os.path.exists(r'%s\result.docx' % save_word_files):
                os.remove(r'%s\result.docx' % save_word_files)
            if os.path.exists(r'%s\result2.docx' % save_word_files):
                os.remove(r'%s\result2.docx' % save_word_files)
            if os.path.exists(r'%s\result3.docx' % save_word_files):
                os.remove(r'%s\result3.docx' % save_word_files)
        except Exception as e:
            word = client.Dispatch("Word.Application")
            word.ActiveDocument.Close()
            self.Delete_Files()
            print(e, '7erorr')

    def Show_paths(self):
        self.cur.execute(''' SELECT * FROM paths WHERE id=1 ''')
        mydata = self.cur.fetchone()
        word_files = mydata[1]
        save_word_files = mydata[2]
        self.lineEdit_14.setText(word_files)
        self.lineEdit_15.setText(save_word_files)

    def Add_Path(self):
        files_path = self.lineEdit_14.text()
        save_files_path = self.lineEdit_15.text()
        self.cur.execute(''' UPDATE paths SET file_path=%s,save_file_path=%s WHERE id=1''',
                         (files_path, save_files_path))
        self.db.commit()
        QMessageBox.information(self, '', 'تم تطبيق المعلومات بنجاح')
        self.Show_paths()

    def Preview(self):
        global if_print
        self.Print_Sale_Data('T')
        self.cur.execute(''' SELECT * FROM paths WHERE id=1 ''')
        mydata = self.cur.fetchone()
        word_files = mydata[1]
        save_word_files = mydata[2]
        warning = QMessageBox.warning(self, '', 'هل انتهيت من المعاينة؟', QMessageBox.Yes | QMessageBox.No)
        try:
            if_print = True
            if os.path.exists(r'%s\result.docx' % save_word_files):
                os.remove(r'%s\result.docx' % save_word_files)
            if os.path.exists(r'%s\result2.docx' % save_word_files):
                os.remove(r'%s\result2.docx' % save_word_files)
            if os.path.exists(r'%s\result3.docx' % save_word_files):
                os.remove(r'%s\result3.docx' % save_word_files)
        except:
            word = client.Dispatch("Word.Application")
            try:
                word.ActiveDocument.Close()
            except:
                try:
                    os.system("TASKKILL /F /IM WINWORD.exe")
                    os.system('start WINWORD.exe')
                except:
                    QMessageBox.information(self, '', 'يرجى اغلاق برنامج MS Word')
            self.Delete_Files()

    def Reset_password(self):
        try:
            user_name = self.lineEdit_7.text()
            self.cur.execute(''' SELECT * FROM adduser ''')
            data = self.cur.fetchall()
            ruser_name = ''
            a = 0
            for row in data:
                if row[1] == user_name:
                    ruser_name = row[1]
                else:
                    a = 5
            if a == 5:
                QMessageBox.information(self, 'info', 'اسم المستخدم الذي ادخلته غير صحيح')
            self.cur.execute(''' SELECT user_email,user_password FROM adduser WHERE user_name=%s ''', (ruser_name,))
            email_data = self.cur.fetchone()
            email = "ameersaad810@gmail.com"
            password = "aahmpredtiddvxlo"
            send_to_email = email_data[0]
            subject = "n"
            message = f'Hello.\n your password is {email_data[1]} '
            msg = MIMEMultipart()
            msg["From"] = email
            msg["To"] = send_to_email
            msg["Subject"] = subject
            msg.attach(MIMEText(message, 'plain'))
            server = smtplib.SMTP("smtp.gmail.com", 587)
            server.starttls()
            server.login(email, password)
            text = msg.as_string()
            server.sendmail(email, send_to_email, text)
            server.quit()
            print('ok')
            QMessageBox.information(self, 'تم بنجاح', f'تم ارسال كلمة المرور الى{send_to_email} ')
        except Exception as e:
            print(e, '97erorr')
            QMessageBox.information(self, 'خطأ', f'هنالك خطأ يرجى\n ارسال هذا النص {e} الى ameersaad810@gmail.com')

    def clear_data_in_sales(self):
        self.Update_addNewItem_Data()
        self.comboBox_17.clear()
        self.tableWidget_5.setRowCount(0)
        self.tableWidget_5.insertRow(0)
        self.comboBox_4.clear()
        self.lineEdit_21.setText('')
        self.spinBox_7.setValue(20)
        self.doubleSpinBox_7.setValue(0)
        self.comboBox_14.setCurrentIndex(0)
        self.comboBox_16.setCurrentIndex(0)
        self.comboBox_17.setCurrentIndex(0)
        self.comboBox_15.setEnabled(True)
        self.comboBox_15.setCurrentIndex(0)
        self.textEdit.setPlainText('')
        self.comboBox_4.setEnabled(True)
        self.spinBox_7.setEnabled(True)
        self.comboBox_14.setEnabled(True)
        self.textEdit.setEnabled(True)
        self.spinBox.setValue(0)
        self.add_clients_to_combo()

    def Show_All_Clients(self):
        global show_all_sales_in_clients_page
        global show_clients_check
        if show_clients_check:
            client_name = self.lineEdit_27.text()
            if client_name != '0' and client_name != '':
                self.cur.execute(''' SELECT * FROM addclient where client_name=%s group by client_name,date(date) ''',
                                 (client_name,))
            data = self.cur.fetchall()
            all_clients = []
            all_clients_id = []
            for index in range(0, len(data)):
                all_clients.append(data[index][1])
                all_clients_id.append(data[index][0])
            self.tableWidget_2.setSortingEnabled(False)
            self.tableWidget_2.setRowCount(0)
            self.tableWidget_2.insertRow(0)
            try:
                if client_name != '0' and client_name != '':
                    for row, form in enumerate(data):
                        for col, item in enumerate(form):
                            if col == 0:
                                self.tableWidget_2.setItem(row, col, QTableWidgetItem(
                                    str(all_clients_id[all_clients.index(data[row][1])])))
                            elif col == 3:
                                self.tableWidget_2.setItem(row, col, QTableWidgetItem(str(data[row][4])))
                            elif col == 4:
                                self.tableWidget_2.setItem(row, col, QTableWidgetItem(str(data[row][5])[:10]))
                            else:
                                self.tableWidget_2.setItem(row, col, QTableWidgetItem(str(item)))
                            col += 1
                        row_pos = self.tableWidget_2.rowCount()
                        self.tableWidget_2.insertRow(row_pos)
                else:
                    show_all_sales_in_clients_page = True
                    self.my_def2()
            except Exception as e:
                print(e, '9gy5erorr')
            self.tableWidget_2.setSortingEnabled(True)
        else:
            self.tableWidget_2.setRowCount(0)
            self.tableWidget_2.insertRow(0)
    def Search_In_History(self):
        actionsd = self.comboBox_24.currentIndex()
        tabley = self.comboBox_20.currentIndex()
        if actionsd != 0 and tabley == 0:
            try:
                self.cur.execute(f'SELECT uid,action,tabled,dates FROM his WHERE action = {actionsd} ORDER BY -dates')
            except Exception as e:
                print(e, '8erorr')
        elif tabley != 0 and actionsd == 0:
            try:
                self.cur.execute(f'SELECT uid,action,tabled,dates FROM his WHERE tabled={tabley} ORDER BY -dates')
            except Exception as e:
                print(e, '9erorr')
        elif actionsd != 0 and tabley != 0:
            try:
                self.cur.execute(
                    f' SELECT uid,action,tabled,dates FROM his WHERE action={actionsd} AND tabled={tabley} ORDER BY -dates')

            except Exception as e:
                print(e, '10erorr')
        else:
            try:
                self.cur.execute(
                    f' SELECT uid,action,tabled,dates FROM his ORDER BY -dates')
            except Exception as e:
                print(e, '11erorr')
        data = self.cur.fetchall()
        self.tableWidget_8.setRowCount(0)
        self.tableWidget_8.insertRow(0)
        for row, form in enumerate(data):
            for col, item in enumerate(form):
                if col == 0:
                    self.tableWidget_8.setItem(row, col, QTableWidgetItem(str(data[row][0])))
                if col == 1:
                    action = ''
                    if data[row][1] == 1:
                        action = 'تسجيل الدخول'
                    if data[row][1] == 2:
                        action = 'تسجيل الخروج'
                    if data[row][1] == 3:
                        action = 'اضافة'
                    if data[row][1] == 4:
                        action = 'تعديل'
                    if data[row][1] == 5:
                        action = 'حذف'
                    if data[row][1] == 6:
                        action = 'بحث'
                    if data[row][1] == 7:
                        action = 'طباعة'
                    self.tableWidget_8.setItem(row, col, QTableWidgetItem(str(action)))
                if col == 2:
                    tables = ''
                    if data[row][2] == 1:
                        tables = 'مبيع يومي'
                    if data[row][2] == 2:
                        tables = 'تحليل'
                    if data[row][2] == 3:
                        tables = 'مشتريات'
                    if data[row][2] == 4:
                        tables = 'مراجعين'
                    if data[row][2] == 5:
                        tables = 'مستخدم'
                    self.tableWidget_8.setItem(row, col, QTableWidgetItem(str(tables)))
                if col == 3:
                    self.tableWidget_8.setItem(row, col, QTableWidgetItem(str(data[row][3])))
                col += 1
            row_pos = self.tableWidget_8.rowCount()
            self.tableWidget_8.insertRow(row_pos)

    def Search_In_All_Sales(self):
        client_name = self.lineEdit_25.text()
        if client_name != '0':
            self.cur.execute(''' SELECT * FROM addclient WHERE client_name=%s AND DATE(date)=%s ORDER BY id''',
                             (client_name, datetime.date.today(),))
        else:
            self.cur.execute(''' SELECT * FROM addclient WHERE DATE(date)=%s ORDER BY -date''',
                             (datetime.date.today(),))
        analyst_data = self.cur.fetchall()
        self.tableWidget_6.setSortingEnabled(True)
        self.tableWidget_6.setRowCount(0)
        self.tableWidget_6.insertRow(0)
        try:
            if client_name != '0' and client_name != '':
                for i, k in enumerate(analyst_data[0]):
                    if i == 3:
                        self.tableWidget_6.setItem(0, i, QTableWidgetItem(str(analyst_data[0][4])))
                    else:
                        self.tableWidget_6.setItem(0, i, QTableWidgetItem(str(k)))
            else:
                self.Show_All_The_Sales()
        except Exception as e:
            print(e, '95erlplorr')
            self.tableWidget_6.setRowCount(0)
            self.tableWidget_6.insertRow(0)
        self.Add_Data_To_history(6, 1)
        self.tableWidget_6.setSortingEnabled(True)

    def Print_Sale_Data(self, prev):
        genuses = self.comboBox_14.currentIndex()
        all_analyst = []
        all_result = []
        date = datetime.datetime.now()
        day = date.year
        month = date.month
        year = date.day
        real_name = ''
        real_doctor = ''
        categorys = []
        for row in range(0, self.tableWidget_5.rowCount() - 1):
            categorys.append(self.tableWidget_5.item(row, 6).text())
            try:
                real_name = self.tableWidget_5.item(row, 0).text()
                analyst = self.tableWidget_5.item(row, 1).text()
                all_analyst.append(analyst)
            except Exception as e:
                print(e, '95ehiororr')
            try:
                result = self.tableWidget_5.cellWidget(row, 2).value()
            except:
                try:
                    result = self.tableWidget_5.cellWidget(row, 2).currentText()
                except:
                    try:
                        result = self.tableWidget_5.item(row, 2).text()
                    except Exception as e:
                        print(e, '95ek;okpororr')

            all_result.append(str(result))
            try:
                real_doctor = self.tableWidget_5.item(row, 3).text()
            except Exception as e:
                print(e, '95ejijijii0rorr')
        # try:
        self.Bio_Word(real_name, real_doctor, all_analyst, all_result, year, month, day, prev, genuses, categorys)
        # except Exception as e:
        #     print(e, '12erorr')
        #     QMessageBox.information(self, 'خطأ',
        #                             f'هنالك خطأ يرجى مراجعة العملية او\n ارسال هذا النص {e} الى ameersaad810@gmail.com')

        if prev != 'T':
            self.Delete_Files()
        self.Update_addNewItem_Data()

    def Delete_Row(self, item):
        try:
            if self.tableWidget_5.rowCount() > 2:

                analyst_name = self.tableWidget_5.item(item, 1).text()
                client_name = self.comboBox_4.currentText()
                warning = QMessageBox.warning(self, 'احذر', f"سوف يتم مسح العنصر{analyst_name} هل انت متأكد؟",
                                              QMessageBox.Yes | QMessageBox.No)
                if warning == QMessageBox.Yes:
                    self.tableWidget_5.removeRow(item)
                    self.cur.execute(''' select MIN(id) from addclient WHERE client_name=%s''',
                                     (client_name,))
                    min_id = self.cur.fetchone()
                    self.cur.execute(
                        ''' DELETE FROM addnewitem WHERE client_name=%s AND analyst_name=%s AND DATE(date)=%s''',
                        (client_name, analyst_name, str(datetime.date.today()),))
                    self.db.commit()
                    self.cur.execute(''' DELETE FROM addclient WHERE client_name=%s AND id!=%s limit 1''',
                                     (client_name, min_id[0],))
                    self.db.commit()
            else:
                analyst_name = self.tableWidget_5.item(item, 1).text()
                QMessageBox.information(self, 'تحذير',
                                        f'لا يمكن حذف كل العناصر اضف عنصر جديد لكي تستطيع حذف{analyst_name}')
        except Exception as e:
            print(e, '13erorr')

    def Delete_Row2(self, item):
        self.tableWidget_3.setSortingEnabled(False)
        try:
            if self.tableWidget_3.rowCount() > 1:
                item_name = self.tableWidget_3.item(item, 0).text()
                item_quantity = self.tableWidget_3.item(item, 1).text()
                item_price = self.tableWidget_3.item(item, 2).text()
                warning = QMessageBox.warning(self, 'احذر', f"سوف يتم مسح العنصر{item_name} هل انت متأكد؟",
                                              QMessageBox.Yes | QMessageBox.No)
                if warning == QMessageBox.Yes:
                    self.cur.execute(
                        ''' DELETE FROM addbuys WHERE item_name=%s AND signal_item_price=%s AND quantity=%s''',
                        (item_name, int(item_price), int(item_quantity),))
                    self.db.commit()
                    self.Show_all_buys()
        except Exception as e:
            print(e, '95erorrlddeee')
        self.tableWidget_3.setSortingEnabled(True)
    def Add_Doctor_Data(self):
        self.comboBox_28.clear()
        self.comboBox_15.clear()
        self.comboBox_28.addItem('---------------')
        self.cur.execute(''' select * from doctor ''')
        data = self.cur.fetchall()
        for i in range(0, len(data)):
            self.comboBox_28.addItem(str(data[i][1]))
            self.comboBox_15.addItem(str(data[i][1]))

    def buttonClicked(self):
        button = self.sender()
        index = self.tableWidget_5.indexAt(button.pos())
        self.Delete_Row(index.row())

    def buttonClicked2(self):
        button = self.sender()
        self.tableWidget_3.setSortingEnabled(False)
        index = self.tableWidget_3.indexAt(button.pos())
        self.Delete_Row2(index.row())
        self.tableWidget_3.setSortingEnabled(True)

    def get_total_price(self):
        total_price = 0
        for row in range(0, self.tableWidget_5.rowCount() - 1):
            try:
                analyst_name = self.tableWidget_5.item(row, 1).text()
            except:
                analyst_name = ''

            try:
                a = self.tableWidget_5.item(row, 4).text()
            except:
                a = 0
            try:
                total_price += int(a)
            except ValueError:
                a = 0
        self.lineEdit_24.setText(str(total_price))

    def Update_addNewItem_Data(self):
        for rowj in range(0, self.tableWidget_5.rowCount() - 1):
            r2_doctor = None
            r2_analyst_name = None
            r2_client_name = None
            price=self.tableWidget_5.item(rowj, 4).text()
            try:
                r2_doctor = self.tableWidget_5.item(rowj, 3).text()
                r2_analyst_name = self.tableWidget_5.item(rowj, 1).text()
                r2_client_name = self.comboBox_4.currentText()
            except:
                r2_client_name = ''
                r2_analyst_name = ''
                r2_doctor = ''
            r2_result = None
            try:
                r2_result = self.tableWidget_5.cellWidget(rowj, 2).value()
            except:
                try:
                    r2_result = self.tableWidget_5.cellWidget(rowj, 2).currentText()
                except:
                    try:
                        r2_result = self.tableWidget_5.cellWidget(rowj, 2).text()
                    except:
                        try:
                            r2_result = self.tableWidget_5.item(rowj, 2).text()
                        except Exception as e:
                            print(e, '99erorr')
                            r2_result = ''
            try:
                self.cur.execute(
                    ''' UPDATE addnewitem SET  analyst_name=%s ,analyst_result=%s,price=%s WHERE client_name=%s AND analyst_name=%s''',
                    (r2_analyst_name, r2_result,price, r2_client_name, r2_analyst_name))
                self.db.commit()
            except:
                print('except')

    # def show_sales_items(self):
    #     self.cur.execute(
    #         ''' SELECT client_name,analyst_name,analyst_result,doctor_name,total_price,sub_category FROM addnewitem WHERE client_name = %s AND DATE(date)=%s''',
    #         (self.comboBox_4.currentText(), datetime.date.today(),))
    #     analyst_data = self.cur.fetchall()
    #     self.tableWidget_5.setRowCount(0)
    #     self.tableWidget_5.insertRow(0)
    #     for row, form in enumerate(analyst_data):
    #         for col, item in enumerate(form):
    #             if col == 1 or col == 0 or col == 3:
    #                 self.tableWidget_5.setItem(row, col, QTableWidgetItem(str(item)))
    #                 my_item = self.tableWidget_5.item(row, col)
    #                 my_item.setFlags(Qt.ItemIsEditable)
    #             else:
    #                 self.tableWidget_5.setItem(row, col, QTableWidgetItem(str(item)))
    #             col += 1
    #         row_pos = self.tableWidget_5.rowCount()
    #         self.tableWidget_5.insertRow(row_pos)
    #     self.tableWidget_5.sortItems(6, Qt.AscendingOrder)
    #     self.tableWidget_5.scrollToBottom()
    #     self.Show_All_The_Sales()
    def Analyst_sub_category(self,analyst_name):
        self.cur.execute(''' select sub_category,price from addanalyst where name=%s''',(analyst_name,))
        data=self.cur.fetchone()
        return data
    def Sales_Page(self, for_loop2=None):
        # print(for_loop2, 'jdijei')
        global clients_name_glo
        not_in_all_analysts_items = False
        for mjrow in range(0, self.tableWidget_5.rowCount() - 1):
            if self.comboBox_16.currentText() == str(self.tableWidget_5.item(mjrow, 1).text()):
                not_in_all_analysts_items = True
        # if self.comboBox_4.currentText()=='':
        #     QMessageBox.information(self,'تحذير','يرجى ادخال اسم المراجع')
        if not not_in_all_analysts_items:
            self.comboBox_4.setEnabled(False)
            self.spinBox_7.setEnabled(False)
            self.comboBox_14.setEnabled(False)
            self.comboBox_15.setEnabled(False)
            self.textEdit.setEnabled(False)
            global client_id_glob
            global chick_if_add_new
            global clients_name_glo
            analyst_name = self.comboBox_16.currentText()
            client_name = self.comboBox_4.currentText()
            client_age = self.spinBox_7.value()
            client_genus = self.comboBox_14.currentText()
            client_doctor = self.comboBox_15.currentText()
            client_genus = self.comboBox_14.currentText()
            analyst_or_clients_notes = self.textEdit.toPlainText()
            analyst_lineEdit_result = self.lineEdit_21.text()
            analyst_combo_result = self.comboBox_17.currentText()
            analyst_number_result = self.doubleSpinBox_7.value()
            client_id = self.spinBox.value()
            self.cur.execute('''SELECT price,category,sub_category FROM addanalyst WHERE name = %s''', (analyst_name,))
            analyst_price = self.cur.fetchone()
            latest_result = 1
            total_price = 0
            if analyst_price != None:
                if analyst_price[1] == 'عدد':
                    latest_result = analyst_number_result
                if analyst_price[1] == 'خيارات':
                    latest_result = analyst_combo_result
                if analyst_price[1] == 'حقل كتابة':
                    latest_result = analyst_lineEdit_result
                if analyst_price[1] == 'خيارات مع تعديل':
                    latest_result = analyst_combo_result
                total_price = int(analyst_price[0])
                if client_name not in clients_name_glo:
                    self.cur.execute('''
                               INSERT INTO addclient (client_name,client_age,client_genus,client_doctor,date)
                               VALUES (%s,%s,%s,%s,%s)
                            ''', (
                        str(client_name), str(client_age), str(client_genus), str(client_doctor),
                        str(datetime.datetime.now())))
                    self.db.commit()
                self.cur.execute('''SELECT id FROM addclient WHERE client_name = %s''', (client_name,))
                real_client_id = self.cur.fetchone()
                self.cur.execute('''
                   INSERT INTO addnewitem (client_name,client_id,client_age,genus,doctor_name,notes,analyst_name,analyst_result,price,total_price,date,sub_category)
                   VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
                ''', (
                    client_name, real_client_id, client_age, client_genus, client_doctor, analyst_or_clients_notes,
                    analyst_name,
                    latest_result,
                    total_price, total_price, datetime.datetime.now(), analyst_price[2],))
                self.db.commit()
                clients_name_glo.append(str(client_name))
                if for_loop2 == False:
                    self.Show_All_one_client_analyst()
                    self.Update_addNewItem_Data()
            else:
                if for_loop2 == False:
                    QMessageBox.information(self, '', 'هذا التحليل غبر موجود')
            if for_loop2 == False:
                chick_if_add_new = True
                self.Add_buttons_combo_spin_to_tableWidget()
                self.tableWidget_5.scrollToBottom()
                self.get_total_price()
                self.my_def2()
                self.Add_Data_To_history(3, 1)
                self.History()
            else:
                print('gnm2')
        else:
            # print('elsdokod')
            # print(for_loop2,'here')
            if for_loop2 == False:
                # print('jj')
                QMessageBox.information(self, 'تحذير', 'هذا العنصر موجود بالفعل')
                self.Add_all_analysts_items()
            else:
                pass
        not_in_all_analysts_items = False
        self.add_client_to_list()
    def Add_buttons_combo_spin_to_tableWidget(self):
        try:
            pass
            # axd1 = self.tableWidget_5.findItems('Fatty drop:GSE', Qt.MatchContains)
            # axd2 = self.tableWidget_5.findItems('Mucuse:GUE', Qt.MatchContains)
            # axd3 = self.tableWidget_5.findItems('Morphology:Pus cells', Qt.MatchContains)
            # self.tableWidget_5.insertRow(axd1[0].row() + 1)
            # # self.tableWidget_5.insertRow(axd1[0].row() + 2)
            # self.tableWidget_5.insertRow(axd2[0].row() + 1)
            # # self.tableWidget_5.insertRow(axd2[0].row() + 2)
            # self.tableWidget_5.insertRow(axd3[0].row() + 1)
            # # self.tableWidget_5.insertRow(axd3[0].row() + 2)
        except Exception as e:
            print(e, '96erorr')
        for rowd in range(0, self.tableWidget_5.rowCount() - 1):
            mypush_button = QPushButton(self)
            mypush_button.setText('حذف')
            mypush_button.setStyleSheet('border-radius:12px;')
            mypush_button.clicked.connect(self.buttonClicked)
            try:
                row_analyst_name = self.tableWidget_5.item(rowd, 1).text()
                if row_analyst_name and row_analyst_name != '':
                    self.tableWidget_5.setCellWidget(rowd, 5, mypush_button)
            except Exception as e:
                print(e, '95erorr')
            all_name_items = []
            try:
                name = self.tableWidget_5.item(rowd, 0).text()
            except Exception as e:
                print(e, '94erorr')
            rs_name = ''
            try:
                rs_name = self.tableWidget_5.item(rowd, 1).text()
            except Exception as e:
                print(e, '93erorr')
                rs_name = ''
            self.cur.execute('''SELECT results,category FROM addanalyst WHERE name=%s''', (rs_name,))
            results_data = self.cur.fetchall()
            mycobmbo = None
            object_type = ''
            if results_data:
                if results_data[0][1] == 'خيارات':
                    mycobmbo = QComboBox(self)
                    list_data = str(results_data[0][0]).replace("'", "")
                    list_data = list_data.split(',')
                    mycobmbo.addItems(list_data)
                    object_type = 'خيارات'
                if results_data[0][1] == 'عدد':
                    mycobmbo = QDoubleSpinBox(self)
                    print('jdije')
                    mycobmbo.setMaximum(99999999999999)
                    mycobmbo.setMinimum(-99999999999999)
                    object_type = 'عدد'
                if results_data[0][1] == 'خيارات مع تعديل':
                    mycobmbo = QComboBox(self)
                    mycobmbo.setEditable(True)
                    list_data = str(results_data[0][0]).replace("'", "")
                    list_data = list_data.split(',')
                    mycobmbo.addItems(list_data)
                    object_type = 'خيارات مع تعديل'
                if results_data[0][1] == 'حقل كتابة':
                    mycobmbo = QLineEdit(self)
                    object_type = 'حقل كتابة'
                try:
                    r2_analyst_name = self.tableWidget_5.item(rowd, 1).text()
                    r2_client_name = self.comboBox_4.currentText()
                    self.cur.execute(
                        ''' SELECT  analyst_result  FROM addnewitem WHERE client_name=%s AND analyst_name=%s ''',
                        (r2_client_name, r2_analyst_name))
                    myrs = self.cur.fetchall()
                    if myrs != None:
                        if object_type == 'خيارات' or object_type == 'خيارات مع تعديل':
                            index = mycobmbo.findText(myrs[0][0], Qt.MatchFixedString)
                            mycobmbo.setCurrentIndex(index)
                        if object_type == 'عدد':
                            # print('Tru')
                            if myrs[0][0] and myrs[0][0] != '':
                                mycobmbo.setValue(float(myrs[0][0]))
                        if object_type == 'حقل كتابة':
                            mycobmbo.setText(str(myrs[0][0]))
                    self.tableWidget_5.setItem(rowd, 2, QTableWidgetItem(str('')))
                    self.tableWidget_5.setCellWidget(rowd, 2, mycobmbo)
                    if object_type == 'خيارات' or object_type == 'خيارات مع تعديل':
                        if mycobmbo.currentText() == '' or mycobmbo.currentText() == ' ':
                            mycobmbo.setCurrentIndex(0)
                except Exception as e:
                    print(e, '91erorr')

    def Show_All_one_client_analyst(self):
        global client_id_glob
        global chick_if_add_new
        global select_by_date
        client_name = self.comboBox_4.currentText()
        if self.spinBox.value() == 0 and chick_if_add_new == False:
            self.tableWidget_5.setRowCount(0)
            self.tableWidget_5.insertRow(0)
            self.lineEdit_21.setText('')
            self.spinBox_7.setValue(0)
            self.doubleSpinBox_7.setValue(0)
            self.comboBox_14.setCurrentIndex(0)
            self.comboBox_17.setCurrentIndex(0)
            self.textEdit.setPlainText('')
            self.comboBox_4.setEnabled(True)
            self.spinBox_7.setEnabled(True)
            self.comboBox_14.setEnabled(True)
            self.textEdit.setEnabled(True)
        if True:
            if select_by_date:
                id = self.searchWidget.spinBox.value()
                from_date = self.searchWidget.dateEdit_6.date()
                to_date = self.searchWidget.dateEdit_5.date()
                self.cur.execute(
                    '''SELECT client_name,analyst_name,analyst_result,doctor_name,total_price,sub_category,client_age,genus,notes FROM addnewitem WHERE client_id = %s AND DATE(date)>=%s AND DATE(date)<=%s''',
                    (id, str(from_date.toPyDate()), str(to_date.toPyDate()),))
            else:
                self.cur.execute(
                    '''SELECT client_name,analyst_name,analyst_result,doctor_name,total_price,sub_category,client_age,genus,notes FROM addnewitem WHERE client_id = %s AND DATE(date)=%s''',
                    (self.spinBox.value(), datetime.date.today(),))
            analyst_data = self.cur.fetchall()
            if analyst_data:
                self.comboBox_15.setCurrentText(str(analyst_data[0][3]))
                self.comboBox_15.setEnabled(False)
                self.comboBox_14.setEnabled(False)
                self.comboBox_4.setCurrentText(analyst_data[0][0])
                self.comboBox_4.setEnabled(False)
                self.spinBox_7.setValue(int(analyst_data[0][6]))
                self.spinBox_7.setEnabled(False)
                if analyst_data[0][7] == 'ذكر':
                    self.comboBox_14.setCurrentIndex(1)
                    self.comboBox_14.setEnabled(False)
                if analyst_data[0][7] == 'انثى':
                    self.comboBox_14.setCurrentIndex(0)
                    self.comboBox_14.setEnabled(False)
                self.textEdit.setPlainText(str(analyst_data[0][8]))
                self.textEdit.setEnabled(False)
                self.tableWidget_5.setRowCount(0)
                self.tableWidget_5.insertRow(0)
                for row, form in enumerate(analyst_data):
                    for col, item in enumerate(form):
                        if col < 5:
                            if col == 1 or col == 0 or col == 3:
                                self.tableWidget_5.setItem(row, col, QTableWidgetItem(str(item)))
                                my_item = self.tableWidget_5.item(row, col)
                                my_item.setFlags(Qt.ItemIsEditable)
                            else:
                                self.tableWidget_5.setItem(row, col, QTableWidgetItem(str(item)))
                        if col == 6:
                            self.tableWidget_5.setItem(row, col, QTableWidgetItem(str(analyst_data[row][5])))
                        col += 1
                    row_pos = self.tableWidget_5.rowCount()
                    self.tableWidget_5.insertRow(row_pos)
                self.tableWidget_5.sortItems(6, Qt.AscendingOrder)
                chick_if_add_new = False
                self.Show_All_The_Sales()
                try:
                    pass
                except Exception as e:
                    print(e, '90erorr')
                self.Add_buttons_combo_spin_to_tableWidget()
            else:
                QMessageBox.information(self, 'Error',
                                        'الرقم الذي ادخلته غير صحيح او غير موجود في مبيعات اليوم يرجى ادخال رقم صحيح او مراجعة صفحة "مبيعات اليوم" للتأكد من الرقم')
            # except Exception as e:
            #     print(e,'1erorr')
            self.get_total_price()
            select_by_date = False
        self.Add_all_analysts_items()

    def Show_Type_of_result_category(self):
        global addTrue
        analyst_name = self.comboBox_16.currentText()
        self.cur.execute('''SELECT category,results FROM addanalyst WHERE name = %s''', (analyst_name,))
        analyst_category = self.cur.fetchone()
        try:
            if analyst_category:
                if analyst_category[0] == 'عدد':
                    self.comboBox_17.hide()
                    self.lineEdit_21.hide()
                    self.doubleSpinBox_7.show()
                if analyst_category[0] == 'خيارات':
                    self.lineEdit_21.hide()
                    self.doubleSpinBox_7.hide()
                    self.comboBox_17.show()
                    self.comboBox_17.setEditable(False)
                if analyst_category[0] == 'حقل كتابة':
                    self.lineEdit_21.show()
                    self.doubleSpinBox_7.hide()
                    self.comboBox_17.hide()
                if analyst_category[0] == 'خيارات مع تعديل':
                    self.lineEdit_21.hide()
                    self.doubleSpinBox_7.hide()
                    # self.comboBox_17.clear()
                    self.comboBox_17.show()
                    self.comboBox_17.setEditable(True)

                try:
                    self.comboBox_17.clear()
                    list_data = str(analyst_category[1]).replace("'", "")
                    list_data = list_data.split(',')
                    for rr in range(0, len(list_data)):
                        self.comboBox_17.addItem(list_data[rr])
                except Exception as e:
                    print(e, '959erorr')

                    # print('erorr here')
        except Exception as e:
            print(e, '2erorr')
            if analyst_name not in mylist:
                if not addTrue:
                    QMessageBox.information(self, 'تحذير', "يرجى اختيار تحليل صحيح")
            # print('error here2')

    def analyst_category_function(self, name):
        global mylist
        # print('starthh')
        # pationt_name = self.comboBox_4.currentText()
        # doctor_name = self.comboBox_15.currentText()
        # analyst_name = self.comboBox_16.currentText()

        # print(mylist)
        if name in mylist:
            # print('name in my list')
            # self.loading=LoadingDialog()
            # self.loading.show()
            self.cur.execute('''SELECT name FROM addanalyst WHERE sub_category=%s ''', (str(name),))
            my_DATA = self.cur.fetchall()
            # print(my_DATA)
            for new in my_DATA:
                self.comboBox_16.setCurrentText(str(new[0]))
                self.Sales_Page()
            self.Show_All_one_client_analyst()
            self.Update_addNewItem_Data()
            self.Add_buttons_combo_spin_to_tableWidget()
            self.tableWidget_5.scrollToBottom()
            self.get_total_price()
            self.my_def2()
            self.Add_Data_To_history(3, 1)
            self.History()
            self.Add_all_analysts_items()
            # self.loading.close()
        # colors = ['yallow', 'brown', 'green', 'milk']
        # RBCs_Pus_cells_GUE_GSE = ['1 - 2', '1 - 3', '2 - 3', '2 - 4', '0 - 1', '0 - 2', '3 - 5', '4 - 6', '5 - 6',
        #                           '5 - 7', '6 - 8', '6 - 7', '+', '++', '+++', '++++', 'Full Field']
        # Consistency = ['Solid', 'Liquid', 'Semi solid', 'Semi liquid', 'Mucoid']
        # G_Lembilia_E_Histolytica = ['Cyst', 'Trophozoite']
        # Epith_cells = ['Few', '+', '++', '+++', '++++']
        # Ova = ['Nil']
        # Appearance = ['Turbid', 'Clear']
        # Reaction = ['Acidic', 'Alkaline']
        # Sugar = ['Nil', '+', '++', '+++', 'Trace']
        # Crystals = ['Am.Urate', 'Am.Phosphatase', 'Uric Acid', 'Ca.Oxalate']
        # Casts = ['Granular cast +', 'Granular cast ++', 'Granular cast +++']
        # Blood_Group = ['A (+ve)', 'B (+ve)', 'AB (+ve)', 'O (+ve)', 'O (-ve)', 'A (-ve)', 'B (-ve)', 'AB (-ve)']
        # Pregnancy_test_in_urine_serum = ['Positive (+ve)', 'Negative (-ve)', 'Weak Positive']
        # Pregnancy_test_in_urine_serum2 = ['Positive (+ve)', 'Negative (-ve)']
        # Blood_Group = ['A(+ve)', 'B(+ve)', 'AB(+ve)', 'O(+ve)', 'O(-ve)', 'A(-ve)', 'B(-ve)', 'AB(-ve)']
        # test = ['Toxoplasma IgG', 'Toxoplasma IgM', 'Cytomegalo Virus IgG', 'Cytomegalo Virus IgM', 'Rubella IgG',
        #         'Rubella IgM', 'Anti - Phspholipin IgG', 'Anti - Phspholipin  IgM', 'Anti - Cardiolipin  IgG',
        #         'Anti - Cardiolipin  IgM', 'Herps   IgG', 'Herpes  IgM']
        # test_choices = ['0.5 Negative', '0.6 Negative', '0.7 Negative', '0.8 Negative', '1.1 Positive', '1.1 Positive',
        #                 '1.2 Positive', '1.3 Positive', '1.4 Positive', '1.5 Positive']
        # Hb = [4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18]
        # motility = [5, 10, 15, 20, 25, 30, 35, 40, 45, 50, 55, 60, 65, 70, 75, 80, 85]

    def Chick_analyst_category(self, sd=None):
        # print(item,'here')
        global mylist
        # print(self.comboBox_16.currentText(), 'jjjjj')
        if self.comboBox_16.currentText() not in mylist and self.comboBox_16.currentIndex() != 0:
            self.analyst_category_function(self.comboBox_16.currentText())
        if self.comboBox_16.currentText() in mylist:
            try:
                ritem = self.comboBox_16.model().itemFromIndex(sd)
                if ritem.checkState() == Qt.Unchecked:
                    ritem.setCheckState(Qt.Checked)
                    # print(ritem.text())
                    self.analyst_category_function(ritem.text())
                    ritem.setCheckState(Qt.Unchecked)
                else:
                    ritem.setCheckState(Qt.Unchecked)
                ritem.setCheckState(Qt.Unchecked)
            except Exception as e:
                print(e, '3erorr')
                # print('erorr 404')

    def get_client_id(self):
        global client_id_glob
        client_id_glob = self.spinBox.value()
        self.Show_All_one_client_analyst()

        self.Add_Data_To_history(6, 1)
        self.History()

    def Show_All_The_Sales(self):
        global show_all_sales_in_clients_page
        self.cur.execute(''' SELECT client_name FROM addclient WHERE DATE(date)=%s ORDER BY -date ''',
                         (datetime.date.today(),))
        analyst_data = self.cur.fetchall()
        all_names = {}
        all_data = []
        id = 0
        for item in analyst_data:
            if item[0] not in all_names:
                all_names[item[0]] = 0
        for s in all_names.keys():
            self.cur.execute(''' SELECT id FROM addclient WHERE client_name=%s ORDER BY id ''', (s,))
            analyst_data2 = self.cur.fetchall()
            all_names[s] = analyst_data2[0][0]
            id = analyst_data2[0][0]
            self.cur.execute(''' SELECT * FROM addclient WHERE id=%s ORDER BY -date ''', (id,))
            latest_data = self.cur.fetchall()
            data = [{'name': s, 'value': latest_data}]
            all_data.append(data)
        self.tableWidget_6.setRowCount(0)
        self.tableWidget_6.insertRow(0)

        l_data = ()
        for ih in all_data:
            hs_data = ih[0]['value']
            l_data += tuple(hs_data)
        for row, form in enumerate(l_data):
            for col, item in enumerate(form):
                if col == 3:
                    self.tableWidget_6.setItem(row, col, QTableWidgetItem(str(l_data[row][4])))
                else:
                    self.tableWidget_6.setItem(row, col, QTableWidgetItem(str(item)))
                col += 1
            row_pos = self.tableWidget_6.rowCount()
            self.tableWidget_6.insertRow(row_pos)

    def Show_All_The_Analysts(self):
        print('dytuope04')
        global from_start
        search_type = self.comboBox_19.currentText()
        search_words = self.lineEdit_26.text()

        if self.comboBox_19.currentIndex() == 1:
            self.cur.execute(
                ''' SELECT name,category,defult,unit,price,sub_category FROM addanalyst WHERE name=%s ''',
                (search_words,))
        if self.comboBox_19.currentIndex() == 2:
            self.cur.execute(
                ''' SELECT name,category,defult,unit,price,sub_category FROM addanalyst WHERE price=%s ''',
                (search_words,))
        if self.comboBox_19.currentIndex() == 3:
            self.cur.execute(
                ''' SELECT name,category,defult,unit,price,sub_category FROM addanalyst WHERE sub_category=%s ''',
                (search_words,))
        if self.comboBox_19.currentIndex() == 0:
            self.cur.execute(
                ''' SELECT name,category,defult,unit,price,sub_category FROM addanalyst ORDER BY sub_category''')

        analyst_data = self.cur.fetchall()
        self.tableWidget_7.setSortingEnabled(False)
        # self.tableWidget_7.isSortingEnabled()
        self.tableWidget_7.setRowCount(0)
        self.tableWidget_7.insertRow(0)
        for row, form in enumerate(analyst_data):
            for col, item in enumerate(form):
                if col == 1:
                    category = ''
                    if analyst_data[row][1] == 'خيارات':
                        category = 'خيارات'
                    if analyst_data[row][1] == 'عدد':
                        category = 'عدد'
                    if analyst_data[row][1] == 'خيارات مع تعديل':
                        category = 'خيارات مع تعديل'
                    if analyst_data[row][1] == 'حقل كتابة':
                        category = 'حقل كتابة'

                    self.tableWidget_7.setItem(row, col, QTableWidgetItem(str(category)))
                if col == 5:
                    if analyst_data[row][5]==None:
                        analyst_data[row][5]=''
                    self.tableWidget_7.setItem(row, col, QTableWidgetItem(str(analyst_data[row][5])))
                if item==None:
                    item = ''
                self.tableWidget_7.setItem(row, col, QTableWidgetItem(str(item)))
                col += 1
            row_pos = self.tableWidget_7.rowCount()
            self.tableWidget_7.insertRow(row_pos)
        if from_start:
            pass
        else:
            self.Add_Data_To_history(6, 2)
            self.History()
        self.tableWidget_7.setSortingEnabled(True)
    def Add_Analyst(self):
        global analysts_name_glo
        analyst_name = self.lineEdit_28.text()
        if analyst_name not in analysts_name_glo:
            analyst_result_category = self.comboBox_22.currentText()
            analyst_price = self.spinBox_5.value()
            categoryes_number = self.comboBox_23.count()
            categoryes = []
            sub_category = self.comboBox_23.currentText()
            date = datetime.datetime.now()
            defult = self.lineEdit_40.text()
            unit = self.lineEdit_48.text()
            results_number = self.comboBox_30.count()
            results = []
            for i in range(0, results_number):
                self.comboBox_30.setCurrentIndex(i)
                results.append(str(self.comboBox_30.currentText()))
            # print(str(results))
            self.cur.execute(
                ''' INSERT INTO addanalyst (name,price,category,sub_category,date,defult,unit,results) VALUES(%s,%s,%s,%s,%s,%s,%s,%s) ''',
                (
                    str(analyst_name), analyst_price, analyst_result_category, sub_category,
                    date, defult, unit, str(results)[1:-1]))
            self.db.commit()
            self.lineEdit_28.setText('')  # analyst_name =
            self.comboBox_22.setCurrentIndex(0)  # analyst_result_category =
            self.spinBox_5.setValue(0)  # analyst_price =
            self.comboBox_23.setCurrentIndex(0)  # sub_category =
            self.lineEdit_40.setText('')
            self.lineEdit_48.setText('')
            QMessageBox.information(self, 'info', 'تم اضافة التحليل بنجاح')
            self.add_Analyst_to_list()
            self.add_Analyst_to_list()
            self.add_Analyst_to_list()
            self.Auto_complete_combo()
            self.Show_all_analysts_in_combo()
            self.Show_All_The_Analysts()
            self.Add_Data_To_history(3, 2)
            self.History()
        else:
            QMessageBox.information(self, '', 'هذا التحليل موجود بالفعل')

    def Show_analyst_in_Edit_Or_Delete(self):
        self.comboBox_31.clear()
        analyst_current_name = self.comboBox_21.currentText()
        analyst_current_index = self.comboBox_21.currentIndex()
        if not analyst_current_index:
            # QMessageBox.information(self, 'تحذير', "يرجى اختيار تحليل صحيح")
            self.comboBox_26.setCurrentText('')  # sub_category =
            self.comboBox_26.setCurrentIndex(0)  # sub_category =
            self.comboBox_25.setCurrentIndex(0)  # sub_category =
            self.lineEdit_29.setText('')  # analyst_name =
            self.comboBox_25.setCurrentText('')  # analyst_result_category =
            self.lineEdit_39.setText('')  # defult =
            self.lineEdit_38.setText('')  # unit =
            self.spinBox_6.setValue(0)  # analyst_price =
        else:
            self.cur.execute(
                ''' SELECT name,defult,unit,price,category,sub_category,results FROM addanalyst WHERE name=%s ''',
                (analyst_current_name,))
            data = self.cur.fetchall()
            if data:
                if data[0][6]:
                    list_data = str(data[0][6]).replace("'", "")
                    list_data = list_data.split(',')
                    # print(list(list_data))
                    # print(list_data)
                    # print(list_data[1])
                    for rr in range(0, len(list_data)):
                        self.comboBox_31.addItem(list_data[rr])  # results =
                        # analyst_choices_results.append(list_data[rr])
                self.comboBox_26.setCurrentText(str(data[0][5]))  # sub_category =
                self.lineEdit_29.setText(str(data[0][0]))  # analyst_name =
                self.comboBox_25.setCurrentText(str(data[0][4]))  # analyst_result_category =
                self.lineEdit_39.setText(data[0][1])  # defult =
                self.lineEdit_38.setText(data[0][2])  # unit =
                self.spinBox_6.setValue(data[0][3])  # analyst_price =
            # self.Add_Data_To_history(6, 2)
            # self.History()
            else:
                # QMessageBox.information(self, 'تحذير', "يرجى اختيار تحليل صحيح")
                self.comboBox_26.setCurrentText('')  # sub_category =
                self.comboBox_26.setCurrentIndex(0)  # sub_category =
                self.comboBox_25.setCurrentIndex(0)  # sub_category =
                self.lineEdit_29.setText('')  # analyst_name =
                self.comboBox_25.setCurrentText('')  # analyst_result_category =
                self.lineEdit_39.setText('')  # defult =
                self.lineEdit_38.setText('')  # unit =
                self.spinBox_6.setValue(0)  # analyst_price =

    def Update_addNewItemAnalysts(self):
        global analyst_name_for_update
        global analyst_name_for_update_before
        print(analyst_name_for_update)
        print(analyst_name_for_update_before)
        self.cur.execute('''UPDATE addnewitem set analyst_name=%s where analyst_name=%s''',
                         (analyst_name_for_update, analyst_name_for_update_before))
        self.db.commit()
        print('complete')
    def Edit_Analyst(self):
        global analysts_name_glo
        global analyst_name_for_update
        global analyst_name_for_update_before
        analyst_name = self.lineEdit_29.text()
        analyst_name_before = self.comboBox_21.currentText()
        if self.comboBox_21.currentText() in analysts_name_glo:
            analyst_result_category = self.comboBox_25.currentText()
            defult = self.lineEdit_39.text()
            unit = self.lineEdit_38.text()
            analyst_price = self.spinBox_6.value()
            sub_category = self.comboBox_26.currentText()
            date = datetime.datetime.now()
            results_number = self.comboBox_31.count()
            results = []
            for i in range(0, results_number):
                self.comboBox_31.setCurrentIndex(i)
                results.append(str(self.comboBox_31.currentText()))
            mysql = '''UPDATE addanalyst SET name=%s,defult=%s,unit=%s,price=%s,category=%s,sub_category=%s,date=%s,results=%s where name=%s'''
            values = (
                str(analyst_name), defult, unit, analyst_price, analyst_result_category, sub_category, date,
                str(results)[1:-1],
                analyst_name_before)
            self.cur.execute(mysql, values)
            self.db.commit()
            analyst_name_for_update = analyst_name
            analyst_name_for_update_before = analyst_name_before
            self.comboBox_21.setCurrentIndex(0)  # analyst_current_name =
            self.lineEdit_29.setText('')  # analyst_name =
            self.comboBox_25.setCurrentIndex(0)  # analyst_result_category =
            self.lineEdit_39.setText('')  # defult =
            self.lineEdit_38.setText('')  # unit =
            self.spinBox_6.setValue(0)  # analyst_price =
            self.comboBox_26.setCurrentIndex(0)  # sub_category =
            QMessageBox.information(self, 'info', 'تم تعديل التحليل بنجاح')
            self.add_Analyst_to_list()
            self.Auto_complete_combo()
            self.Show_all_analysts_in_combo()
            self.Show_All_The_Analysts()
            self.Add_Data_To_history(4, 2)
            self.History()
        else:
            QMessageBox.information(self, '', 'هذا التحليل غير موجود')
        thread_0 = Thread(target=self.Update_addNewItemAnalysts)
        thread_0.start()

    def Delete_Analyst(self):
        warning = QMessageBox.warning(self, 'احذر', "هل انت متأكد من انك تريد مسح التحليل",
                                      QMessageBox.Yes | QMessageBox.No)
        if warning == QMessageBox.Yes:
            analyst_current_name = self.comboBox_21.currentText()
            sql = ''' DELETE FROM addanalyst WHERE name=%s '''
            self.cur.execute(sql, [(analyst_current_name)])
            self.db.commit()
            QMessageBox.information(self, 'info', 'تم حذف التحليل بنجاح')
            self.Show_all_analysts_in_combo()
            self.Add_Data_To_history(5, 2)
            self.History()

    def Show_all_analysts_in_combo(self):
        global mylist
        global addTrue
        addTrue = True
        self.cur.execute(
            ''' SELECT name FROM addanalyst ORDER BY sub_category''')
        data = self.cur.fetchall()
        self.comboBox_21.clear()
        self.comboBox_16.clear()
        self.comboBox_26.clear()
        self.comboBox_23.clear()
        self.comboBox_26.addItem('----------------')
        self.comboBox_23.addItem('----------------')
        self.comboBox_21.addItem('----------------')
        self.comboBox_16.addItem('----------------')
        self.comboBox_5.addItem('الكل')

        self.comboBox_16.addItems(mylist)
        self.comboBox_5.addItems(mylist)
        self.comboBox_26.addItems(mylist)
        self.comboBox_23.addItems(mylist)
        for item in data:
            self.comboBox_21.addItem(str(item[0]))
            self.comboBox_16.addItem(str(item[0]))
            self.comboBox_5.addItem(str(item[0]))
        for ii in range(0, self.comboBox_16.count()):
            self.comboBox_16.setCurrentIndex(ii)
            if self.comboBox_16.currentText() in mylist:
                # self.comboBox_16.setItemChecked(self.comboBox_16.currentIndex(),False)
                myitem = self.comboBox_16.model().item(self.comboBox_16.currentIndex(), self.comboBox_16.modelColumn())
                myitem.setCheckState(Qt.Unchecked)
        self.comboBox_16.setCurrentIndex(0)
        addTrue = False
        # self.handel_buttons()

    def Clients_Page(self):
        global search_info_by_date
        self.tableWidget_4.setSortingEnabled(False)
        self.tableWidget_9.setSortingEnabled(False)
        self.tableWidget_4.setRowCount(0)
        self.tableWidget_4.insertRow(0)
        self.tableWidget_9.setRowCount(0)
        self.tableWidget_9.insertRow(0)
        id = self.spinBox_2.value()
        # date = self.dateEdit_3.date()

        if search_info_by_date:
            id = self.searchWidget2.spinBox.value()
            from_date = str(self.searchWidget2.dateEdit_6.date().toPyDate())
            to_date = str(self.searchWidget2.dateEdit_5.date().toPyDate())
            self.cur.execute(
                ''' SELECT price,analyst_name,analyst_result,client_name,date FROM addnewitem WHERE client_id=%s AND DATE(date)>=%s AND DATE(date)<=%s''',
                (str(id), from_date, to_date,))
            client_analyst_data = self.cur.fetchall()
            self.cur.execute(
                ''' SELECT client_name,client_age,client_genus,client_doctor FROM addclient WHERE id=%s AND DATE(date)>=%s AND DATE(date)<=%s''',
                (str(id), from_date, to_date,))
            client_data = self.cur.fetchall()
        else:
            self.cur.execute(
                ''' SELECT price,analyst_name,analyst_result,client_name,date FROM addnewitem WHERE client_id=%s''',
                (str(id),))
            client_analyst_data = self.cur.fetchall()
            self.cur.execute(''' SELECT client_name,client_age,client_genus,client_doctor FROM addclient WHERE id=%s''',
                             (str(id),))
            client_data = self.cur.fetchall()
        num = 0
        all_client_analyst = []
        total = 0
        for i in client_analyst_data:
            num += 1
        for price in range(0, num):
            total += client_analyst_data[price][0]
        for j in range(0, num):
            all_client_analyst.append(str(client_analyst_data[j][1]))
        self.tableWidget_4.setRowCount(0)
        self.tableWidget_4.insertRow(0)
        for row, form in enumerate(client_data):
            for col, item in enumerate(form):
                self.tableWidget_4.setItem(row, 4, QTableWidgetItem(str(num)))
                self.tableWidget_4.setItem(row, 5, QTableWidgetItem(str(','.join(all_client_analyst))))
                self.tableWidget_4.setItem(row, 6, QTableWidgetItem(str(total)))
                self.tableWidget_4.setItem(row, col, QTableWidgetItem(str(item)))
                col += 1

            row_pos = self.tableWidget_4.rowCount()
            self.tableWidget_4.insertRow(row_pos)
        self.tableWidget_9.setRowCount(0)
        self.tableWidget_9.insertRow(0)
        for row, form in enumerate(client_analyst_data):
            for col, item in enumerate(form):
                if col == 0:
                    self.tableWidget_9.setItem(row, col, QTableWidgetItem(str(client_analyst_data[0][3])))
                if col == 1:
                    self.tableWidget_9.setItem(row, col, QTableWidgetItem(str(client_analyst_data[row][1])))
                if col == 2:
                    self.tableWidget_9.setItem(row, col, QTableWidgetItem(str(client_analyst_data[row][2])))
                if col == 3:
                    self.tableWidget_9.setItem(row, col, QTableWidgetItem(str(client_analyst_data[row][4])))
                col += 1

            row_pos = self.tableWidget_9.rowCount()
            self.tableWidget_9.insertRow(row_pos)
        if search_info_by_date:
            self.searchWidget2.close()
        self.Add_Data_To_history(6, 5)
        self.History()
        search_info_by_date = False
        self.tableWidget_4.setSortingEnabled(True)
        self.tableWidget_9.setSortingEnabled(True)

    def Add_Buys(self):
        item_name = self.lineEdit_13.text()
        quantity = self.spinBox_3.value()
        signal_item_price = self.spinBox_8.value()
        total_item_price = signal_item_price * quantity
        date_create_before = str(self.dateEdit.date().toPyDate())
        self.cur.execute(
            ''' INSERT INTO addbuys (item_name,signal_item_price,total_price,quantity,date) VALUES (%s,%s,%s,%s,%s)'''
            , (item_name, signal_item_price, total_item_price, quantity, date_create_before))
        self.db.commit()
        self.Show_all_buys()
        self.lineEdit_13.setText('')
        self.spinBox_3.setValue(0)
        self.spinBox_8.setValue(0)
        self.Add_Data_To_history(3, 4)
        self.History()

    def Show_all_buys(self):
        self.cur.execute(''' SELECT item_name,quantity,signal_item_price,total_price,date FROM addbuys ''')
        data = self.cur.fetchall()
        self.tableWidget_3.setSortingEnabled(False)
        self.tableWidget_3.setRowCount(0)
        self.tableWidget_3.insertRow(0)
        for row, form in enumerate(data):
            for col, item in enumerate(form):
                self.tableWidget_3.setItem(row, col, QTableWidgetItem(str(item)))
                col += 1
            row_pos = self.tableWidget_3.rowCount()
            self.tableWidget_3.insertRow(row_pos)
        for rowd in range(0, self.tableWidget_3.rowCount() - 1):
            mypush_button = QPushButton(self)
            mypush_button.setText('حذف')
            mypush_button.setStyleSheet('border-radius:12px;')
            mypush_button.clicked.connect(self.buttonClicked2)
            self.tableWidget_3.setCellWidget(rowd, 5, mypush_button)
        self.tableWidget_3.setSortingEnabled(True)
    def Show_statics(self):
        all_categorys = []
        self.cur.execute(''' SELECT sub_category FROM addanalyst ''')
        datagh = self.cur.fetchall()
        for ijko in datagh:
            if ijko[0] not in all_categorys:
                all_categorys.append(ijko[0])
        from_date = self.dateEdit_6.date().toPyDate()
        to_date = self.dateEdit_5.date().toPyDate()
        item = self.comboBox_5.currentText()
        if self.comboBox_5.currentIndex() == 0:
            self.cur.execute(''' SELECT sum(total_price) FROM addbuys WHERE DATE(date)>=%s AND DATE(date)<=%s''',
                             (str(from_date), str(to_date),))
            data = self.cur.fetchone()
            if data:
                total_buys_price = data[0]
            self.cur.execute(
                ''' select sum(price) AS Price from addnewitem where DATE(date)>=%s and DATE(date)<=%s Group by DATE(date) Order by Price DESC  ''',
                (str(from_date), str(to_date)))
            data2 = self.cur.fetchall()

            total_sales_price = 0
            total_sales_items = 0
            if data2:
                for i in range(0, len(data2)):
                    total_sales_price += data2[i][0]
                self.lineEdit_33.setText(str(total_sales_price))
            if data:
                self.lineEdit_30.setText(str(total_buys_price))
            if data2 and data:
                self.lineEdit_31.setText(str(total_sales_price - total_buys_price))
                self.lineEdit_32.setText(str((total_sales_price / total_buys_price) * 100)[0] + '%')
            # self.cur.execute(
            #     ''' select count(price) AS Price from addnewitem where DATE(date)>=%s and DATE(date)<=%s Group by DATE(date) Order by Price DESC  ''',(str(from_date), str(to_date)))
            # data2 = self.cur.fetchone()
            self.cur.execute(''' select id from addnewitem where DATE(date)>=%s and DATE(date)<=%s ''',
                             (str(from_date), str(to_date)))
            data = self.cur.fetchall()
            sales_num = 0
            for iueh in data:
                sales_num += 1
            # print(data, ';k')
            self.lineEdit_36.setText(str(sales_num))
            self.Add_Data_To_history(6, 8)
            self.History()
        else:
            if item not in all_categorys:
                self.cur.execute(''' SELECT sum(total_price) FROM addbuys WHERE DATE(date)>=%s AND DATE(date)<=%s''',
                                 (str(from_date), str(to_date),))
                data = self.cur.fetchone()
                print(data)
                if data:
                    total_buys_price = data[0]
                self.cur.execute(
                    ''' select sum(price) AS Price from addnewitem where DATE(date)>=%s and DATE(date)<=%s AND analyst_name=%s Group by DATE(date) Order by Price DESC  ''',
                    (str(from_date), str(to_date), item))
                data2 = self.cur.fetchall()
                total_sales_price = 0
                total_sales_items = 0
                if data2:
                    for i in range(0, len(data2)):
                        total_sales_price += data2[i][0]
                    self.lineEdit_33.setText(str(total_sales_price))
                if data:
                    self.lineEdit_30.setText(str(total_buys_price))
                if data:
                    if data2:
                        self.lineEdit_31.setText(str(total_sales_price - total_buys_price))
                        self.lineEdit_32.setText(str((total_sales_price / total_buys_price) * 100)[0] + '%')
                self.cur.execute(
                    ''' select count(analyst_name) from addnewitem where DATE(date)>=%s and DATE(date)<=%s AND analyst_name=%s ''',
                    (str(from_date), str(to_date), item))
                data = self.cur.fetchall()
                if data:
                    self.lineEdit_36.setText(str(data[0][0]))
            else:
                self.cur.execute(''' SELECT sum(total_price) FROM addbuys WHERE DATE(date)>=%s AND DATE(date)<=%s''',
                                 (str(from_date), str(to_date),))
                data = self.cur.fetchone()
                total_buys_price = 0
                if data:
                    total_buys_price = data[0]
                self.cur.execute(
                    ''' select sum(price) AS Price from addnewitem where DATE(date)>=%s and DATE(date)<=%s AND sub_category=%s Group by DATE(date) Order by Price DESC  ''',
                    (str(from_date), str(to_date), item))
                data2 = self.cur.fetchall()
                total_sales_price = 0
                if data2:
                    for i in range(0, len(data2)):
                        total_sales_price += data2[i][0]
                        self.lineEdit_30.setText(str(total_buys_price))
                if data2:
                    self.lineEdit_33.setText(str(total_sales_price))
                if data and data2 and total_buys_price:
                    self.lineEdit_31.setText(str(total_sales_price - total_buys_price))
                    self.lineEdit_32.setText(str((total_sales_price / total_buys_price) * 100)[0] + '%')
                else:
                    self.lineEdit_30.setText('0')
                    self.lineEdit_33.setText('0')
                    self.lineEdit_31.setText(str('0'))
                    self.lineEdit_32.setText('0')
                self.cur.execute(
                    ''' select count(sub_category) from addnewitem where DATE(date)>=%s and DATE(date)<=%s AND sub_category=%s ''',
                    (str(from_date), str(to_date), item))
                data = self.cur.fetchall()
                if data:
                    self.lineEdit_36.setText(str(data[0][0]))
                else:
                    self.lineEdit_36.setText('0')

    def Show_default_statics(self):
        print('run')
        if self.tabWidget_4.currentIndex() == 4:
            print('yes')
            try:
                self.cur.execute(
                    '''select analyst_name,count(price) AS Price from addnewitem Group by analyst_name Order by Price DESC limit 5''')
                data = self.cur.fetchall()
                # print(max(analysts_counts))
                self.comboBox_7.clear()
                for i in range(0, len(data)):
                    self.comboBox_7.addItem(data[i][0] + '     ' + str(data[i][1]))
                self.cur.execute(
                    '''select sub_category,count(sub_category) AS Price from addnewitem Group by sub_category Order by Price DESC''')
                data = self.cur.fetchall()

                self.comboBox_8.clear()
                if data:
                    for i7 in range(0, len(data)):
                        self.comboBox_8.addItem(data[i7][0] + '       ' + str(data[i7][1]))
                self.cur.execute(
                    ''' select analyst_name,sum(price) AS Price from addnewitem Group by analyst_name Order by Price DESC limit 5''')
                data = self.cur.fetchall()
                self.comboBox.clear()
                for iqw in range(0, len(data)):
                    self.comboBox.addItem(data[iqw][0] + '     ' + str(data[iqw][1]))
                self.cur.execute(
                    ''' select sub_category,sum(price) AS Price from addnewitem Group by sub_category Order by Price DESC''')
                data = self.cur.fetchall()
                # print(data)
                self.comboBox_6.clear()
                for iqw2 in range(0, len(data)):
                    self.comboBox_6.addItem(data[iqw2][0] + '     ' + str(data[iqw2][1]))
                self.cur.execute(
                    ''' select date(date),count(date(date)) AS Price from addnewitem Group by date(date) Order by Price DESC limit 10 ''')
                data = self.cur.fetchall()
                self.comboBox_9.clear()
                for iqw3 in range(0, len(data)):
                    self.comboBox_9.addItem(str(data[iqw3][0]) + '     ' + str(data[iqw3][1]))
                self.cur.execute(
                    ''' select date(date),sum(price) AS Price from addnewitem Group by date(date) Order by Price DESC limit 10 ''')
                data = self.cur.fetchall()
                self.comboBox_10.clear()
                for iqw4 in range(0, len(data)):
                    self.comboBox_10.addItem(str(data[iqw4][0]) + '     ' + str(data[iqw4][1]))
            except Exception as e:
                print(e, '98erorr')

    def Add_Data_To_history(self, action, table):
        global user_id
        self.cur.execute(
            ''' INSERT INTO his VALUES (DEFAULT,%s,%s,%s,%s,%s)''',
            (user_id, action, table, datetime.datetime.now(), 1))
        self.db.commit()

    def History(self):
        self.cur.execute('''SELECT uid,action,tabled,dates FROM his ORDER BY -dates''')
        analyst_data = self.cur.fetchall()
        self.tableWidget_8.setRowCount(0)
        self.tableWidget_8.insertRow(0)
        for row, form in enumerate(analyst_data):
            for col, item in enumerate(form):
                if col == 0:
                    self.tableWidget_8.setItem(row, col, QTableWidgetItem(str(analyst_data[row][0])))
                if col == 1:
                    action = ''
                    if analyst_data[row][1] == 1:
                        action = 'تسجيل الدخول'
                    if analyst_data[0][1] == 2:
                        action = 'تسجيل الخروج'
                    if analyst_data[row][1] == 3:
                        action = 'اضافة'
                    if analyst_data[row][1] == 4:
                        action = 'تعديل'
                    if analyst_data[row][1] == 5:
                        action = 'حذف'
                    if analyst_data[row][1] == 6:
                        action = 'بحث'
                    if analyst_data[row][1] == 7:
                        action = 'طباعة'
                    if analyst_data[row][1] == 8:
                        action = 'معاينة'
                    self.tableWidget_8.setItem(row, col, QTableWidgetItem(str(action)))
                if col == 2:
                    tables = ''
                    if analyst_data[row][2] == 1:
                        tables = 'مبيع يومي'
                    if analyst_data[row][2] == 2:
                        tables = 'تحليل'
                    if analyst_data[row][2] == 3:
                        tables = 'مشتريات'
                    if analyst_data[row][2] == 4:
                        tables = 'مراجعين'
                    if analyst_data[row][2] == 5:
                        tables = 'مستخدم'
                    if analyst_data[row][2] == 6:
                        tables = 'طبيب'
                    if analyst_data[row][2] == 7:
                        tables = 'تقرير'
                    if analyst_data[row][2] == 8:
                        tables = 'احصائيات'
                    self.tableWidget_8.setItem(row, col, QTableWidgetItem(str(tables)))
                if col == 3:
                    self.tableWidget_8.setItem(row, col, QTableWidgetItem(str(analyst_data[row][3])))

                col += 1
            row_pos = self.tableWidget_8.rowCount()
            self.tableWidget_8.insertRow(row_pos)


    def Log_In_Chieck(self):
        global user_id
        global Edit_Doctor
        global Delete_Doctor
        global edit_employee_check
        global show_clients_check
        user_name = self.lineEdit.text()
        user_password = self.lineEdit_2.text()
        self.cur.execute(''' SELECT id,user_password,user_name FROM adduser WHERE user_name=%s ''', (user_name,))

        data = self.cur.fetchone()
        if data != None:
            if data[2] == user_name and data[1] == user_password:
                self.groupBox.setEnabled(True)
                user_id = data[2]
                try:
                    self.cur.execute(f' SELECT * FROM userper WHERE employee_name=%s ', (str(user_name),))
                    PerData = self.cur.fetchone()
                    if PerData:
                        if not PerData[27]:
                            self.pushButton_47.setEnabled(False)
                        else:
                            self.pushButton_47.setEnabled(True)
                        if not PerData[28]:
                            Edit_Doctor = False
                            self.pushButton_48.setEnabled(False)
                        else:
                            Edit_Doctor = True
                            self.pushButton_48.setEnabled(True)
                        if not PerData[29]:
                            Delete_Doctor = False
                            self.pushButton_49.setEnabled(False)
                        else:
                            Delete_Doctor = True
                            self.pushButton_49.setEnabled(True)
                        if not PerData[2]:
                            self.pushButton_17.setEnabled(False)
                            self.pushButton_44.setEnabled(False)
                        else:
                            self.pushButton_17.setEnabled(True)
                            self.pushButton_44.setEnabled(True)
                        if not PerData[3]:
                            self.pushButton.setEnabled(False)
                        else:
                            self.pushButton.setEnabled(True)
                        if not PerData[4]:
                            self.pushButton_2.setEnabled(False)
                        else:
                            self.pushButton_2.setEnabled(True)
                        if not PerData[5]:
                            self.pushButton_23.setEnabled(False)
                            self.comboBox.clear()
                            self.comboBox_6.clear()
                            self.comboBox_7.clear()
                            self.comboBox_8.clear()
                            self.comboBox_9.clear()
                            self.comboBox_10.clear()
                        else:
                            self.Show_default_statics()
                            self.pushButton_23.setEnabled(True)
                        if not PerData[6]:
                            self.pushButton_5.setEnabled(False)
                        else:
                            self.pushButton_5.setEnabled(True)

                        if not PerData[7]:
                            self.pushButton_3.setEnabled(False)
                        else:
                            self.pushButton_3.setEnabled(True)
                        if not PerData[8]:
                            self.pushButton_4.setEnabled(False)
                        else:
                            self.pushButton_4.setEnabled(True)

                        if not PerData[9]:
                            self.groupBox_4.setEnabled(False)
                        else:
                            self.groupBox_4.setEnabled(True)

                        if not PerData[10]:
                            self.pushButton_27.setEnabled(False)
                        else:
                            self.pushButton_27.setEnabled(True)

                        if not PerData[11]:
                            self.pushButton_28.setEnabled(False)
                        else:
                            self.pushButton_28.setEnabled(True)

                        if not PerData[12]:
                            self.pushButton_15.setEnabled(False)
                            self.pushButton_13.setEnabled(False)
                            self.pushButton_9.setEnabled(False)
                            self.pushButton_14.setEnabled(False)
                            self.pushButton_11.setEnabled(False)
                        else:
                            self.pushButton_15.setEnabled(True)
                            self.pushButton_13.setEnabled(True)
                            self.pushButton_9.setEnabled(True)
                            self.pushButton_14.setEnabled(True)
                            self.pushButton_11.setEnabled(True)
                        if not PerData[13]:
                            self.pushButton_16.setEnabled(False)
                        else:
                            self.pushButton_16.setEnabled(True)

                        if not PerData[14]:
                            self.pushButton_31.setEnabled(False)
                            self.groupBox_14.setEnabled(False)
                        else:
                            self.pushButton_31.setEnabled(True)
                            self.groupBox_14.setEnabled(True)
                        if not PerData[15]:
                            self.pushButton_24.setEnabled(False)
                        else:
                            self.pushButton_24.setEnabled(True)

                        if not PerData[16]:
                            self.pushButton_22.setEnabled(False)
                        else:
                            self.pushButton_22.setEnabled(True)

                        if not PerData[17]:
                            self.pushButton_53.setEnabled(False)
                        else:
                            self.pushButton_53.setEnabled(True)

                        if not PerData[18]:
                            self.pushButton_25.setEnabled(False)
                        else:
                            self.pushButton_25.setEnabled(True)

                        if not PerData[19]:
                            self.pushButton_37.setEnabled(False)
                        else:
                            self.pushButton_37.setEnabled(True)

                        if not PerData[20]:
                            self.pushButton_43.setEnabled(False)
                        else:
                            self.pushButton_43.setEnabled(True)
                        if not PerData[21]:
                            show_clients_check = False

                        else:
                            show_clients_check = True
                        if not PerData[22]:
                            self.pushButton_36.setEnabled(False)
                        else:
                            self.pushButton_36.setEnabled(True)

                        if not PerData[23]:
                            self.pushButton_29.setEnabled(False)
                            self.pushButton_42.setEnabled(False)
                        else:
                            self.pushButton_29.setEnabled(True)
                            self.pushButton_42.setEnabled(True)

                        if not PerData[24]:
                            self.pushButton_34.setEnabled(False)
                        else:
                            self.pushButton_34.setEnabled(True)

                        if not PerData[25]:
                            self.pushButton_18.setEnabled(False)
                        else:
                            self.pushButton_18.setEnabled(True)

                        if not PerData[26]:
                            self.pushButton_36.setEnabled(False)
                        else:
                            self.pushButton_36.setEnabled(True)
                except Exception as e:
                    print(e, '5erorr')
                self.lineEdit.setText('')
                self.lineEdit_2.setText('')
                self.tabWidget.setCurrentIndex(3)
                self.Add_Data_To_history(1, 5)
                self.History()

            else:
                warning = QMessageBox.warning(self, '',
                                              "كلمة المرور او اسم المستخدم غير صحيحة هل تريد استعادة كلمة المرور؟",
                                              QMessageBox.Yes | QMessageBox.No)
                if warning == QMessageBox.Yes:
                    self.Open_ResetPassword_Page()
            self.pushButton_44.setEnabled(False)
            self.pushButton_17.setEnabled(False)
            self.comboBox_16.setEnabled(False)
        else:
            warning = QMessageBox.warning(self, '',
                                          "كلمة المرور او اسم المستخدم غير صحيحة هل تريد استعادة كلمة المرور؟",
                                          QMessageBox.Yes | QMessageBox.No)
            if warning == QMessageBox.Yes:
                self.Open_ResetPassword_Page()

    def Add_all_employee_to_comboBox(self):
        self.comboBox_3.clear()
        self.cur.execute(''' SELECT user_name FROM adduser ''')
        data = self.cur.fetchall()
        self.comboBox_3.addItem('-----------------')
        for i in data:
            self.comboBox_3.addItem(i[0])

    def False_checkState(self):
        self.checkBox_58.setCheckState(False)
        self.checkBox_57.setCheckState(False)
        self.checkBox_56.setCheckState(False)
        self.checkBox_8.setCheckState(False)
        self.checkBox_15.setCheckState(False)
        self.checkBox_9.setCheckState(False)
        self.checkBox_55.setCheckState(False)
        self.checkBox_44.setCheckState(False)
        self.checkBox_40.setCheckState(False)
        self.checkBox_13.setCheckState(False)
        self.checkBox_16.setCheckState(False)
        # self.checkBox_3.setCheckState(False)
        self.checkBox_14.setCheckState(False)
        self.checkBox_7.setCheckState(False)
        self.checkBox_11.setCheckState(False)
        self.checkBox_12.setCheckState(False)
        self.checkBox_46.setCheckState(False)
        self.checkBox_49.setCheckState(False)
        self.checkBox_50.setCheckState(False)
        self.checkBox_47.setCheckState(False)
        self.checkBox_48.setCheckState(False)
        self.checkBox_51.setCheckState(False)
        self.checkBox_52.setCheckState(False)
        self.checkBox_53.setCheckState(False)
        self.checkBox_54.setCheckState(False)
        self.checkBox_44.setCheckState(False)
        self.checkBox_45.setCheckState(False)
        self.checkBox_42.setCheckState(False)
        self.checkBox_43.setCheckState(False)
        self.checkBox_41.setCheckState(False)

    def Add_employee(self):
        global Edit_employee
        employee_name = self.comboBox_3.currentText()
        employee_password = self.lineEdit_17.text()
        employee_email = self.lineEdit_3.text()
        if self.sender().text() != 'حفظ':
            self.cur.execute(''' INSERT INTO adduser (user_name,user_password,user_email,date) VALUES (%s,%s,%s,%s) ''',
                             (employee_name, employee_password, employee_email, str(datetime.datetime.now())))
            self.db.commit()
            QMessageBox.information(self, '', "تم اضافة الموظف بنجاح")
            self.Add_Data_To_history(3, 5)
            self.History()
        else:
            self.cur.execute(
                ''' UPDATE adduser SET user_name=%s,user_password=%s,user_email=%s,date=%s WHERE user_name=%s''',
                (employee_name, employee_password, employee_email, str(datetime.datetime.now()), employee_name))
            self.db.commit()
            QMessageBox.information(self, '', "تم تعديل الموظف بنجاح")
            self.Add_Data_To_history(4, 5)
            self.History()
        self.lineEdit_17.setText('')
        self.lineEdit_3.setText('')
        self.comboBox_3.setCurrentIndex(0)
        self.comboBox_2.setCurrentIndex(0)
        self.frame.show()
        self.Add_all_employee_to_comboBox()
        self.False_checkState()
    def Show_employee_data(self):
        global Edit_employee
        if self.comboBox_3.currentIndex() != 0:

            self.False_checkState()
            # try:
            self.cur.execute(''' SELECT user_name,user_password,user_email FROM adduser WHERE user_name=%s ''',
                             (self.comboBox_3.currentText(),))
            data = self.cur.fetchone()
            if data:
                Edit_employee = True
                self.pushButton_53.hide()
                self.pushButton_25.show()
                # self.comboBox_3.setCurrentText(data[0])
                self.lineEdit_17.setText(data[1])
                self.lineEdit_3.setText(data[2])
                self.cur.execute(''' SELECT * FROM userper WHERE employee_name=%s ''', (self.comboBox_3.currentText(),))
                perData = self.cur.fetchone()
                if perData:

                    if perData[27]:
                        self.checkBox_58.setCheckState(True)
                    if perData[28]:
                        self.checkBox_57.setCheckState(True)
                    if perData[29]:
                        self.checkBox_56.setCheckState(True)
                    if perData[3]:
                        self.checkBox_9.setCheckState(True)
                    if perData[6]:
                        self.checkBox_55.setCheckState(True)
                    if perData[4]:
                        self.checkBox_40.setCheckState(True)
                    if perData[26]:
                        self.checkBox_13.setCheckState(True)
                    if perData[2]:
                        self.checkBox_8.setCheckState(True)
                    if perData[24]:
                        self.checkBox_15.setCheckState(True)
                    if perData[25]:
                        self.checkBox_16.setCheckState(True)
                    if perData[20]:
                        self.checkBox_14.setCheckState(True)
                    if perData[21]:
                        self.checkBox_7.setCheckState(True)
                    if perData[22]:
                        self.checkBox_11.setCheckState(True)
                    if perData[23]:
                        self.checkBox_12.setCheckState(True)
                    if perData[8]:
                        self.checkBox_46.setCheckState(True)
                    if perData[12]:
                        self.checkBox_49.setCheckState(True)
                    if perData[15]:
                        self.checkBox_50.setCheckState(True)
                    if perData[14]:
                        self.checkBox_47.setCheckState(True)
                    if perData[13]:
                        self.checkBox_48.setCheckState(True)
                    if perData[17]:
                        self.checkBox_51.setCheckState(True)
                    if perData[18]:
                        self.checkBox_52.setCheckState(True)
                    if perData[19]:
                        self.checkBox_53.setCheckState(True)
                    if perData[5]:
                        self.checkBox_54.setCheckState(True)
                    if perData[7]:
                        self.checkBox_44.setCheckState(True)
                    if perData[16]:
                        self.checkBox_45.setCheckState(True)
                    # if self.checkBox_40.isChecked() == True:
                    #     analysts_page = 1
                    if perData[9]:
                        self.checkBox_42.setCheckState(True)
                    if perData[10]:
                        self.checkBox_43.setCheckState(True)
                    if perData[11]:
                        self.checkBox_41.setCheckState(True)
            else:
                self.pushButton_25.hide()
                self.pushButton_53.show()
                self.comboBox_2.setCurrentIndex(0)
                self.False_checkState()
                self.lineEdit_17.setText('')
                self.lineEdit_3.setText('')
        else:
            self.comboBox_2.setCurrentIndex(0)
            self.False_checkState()
            self.lineEdit_17.setText('')
            self.lineEdit_3.setText('')

    def Show_permissions(self):
        if self.comboBox_2.currentIndex() != 0:
            if self.comboBox_2.currentIndex() == 1:
                self.frame.hide()
                self.groupBox_21.show()
                self.groupBox_16.hide()
                self.groupBox_20.hide()
                self.groupBox_19.hide()
                self.groupBox_15.hide()
            if self.comboBox_2.currentIndex() == 2:
                self.frame.hide()
                self.groupBox_15.show()
                self.groupBox_21.hide()
                self.groupBox_16.hide()
                self.groupBox_20.hide()
                self.groupBox_19.hide()
            if self.comboBox_2.currentIndex() == 3:
                self.frame.hide()
                self.groupBox_16.show()
                self.groupBox_15.hide()
                self.groupBox_21.hide()
                self.groupBox_20.hide()
                self.groupBox_19.hide()
            if self.comboBox_2.currentIndex() == 4:
                self.frame.hide()
                self.groupBox_19.show()
                self.groupBox_15.hide()
                self.groupBox_21.hide()
                self.groupBox_16.hide()
                self.groupBox_20.hide()
            if self.comboBox_2.currentIndex() == 5:
                self.frame.hide()
                self.groupBox_20.show()
                self.groupBox_15.hide()
                self.groupBox_21.hide()
                self.groupBox_16.hide()
                self.groupBox_19.hide()
        else:
            self.frame.show()

    def Add_permissions(self):
        employee_name = self.comboBox_3.currentText()
        sales_page = 0
        analysts_page = 0
        settings_page = 0
        clients_page = 0
        history_page = 0
        add_sales = 0
        print_report = 0
        prev_report = 0
        search_by_date_in_sales = 0
        add_employee = 0
        edit_employee = 0
        delete_employee = 0
        add_doctor = 0
        edit_doctor = 0
        delete_doctor = 0
        add_analyst = 0
        edit_analyst = 0
        delete_analyst = 0
        show_client_info = 0
        change_path = 0
        change_theme = 0
        edit_report = 0
        add_buys = 0
        statics = 0
        search_in_all_sales = 0
        show_all_clients = 0
        search_client = 0
        delete_all_history = 0
        if self.checkBox_58.isChecked() == True:
            add_doctor = 1
        if self.checkBox_57.isChecked() == True:
            edit_doctor = 1
        if self.checkBox_56.isChecked() == True:
            delete_doctor = 1
        if self.checkBox_9.isChecked() == True:
            sales_page = 1
        if self.checkBox_55.isChecked() == True:
            clients_page = 1
        if self.checkBox_44.isChecked() == True:
            history_page = 1
        if self.checkBox_40.isChecked() == True:
            analysts_page = 1
        if self.checkBox_13.isChecked() == True:
            search_in_all_sales = 1
        if self.checkBox_8.isChecked() == True:
            add_sales = 1
        if self.checkBox_15.isChecked() == True:
            prev_report = 1
        if self.checkBox_16.isChecked() == True:
            print_report = 1
        if self.checkBox_13.isChecked() == True:
            search_in_all_sales = 1
        if self.checkBox_14.isChecked() == True:
            search_by_date_in_sales = 1
        if self.checkBox_7.isChecked() == True:
            show_all_clients = 1
        if self.checkBox_11.isChecked() == True:
            search_client = 1
        if self.checkBox_12.isChecked() == True:
            show_client_info = 1
        if self.checkBox_46.isChecked() == True:
            settings_page = 1
        if self.checkBox_49.isChecked() == True:
            change_theme = 1
        if self.checkBox_50.isChecked() == True:
            change_path = 1
        if self.checkBox_47.isChecked() == True:
            edit_report = 1
        if self.checkBox_48.isChecked() == True:
            add_buys = 1  #
        if self.checkBox_51.isChecked() == True:
            add_employee = 1
        if self.checkBox_52.isChecked() == True:
            edit_employee = 1
        if self.checkBox_53.isChecked() == True:
            delete_employee = 1
        if self.checkBox_54.isChecked() == True:
            statics = 1
        if self.checkBox_44.isChecked() == True:
            history_page = 1
        if self.checkBox_45.isChecked() == True:
            delete_all_history = 1
        if self.checkBox_40.isChecked() == True:
            analysts_page = 1
        if self.checkBox_42.isChecked() == True:
            add_analyst = 1
        if self.checkBox_43.isChecked() == True:
            edit_analyst = 1
        if self.checkBox_41.isChecked() == True:
            delete_analyst = 1
        if self.sender().text()=='اضافة':
            self.cur.execute(
                ''' INSERT INTO userper (employee_name,add_sale_item_page,sales_page,analyst_page,statics,clients_page,history_page,settings_page,add_analyst,edit_analyst,delete_analyst,change_theme,add_buys,edit_report,change_path,delet_history,add_user,edit_user,delete_user,search_by_date_in_sales_page,show_clients,search_client,search_client_info,prev_report,print_report,search_in_all_sales,add_doctor,edit_doctor,delete_doctor) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s) ''',
                (
                    employee_name, add_sales, sales_page, analysts_page, statics, clients_page, history_page,
                    settings_page,
                    add_analyst, edit_analyst, delete_analyst, change_theme, add_buys, edit_report, change_path,
                    delete_all_history, add_employee, edit_employee, delete_employee, search_by_date_in_sales,
                    show_all_clients, search_client, show_client_info, prev_report, print_report, search_in_all_sales,
                    add_doctor, edit_doctor, delete_doctor,))
            self.db.commit()
            QMessageBox.information(self, 'info', 'تم اضافة الموظف بنجاح')
        else:
            self.cur.execute(
                ''' UPDATE userper SET employee_name=%s,add_sale_item_page=%s,sales_page=%s,analyst_page=%s,statics=%s,clients_page=%s,history_page=%s,settings_page=%s,add_analyst=%s,edit_analyst=%s,delete_analyst=%s,change_theme=%s,add_buys=%s,edit_report=%s,change_path=%s,delet_history=%s,add_user=%s,edit_user=%s,delete_user=%s,search_by_date_in_sales_page=%s,show_clients=%s,search_client=%s,search_client_info=%s,prev_report=%s,print_report=%s,search_in_all_sales=%s,add_doctor=%s,edit_doctor=%s,delete_doctor=%s WHERE employee_name=%s ''',
                (
                    employee_name, add_sales, sales_page, analysts_page, statics, clients_page, history_page,
                    settings_page,
                    add_analyst, edit_analyst, delete_analyst, change_theme, add_buys, edit_report, change_path,
                    delete_all_history, add_employee, edit_employee, delete_employee, search_by_date_in_sales,
                    show_all_clients, search_client, show_client_info, prev_report, print_report, search_in_all_sales,
                    add_doctor, edit_doctor, delete_doctor,
                    self.comboBox_3.currentText()))
            self.db.commit()
            QMessageBox.information(self, 'info', 'تم تعديل بيانات الموظف بنجاح')

    def Delete_employee(self):
        warning = QMessageBox.warning(self, 'احذر', f"هل انت متأكد من انك تريد مسح الموظف {self.comboBox_4.text()}",
                                      QMessageBox.Yes | QMessageBox.No)
        if warning == QMessageBox.Yes:
            self.cur.execute(''' DELETE FROM adduser WHERE user_name=%s ''', (self.comboBox_4.text(),))
            self.cur.execute(''' DELETE FROM userper WHERE employee_name=%s ''', (self.comboBox_4.text(),))
            self.db.commit()
            self.False_checkState()
            self.lineEdit_17.setText('')
            self.lineEdit_3.setText('')
            self.comboBox_3.setCurrentIndex(0)
            self.comboBox_2.setCurrentIndex(0)
            self.Add_Data_To_history(5, 5)
            self.History()
        QMessageBox.information(self, 'info', 'تم حذف الموظف بنجاح')

    def Delete_All_History_Data(self):
        sql = '''DELETE FROM his'''
        self.cur.execute(sql)
        self.db.commit()
        QMessageBox.information(self, 'info', 'تم حذف محتويات السجل بنجاح')
        self.tableWidget_8.setRowCount(0)
        self.tableWidget_8.insertRow(0)

    def Open_Sales_Page(self):
        self.tabWidget.setCurrentIndex(3)

    def Open_Login_Page(self):
        self.tabWidget.setCurrentIndex(0)

    def Open_Settings_Page(self):
        self.tabWidget.setCurrentIndex(6)

    # def Open_Print_Page(self):
    #     self.tabWidget.setCurrentIndex(7)

    def Open_History_Page(self):
        self.tabWidget.setCurrentIndex(5)

    def Open_clients_Page(self):
        self.tabWidget.setCurrentIndex(2)

    def Open_Analyse_Page(self):
        self.tabWidget.setCurrentIndex(4)

    def Open_ResetPassword_Page(self):
        self.tabWidget.setCurrentIndex(1)

    def Light_Blue_Theme(self):
        style = open('thems/light_blue.css', 'r')
        style = style.read()
        self.setStyleSheet(style)

    def Dark_Blue_Theme(self):
        style = open('thems/darkblue.css', 'r')
        style = style.read()
        self.setStyleSheet(style)

    def Dark_Gray_Theme(self):
        style = open('thems/darkgray.css', 'r')
        style = style.read()
        self.setStyleSheet(style)

    def Dark_Orange_Theme(self):
        style = open('thems/darkorange.css', 'r')
        style = style.read()
        self.setStyleSheet(style)

    def Dark_Theme(self):
        style = open('thems/qdark.css', 'r')
        style = style.read()
        self.setStyleSheet(style)

    def Bio_Word(self, name, doctor, analysts, results, year, month, day, prev, genus, categorys):
        global if_print
        please = []
        categorys = []
        all_items = []
        all_items_for_units_defults = []
        all_items_index = []
        self.cur.execute(''' SELECT * FROM paths WHERE id=1 ''')
        mydata = self.cur.fetchone()
        word_files = mydata[1]
        save_word_files = mydata[2]
        self.cur.execute(''' SELECT * FROM  word WHERE id=1''')
        word_data = self.cur.fetchone()
        self.cur.execute(''' SELECT * FROM doctor WHERE name=%s ''', (doctor,))
        doctor_data = self.cur.fetchone()
        doctor_genus = doctor_data[2]
        files = 0
        units = []
        defults = []
        number_of_analysts = len(analysts)
        for my_index, ii in enumerate(analysts):
            self.cur.execute(''' SELECT sub_category FROM addanalyst WHERE name=%s ''', (ii,))
            data1 = self.cur.fetchall()
            if data1:
                analysts[my_index] = str(analysts[my_index] + data1[0][0])
                if data1[0][0] not in categorys:
                    categorys.append(data1[0][0])
        count = 1
        for c in categorys:
            all_items.append(str(c))
            all_items_for_units_defults.append(str(c))
            all_items_index.append(count)
            for sub in analysts:
                if c in sub:
                    count += 1
                    all_items.append(str('   ' + sub))
                    all_items_for_units_defults.append(str(sub))
                    all_items_index.append(count)
            count += 1
        for iplz in all_items_for_units_defults:
            for inm in categorys:
                if inm in iplz:
                    iplz = iplz[:-len(inm)]
                    self.cur.execute(''' SELECT unit,defult FROM addanalyst WHERE name=%s ''', (str(iplz),))
                    myplzdata = self.cur.fetchall()
                    if myplzdata:
                        if myplzdata[0][0]:
                            units.append(myplzdata[0][0])
                        else:
                            units.append('')
                        if myplzdata[0][1]:
                            defults.append(myplzdata[0][1])
                        else:
                            defults.append('')
                    else:
                        units.append('')
                        defults.append('')
        for my_index2, rowgh in enumerate(analysts):
            for item in categorys:
                if item in analysts[my_index2]:
                    analysts[my_index2] = str(analysts[my_index2][:-len(item)])
        for my_index3, rowgh2 in enumerate(all_items):
            for item2 in categorys:
                if item2 in all_items[my_index3]:
                    if all_items[my_index3].startswith('   '):
                        all_items[my_index3] = str(all_items[my_index3][:-len(item2)])
        f = open(r'%s\test-mydocx.docx' % word_files, 'rb')
        f.read()
        document = Document(f)
        if len(all_items) >= 20:
            files = 2
            f = open(r'%s\test-mydocx.docx' % word_files, 'rb')
            f.read()
            document = Document(f)
            f2 = open(r'%s\test-mydocx2.docx' % word_files, 'rb')
            f2.read()
            document2 = Document(f2)
        if len(all_items) > 40 and len(all_items) < 60:
            files = 3
            f = open(r'%s\test-mydocx.docx' % word_files, 'rb')
            f.read()
            document = Document(f)
            f2 = open(r'%s\test-mydocx2.docx' % word_files, 'rb')
            f2.read()
            document2 = Document(f2)
            f3 = open(r'%s\test-mydocx3.docx' % word_files, 'rb')
            f3.read()
            document3 = Document(f3)
        category_count = 0
        for table_index, i in enumerate(document.tables):
            for row_index, k in enumerate(i.rows):
                for cell_index, j in enumerate(k.cells):
                    for n in j.paragraphs:
                        if n.text == 'Date     /    /20':
                            n.text = f'Date  {year}/{month}/{day}'
                            run = n.runs
                            font = run[0].font
                            font.name = 'Tahoma'
                            font.bold = True
                            font.size = Pt(12)
                        if n.text == 'employee1':
                            n.text = str(word_data[4])
                            run = n.runs
                            font = run[0].font
                            font.name = 'Times New Roman'
                            font.bold = True
                            font.size = Pt(16)
                        if n.text == 'employee2':
                            n.text = str(word_data[5])
                            run = n.runs
                            font = run[0].font
                            font.name = 'Times New Roman'
                            font.bold = True
                            font.size = Pt(16)
                        if n.text == 'employeeshahada1':
                            n.text = str(word_data[6])
                            run = n.runs
                            font = run[0].font
                            font.name = 'Times New Roman'
                            font.bold = True
                            font.size = Pt(12)
                        if n.text == 'employeeshahada2':
                            n.text = str(word_data[7])
                            run = n.runs
                            font = run[0].font
                            font.name = 'Times New Roman'
                            font.bold = True
                            font.size = Pt(12)
                        if n.text == 'shopname':
                            n.text = str(word_data[1])
                            run = n.runs
                            font = run[0].font
                            font.name = 'Times New Roman'
                            font.bold = True
                            font.size = Pt(26)
                        if n.text == 'phone1':
                            n.text = str(word_data[2])
                            run = n.runs
                            font = run[0].font
                            font.name = 'Times New Roman'
                            font.bold = True
                            font.size = Pt(14)
                        if n.text == 'phone2':
                            n.text = str(word_data[3])
                            run = n.runs
                            font = run[0].font
                            font.name = 'Times New Roman'
                            font.bold = True
                            font.size = Pt(14)
                        if n.text == 'gps':
                            n.text = str(word_data[8])
                            run = n.runs
                            font = run[0].font
                            font.name = 'Times New Roman'
                            font.bold = True
                            font.size = Pt(13)
                        if n.text == 'client_name':
                            if genus == 1:
                                n.text = f'{word_data[9]}'
                            else:
                                n.text = f'{word_data[13]}'
                            run = n.runs
                            font = run[0].font
                            font.name = 'Times New Roman'
                            font.bold = True
                            font.size = Pt(12)
                        if n.text == 'rname':
                            n.text = str(name)
                            run = n.runs
                            font = run[0].font
                            font.name = 'Times New Roman'
                            font.bold = True
                            font.size = Pt(12)
                        if n.text == 'doctor':
                            if doctor_genus == 'female':
                                n.text = f'{word_data[15]}'
                            else:
                                n.text = f'{word_data[11]}'
                            run = n.runs
                            font = run[0].font
                            font.name = 'Times New Roman'
                            font.bold = True
                            font.size = Pt(12)
                        if n.text == 'rdoctor':
                            n.text = str(doctor)
                            run = n.runs
                            font = run[0].font
                            font.name = 'Times New Roman'
                            font.bold = True
                            font.size = Pt(12)
                        if n.text == 'lqb1':
                            if genus == 1:
                                n.text = str(word_data[10])
                            else:
                                n.text = str(word_data[14])
                            run = n.runs
                            font = run[0].font
                            font.name = 'Times New Roman'
                            font.bold = True
                            font.size = Pt(12)
                        if n.text == 'lqb2':
                            if doctor_genus == 'female':
                                n.text = str(word_data[16])
                            else:
                                n.text = str(word_data[12])
                            run = n.runs
                            font = run[0].font
                            font.name = 'Times New Roman'
                            font.bold = True
                            font.size = Pt(12)
                        for row in range(0, len(all_items)):
                            if n.text == str(all_items_index[row]) + 'r':
                                n.text = str(all_items[row]) + 'k'
                                # print(n.text,len(n.text))
                                if str(n.text)[:-1] not in categorys:
                                    run = n.runs
                                    font = run[0].font
                                    font.bold = True
                                    font.size = Pt(11)
                                    font.name = 'Tahoma'
                                else:
                                    n.text = ''
                                if str(all_items[row]) in categorys:
                                    category_count += 1

                                if n.text == str(all_items[row]) + 'k':
                                    n.text = str(results[row - category_count])
                                    run = n.runs
                                    font = run[0].font
                                    font.bold = False
                                    font.size = Pt(11)
                                    font.name = 'Tahoma'
                            if n.text == str(all_items_index[row]):
                                if str(all_items[row]) not in please:
                                    n.text = str(all_items[row])
                                    run = n.runs
                                    font = run[0].font
                                    font.bold = True
                                    font.size = Pt(13)
                                    font.name = 'Tahoma'
                                    if str(n.text) in categorys:
                                        for igq in range(0, 5):
                                            i.rows[row_index].cells[igq]._tc.get_or_add_tcPr().append(
                                                parse_xml(r'<w:shd {} w:fill="d9d9d9"/>'.format(nsdecls('w'))))
                                    else:
                                        run = n.runs
                                        font = run[0].font
                                        font.bold = True
                                        font.size = Pt(11)
                                        font.name = 'Tahoma'
                                    please.append(str(all_items[row]))

                            if n.text == str(all_items_index[row]) + 'unit':
                                n.text = str(units[row])
                                run1 = n
                                font1 = run1.runs[0].font
                                font1.bold = True
                                font1.size = Pt(12)
                                font1.name = 'Times New Roman'
                            if n.text == str(all_items_index[row]) + 'defult':
                                n.text = str(defults[row])
                                run1 = n
                                font1 = run1.runs[0].font
                                font1.bold = True
                                font1.size = Pt(12)
                                font1.name = 'Times New Roman'
        for iq1 in document.tables:
            for kq1 in iq1.rows:
                for jq1 in kq1.cells:
                    for nq1 in jq1.paragraphs:
                        for myq in range(0, 21):
                            if nq1.text == str(myq):
                                if nq1.runs[0].font.underline == True:
                                    nq1.text = ''
                            if nq1.text == str(myq) + 'unit':
                                if nq1.runs[0].font.underline == True:
                                    nq1.text = ''
                            if nq1.text == str(myq) + 'defult':
                                if nq1.runs[0].font.underline == True:
                                    nq1.text = ''
                            if nq1.text == str(myq) + 'r':
                                if nq1.runs[0].font.underline == True:
                                    nq1.text = ''
        if len(all_items) >= 20:
            # print(results)
            please2 = []
            for i2 in document2.tables:
                for row_index2, k2 in enumerate(i2.rows):
                    for cell_index2, j2 in enumerate(k2.cells):
                        for n2 in j2.paragraphs:
                            if n2.text == 'Date     /    /20':
                                n2.text = f'Date  {year}/{month}/{day}'
                                run = n2.runs
                                font = run[0].font
                                font.name = 'Tahoma'
                                font.bold = True
                                font.size = Pt(12)
                            if n2.text == 'employee1':
                                n2.text = str(word_data[4])
                                run = n2.runs
                                font = run[0].font
                                font.name = 'Times New Roman'
                                font.bold = True
                                font.size = Pt(16)
                            if n2.text == 'employee2':
                                n2.text = str(word_data[5])
                                run = n2.runs
                                font = run[0].font
                                font.name = 'Times New Roman'
                                font.bold = True
                                font.size = Pt(16)
                            if n2.text == 'employeeshahada1':
                                n2.text = str(word_data[6])
                                run = n2.runs
                                font = run[0].font
                                font.name = 'Times New Roman'
                                font.bold = True
                                font.size = Pt(12)
                            if n2.text == 'employeeshahada2':
                                n2.text = str(word_data[7])
                                run = n2.runs
                                font = run[0].font
                                font.name = 'Times New Roman'
                                font.bold = True
                                font.size = Pt(12)
                            if n2.text == 'shopname':
                                n2.text = str(word_data[1])
                                run = n2.runs
                                font = run[0].font
                                font.name = 'Times New Roman'
                                font.bold = True
                                font.size = Pt(26)
                            if n2.text == 'phone1':
                                n2.text = str(word_data[2])
                                run = n2.runs
                                font = run[0].font
                                font.name = 'Times New Roman'
                                font.bold = True
                                font.size = Pt(14)
                            if n2.text == 'phone2':
                                n2.text = str(word_data[3])
                                run = n2.runs
                                font = run[0].font
                                font.name = 'Times New Roman'
                                font.bold = True
                                font.size = Pt(14)
                            if n2.text == 'gps':
                                n2.text = str(word_data[8])
                                run = n2.runs
                                font = run[0].font
                                font.name = 'Times New Roman'
                                font.bold = True
                                font.size = Pt(13)
                            if n2.text == 'client_name':
                                if genus == 1:
                                    n2.text = f'{word_data[9]}'
                                else:
                                    n2.text = f'{word_data[13]}'
                                run = n2.runs
                                font = run[0].font
                                font.name = 'Times New Roman'
                                font.bold = True
                                font.size = Pt(12)
                            if n2.text == 'rname':
                                n2.text = str(name)
                                run = n2.runs
                                font = run[0].font
                                font.name = 'Times New Roman'
                                font.bold = True
                                font.size = Pt(12)
                            if n2.text == 'doctor':
                                if doctor_genus == 'female':
                                    n2.text = f'{word_data[15]}'
                                else:
                                    n2.text = f'{word_data[11]}'
                                run = n2.runs
                                font = run[0].font
                                font.name = 'Times New Roman'
                                font.bold = True
                                font.size = Pt(12)
                            if n2.text == 'rdoctor':
                                n2.text = str(doctor)
                                run = n2.runs
                                font = run[0].font
                                font.name = 'Times New Roman'
                                font.bold = True
                                font.size = Pt(12)
                            if n2.text == 'lqb1':
                                if genus == 1:
                                    n2.text = str(word_data[10])
                                else:
                                    n2.text = str(word_data[14])
                                run = n2.runs
                                font = run[0].font
                                font.name = 'Times New Roman'
                                font.bold = True
                                font.size = Pt(12)
                            if n2.text == 'lqb2':
                                if doctor_genus == 'female':
                                    n2.text = str(word_data[16])
                                else:
                                    n2.text = str(word_data[12])
                                run = n2.runs
                                font = run[0].font
                                font.name = 'Times New Roman'
                                font.bold = True
                                font.size = Pt(12)
                            for row3 in range(20, len(all_items)):
                                if n2.text == str(all_items_index[row3 - 20]) + 'r':
                                    n2.text = str(all_items[row3]) + 'k'
                                    if str(n2.text)[:-1] not in categorys:
                                        run = n2.runs
                                        font = run[0].font
                                        font.bold = True
                                        font.size = Pt(11)
                                        font.name = 'Tahoma'
                                    else:
                                        n2.text = ''
                                    if str(all_items[row3]) in categorys:
                                        category_count += 1
                                    if n2.text == str(all_items[row3]) + 'k':
                                        if str(n2.text)[2:-1] in categorys:
                                            n2.text = ''
                                        else:
                                            n2.text = results[row3 - category_count]
                                            run = n2.runs
                                            font = run[0].font
                                            font.bold = False
                                            font.size = Pt(11)
                                            font.name = 'Tahoma'
                                if n2.text == str(all_items_index[row3 - 20]):
                                    if str(all_items[row3]) not in please2:
                                        n2.text = str(all_items[row3])
                                        run = n2.runs
                                        font = run[0].font
                                        font.bold = True
                                        font.size = Pt(13)
                                        font.name = 'Tahoma'
                                        if n2.text in categorys:
                                            for igq2 in range(0, 5):
                                                i2.rows[row_index2].cells[igq2]._tc.get_or_add_tcPr().append(
                                                    parse_xml(r'<w:shd {} w:fill="d9d9d9"/>'.format(nsdecls('w'))))
                                        else:
                                            font.size = Pt(11)
                                        please2.append(n2.text)
                                    else:
                                        run = n2.runs
                                        font = run[0].font
                                        font.bold = True
                                        font.size = Pt(11)
                                        font.name = 'Tahoma'

                                if n2.text == str(all_items_index[row3 - 20]) + 'unit':
                                    n2.text = str(units[row3])
                                    run1 = n2
                                    font1 = run1.runs[0].font
                                    font1.bold = True
                                    font1.size = Pt(12)
                                    font1.name = 'Times New Roman'
                                if n2.text == str(all_items_index[row3 - 20]) + 'defult':
                                    n2.text = str(defults[row3])
                                    run1 = n2
                                    font1 = run1.runs[0].font
                                    font1.bold = True
                                    font1.size = Pt(12)
                                    font1.name = 'Times New Roman'
            for iq2 in document2.tables:
                for kq2 in iq2.rows:
                    for jq2 in kq2.cells:
                        for nq2 in jq2.paragraphs:
                            for myq2 in range(0, 26):
                                if nq2.text == str(myq2):
                                    if nq2.runs[0].font.underline == True:
                                        nq2.text = ''
                                if nq2.text == str(myq2) + 'unit':
                                    if nq2.runs[0].font.underline == True:
                                        nq2.text = ''
                                if nq2.text == str(myq2) + 'defult':
                                    if nq2.runs[0].font.underline == True:
                                        nq2.text = ''
                                if nq2.text == str(myq2) + 'r':
                                    if nq2.runs[0].font.underline == True:
                                        nq2.text = ''

        if len(all_items) >= 40:
            please3 = []
            for i3 in document3.tables:
                for row_index3, k3 in enumerate(i3.rows):
                    for cell_index3, j3 in enumerate(k3.cells):
                        for n3 in j3.paragraphs:
                            if n3.text == 'Date     /    /20':
                                n3.text = f'Date  {year}/{month}/{day}'
                                run = n3.runs
                                font = run[0].font
                                font.name = 'Tahoma'
                                font.bold = True
                                font.size = Pt(12)
                            if n3.text == 'employee1':
                                n3.text = str(word_data[4])
                                run = n3.runs
                                font = run[0].font
                                font.name = 'Times New Roman'
                                font.bold = True
                                font.size = Pt(16)
                            if n3.text == 'employee2':
                                n3.text = str(word_data[5])
                                run = n3.runs
                                font = run[0].font
                                font.name = 'Times New Roman'
                                font.bold = True
                                font.size = Pt(16)
                            if n3.text == 'employeeshahada1':
                                n3.text = str(word_data[6])
                                run = n3.runs
                                font = run[0].font
                                font.name = 'Times New Roman'
                                font.bold = True
                                font.size = Pt(12)
                            if n3.text == 'employeeshahada2':
                                n3.text = str(word_data[7])
                                run = n3.runs
                                font = run[0].font
                                font.name = 'Times New Roman'
                                font.bold = True
                                font.size = Pt(12)
                            if n3.text == 'shopname':
                                n3.text = str(word_data[1])
                                run = n3.runs
                                font = run[0].font
                                font.name = 'Times New Roman'
                                font.bold = True
                                font.size = Pt(26)
                            if n3.text == 'phone1':
                                n3.text = str(word_data[2])
                                run = n3.runs
                                font = run[0].font
                                font.name = 'Times New Roman'
                                font.bold = True
                                font.size = Pt(14)
                            if n3.text == 'phone2':
                                n3.text = str(word_data[3])
                                run = n3.runs
                                font = run[0].font
                                font.name = 'Times New Roman'
                                font.bold = True
                                font.size = Pt(14)
                            if n3.text == 'gps':
                                n3.text = str(word_data[8])
                                run = n3.runs
                                font = run[0].font
                                font.name = 'Times New Roman'
                                font.bold = True
                                font.size = Pt(13)
                            if n3.text == 'client_name':
                                if genus == 1:
                                    n3.text = f'{word_data[9]}'
                                else:
                                    n3.text = f'{word_data[13]}'
                                run = n3.runs
                                font = run[0].font
                                font.name = 'Times New Roman'
                                font.bold = True
                                font.size = Pt(12)
                            if n3.text == 'rname':
                                n3.text = str(name)
                                run = n3.runs
                                font = run[0].font
                                font.name = 'Times New Roman'
                                font.bold = True
                                font.size = Pt(12)
                            if n3.text == 'doctor':
                                if doctor_genus == 'female':
                                    n3.text = f'{word_data[15]}'
                                else:
                                    n3.text = f'{word_data[11]}'
                                run = n3.runs
                                font = run[0].font
                                font.name = 'Times New Roman'
                                font.bold = True
                                font.size = Pt(12)
                            if n3.text == 'rdoctor':
                                n3.text = str(doctor)
                                run = n3.runs
                                font = run[0].font
                                font.name = 'Times New Roman'
                                font.bold = True
                                font.size = Pt(12)
                            if n3.text == 'lqb1':
                                if genus == 1:
                                    n3.text = str(word_data[10])
                                else:
                                    n3.text = str(word_data[14])
                                run = n3.runs
                                font = run[0].font
                                font.name = 'Times New Roman'
                                font.bold = True
                                font.size = Pt(12)
                            if n3.text == 'lqb2':
                                if doctor_genus == 'female':
                                    n3.text = str(word_data[16])
                                else:
                                    n3.text = str(word_data[12])
                                run = n3.runs
                                font = run[0].font
                                font.name = 'Times New Roman'
                                font.bold = True
                                font.size = Pt(12)
                            for row4 in range(40, len(all_items)):
                                if n3.text == str(all_items_index[row4 - 40]) + 'r':
                                    n3.text = str(all_items[row4]) + 'k'
                                    if str(n3.text)[:-1] not in categorys:
                                        run = n3.runs
                                        font = run[0].font
                                        font.bold = True
                                        font.size = Pt(11)
                                        font.name = 'Tahoma'
                                    else:
                                        n3.text = ''
                                    if str(all_items[row4]) in categorys:
                                        category_count += 1
                                    if n3.text == str(all_items[row4]) + 'k':
                                        if str(n3.text)[2:-1] in categorys:
                                            n3.text = ''
                                        else:
                                            n3.text = results[row4 - category_count-1]
                                            run = n3.runs
                                            font = run[0].font
                                            font.bold = False
                                            font.size = Pt(11)
                                            font.name = 'Tahoma'
                                if n3.text == str(all_items_index[row4 - 40]):
                                    if str(all_items[row4]) not in please3:
                                        n3.text = str(all_items[row4])
                                        run = n3.runs
                                        font = run[0].font
                                        font.bold = True
                                        font.size = Pt(13)
                                        font.name = 'Tahoma'
                                        if n3.text in categorys:
                                            for igq3 in range(0, 5):
                                                i3.rows[row_index3].cells[igq3]._tc.get_or_add_tcPr().append(
                                                    parse_xml(r'<w:shd {} w:fill="d9d9d9"/>'.format(nsdecls('w'))))
                                        else:
                                            font.size = Pt(11)
                                        please3.append(n3.text)
                                    else:
                                        run = n3.runs
                                        font = run[0].font
                                        font.bold = True
                                        font.size = Pt(11)
                                        font.name = 'Tahoma'
                                if n3.text == str(all_items_index[row4 - 40]) + 'unit':
                                    n3.text = str(units[row4])
                                    run1 = n3
                                    font1 = run1.runs[0].font
                                    font1.bold = True
                                    font1.size = Pt(12)
                                    font1.name = 'Times New Roman'
                                if n3.text == str(all_items_index[row4 - 40]) + 'defult':
                                    n3.text = str(defults[row4])
                                    run1 = n3
                                    font1 = run1.runs[0].font
                                    font1.bold = True
                                    font1.size = Pt(12)
                                    font1.name = 'Times New Roman'
            for iq3 in document3.tables:
                for kq3 in iq3.rows:
                    for jq3 in kq3.cells:
                        for nq3 in jq3.paragraphs:
                            for myq3 in range(0, 26):
                                if nq3.text == str(myq3):
                                    nq3.text = ''
                                if nq3.text == str(myq3) + 'unit':
                                    nq3.text = ''
                                if nq3.text == str(myq3) + 'defult':
                                    nq3.text = ''
                                if nq3.text == str(myq3) + 'r':
                                    nq3.text = ''
        document.save(r'%s\result.docx' % save_word_files)
        f.close()
        # if files == 2:
        print('2 files')
        document2.save(r'%s\result2.docx' % save_word_files)
        f2.close()
        if files == 3:
            print('files3')
            document3.save(r'%s\result3.docx' % save_word_files)
            f3.close()
        try:
            if prev == 'T':
                word = client.Dispatch("Word.Application")
                if os.path.exists(r'%s\result.docx' % save_word_files):
                    word.Documents.Open(r'%s\result.docx' % save_word_files)
                if os.path.exists(r'%s\result2.docx' % save_word_files):
                    word.Documents.Open(r'%s\result2.docx' % save_word_files)
                if os.path.exists(r'%s\result3.docx' % save_word_files):
                    word.Documents.Open(r'%s\result3.docx' % save_word_files)
                self.Add_Data_To_history(7, 6)
                self.History()
            else:
                os.system("TASKKILL /F /IM WINWORD.exe")
                word = client.Dispatch("Word.Application")
                if os.path.exists(r'%s\result.docx' % save_word_files):
                    word.Documents.Open(r'%s\result.docx' % save_word_files)
                    word.ActiveDocument.PrintOut(Background=False)
                if os.path.exists(r'%s\result2.docx' % save_word_files):
                    word.Documents.Open(r'%s\result2.docx' % save_word_files)
                    word.ActiveDocument.PrintOut(Background=False)
                if os.path.exists(r'%s\result3.docx' % save_word_files):
                    word.Documents.Open(r'%s\result3.docx' % save_word_files)
                    word.ActiveDocument.PrintOut(Background=False)
                self.Add_Data_To_history(7, 7)
                self.Delete_Files()
        except Exception as e:
            print(e, '6erorr')


def main():
    app = QApplication(sys.argv)
    app.processEvents()
    window = mainapp()
    window.show()
    app.exec_()


if __name__ == '__main__':
    main()
