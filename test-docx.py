from docx import *
from shutil import copyfile

# def Bio_Word(word):
from docx.shared import Pt
# #
# # # # f = open('word/مختبر بغـداد.docx', 'rb')
# # # # f.read()
# # # #
# # # # document = Document(f)
# # # #
# # # # for i in document.tables:
# # # #     for k in i.rows:
# # # #         for j in k.cells:
# # # #
# # # #             for n in j.paragraphs:
# # # #                 if n.text == '     أســم الـمــريــض    :':
# # # #                     n.text = f'     أســم الـمــريــض    :{name}'
# # # #                     n.style.font.name = 'Monotype Koufi'
# # # #                     n.style.font.bold = True
# # # #                     n.style.font.size = Pt(14)
# # # #                 if n.text == 'حـضـرة الـدكتـور    الـفاضـــل :                                                                                            الـمـحـتـرم':
# # # #                     n.text = f'حـضـرة الـدكتـور    الـفاضـــل :{doctor}                                                                                            الـمـحـتـرم'
# # # #                     n.style.font.name = 'Monotype Koufi'
# # # #                     n.style.font.bold = True
# # # #                     n.style.font.size = Pt(14)
# # # #                 if n.text == '0r':
# # # #                     for row in range(0, len(analysts)):
# # # #                         analyst_and_result = {
# # # #                             'analyst': analysts[row],
# # # #                             'result': results[row]
# # # #                         }
# # # #                         if analyst_and_result['analyst'] == 'Toxoplasma IgG':
# # # #                             k = analyst_and_result['result']
# # # #                         n.text =f'  {5}'
# # # #                         n.style.font.name = 'Tahoma'
# # # #                         n.style.font.bold = True
# # # #                         n.style.font.size = Pt(12)
# # # #                 if n.text == '1r':
# # # #                     for row in range(0, len(analysts)):
# # # #                         analyst_and_result = {
# # # #                             'analyst': analysts[row],
# # # #                             'result': results[row]
# # # #                         }
# # # #                         if analyst_and_result['analyst'] == 'Toxoplasma IgM':
# # # #                             k = analyst_and_result['result']
# # # #                         n.text =f'  {5}'
# # # #                         n.style.font.name = 'Tahoma'
# # # #                         n.style.font.bold = True
# # # #                         n.style.font.size = Pt(12)
# # # #                 if n.text == '2r':
# # # #                     for row in range(0, len(analysts)):
# # # #                         analyst_and_result = {
# # # #                             'analyst': analysts[row],
# # # #                             'result': results[row]
# # # #                         }
# # # #                         if analyst_and_result['analyst'] == 'Cytomegalo Virus IgG':
# # # #                             k = analyst_and_result['result']
# # # #                         n.text =f'  {5}'
# # # #                         n.style.font.name = 'Tahoma'
# # # #                         n.style.font.bold = True
# # # #                         n.style.font.size = Pt(12)
# # # #                 if n.text == '3r':
# # # #                     for row in range(0, len(analysts)):
# # # #                         analyst_and_result = {
# # # #                             'analyst': analysts[row],
# # # #                             'result': results[row]
# # # #                         }
# # # #                         if analyst_and_result['analyst'] == 'Cytomegalo Virus IgM':
# # # #                             k = analyst_and_result['result']
# # # #                         n.text =f'  {5}'
# # # #                         n.style.font.name = 'Tahoma'
# # # #                         n.style.font.bold = True
# # # #                         n.style.font.size = Pt(12)
# # # #                 if n.text == '4r':
# # # #                     for row in range(0, len(analysts)):
# # # #                         analyst_and_result = {
# # # #                             'analyst': analysts[row],
# # # #                             'result': results[row]
# # # #                         }
# # # #                         if analyst_and_result['analyst'] == 'Rubella IgG':
# # # #                             k = analyst_and_result['result']
# # # #                         n.text =f'  {5}'
# # # #                         n.style.font.name = 'Tahoma'
# # # #                         n.style.font.bold = True
# # # #                         n.style.font.size = Pt(12)
# # # #                 if n.text == '5r':
# # # #                     for row in range(0, len(analysts)):
# # # #                         analyst_and_result = {
# # # #                             'analyst': analysts[row],
# # # #                             'result': results[row]
# # # #                         }
# # # #                         if analyst_and_result['analyst'] == 'Rubella IgM':
# # # #                             k = analyst_and_result['result']
# # # #                         n.text =f'  {5}'
# # # #                         n.style.font.name = 'Tahoma'
# # # #                         n.style.font.bold = True
# # # #                         n.style.font.size = Pt(12)
# # # #                 if n.text == '6r':
# # # #                     for row in range(0, len(analysts)):
# # # #                         analyst_and_result = {
# # # #                             'analyst': analysts[row],
# # # #                             'result': results[row]
# # # #                         }
# # # #                         if analyst_and_result['analyst'] == 'Anti - Phspholipin IgG':
# # # #                             k = analyst_and_result['result']
# # # #                         n.text =f'  {5}'
# # # #                         n.style.font.name = 'Tahoma'
# # # #                         n.style.font.bold = True
# # # #                         n.style.font.size = Pt(12)
# # # #                 if n.text == '7r':
# # # #                     for row in range(0, len(analysts)):
# # # #                         analyst_and_result = {
# # # #                             'analyst': analysts[row],
# # # #                             'result': results[row]
# # # #                         }
# # # #                         if analyst_and_result['analyst'] == 'Anti - Phspholipin  IgM':
# # # #                             k = analyst_and_result['result']
# # # #                         n.text =f'  {5}'
# # # #                         n.style.font.name = 'Tahoma'
# # # #                         n.style.font.bold = True
# # # #                         n.style.font.size = Pt(12)
# # # #                 if n.text == '8r':
# # # #                     for row in range(0, len(analysts)):
# # # #                         analyst_and_result = {
# # # #                             'analyst': analysts[row],
# # # #                             'result': results[row]
# # # #                         }
# # # #                         if analyst_and_result['analyst'] == 'Anti - Cardiolipin  IgG':
# # # #                             k = analyst_and_result['result']
# # # #                         n.text =f'  {5}'
# # # #                         n.style.font.name = 'Tahoma'
# # # #                         n.style.font.bold = True
# # # #                         n.style.font.size = Pt(12)
# # # #                 if n.text == '9r':
# # # #                     for row in range(0, len(analysts)):
# # # #                         analyst_and_result = {
# # # #                             'analyst': analysts[row],
# # # #                             'result': results[row]
# # # #                         }
# # # #                         if analyst_and_result['analyst'] == 'Anti - Cardiolipin  IgM':
# # # #                             k = analyst_and_result['result']
# # # #                         n.text =f'  {5}'
# # # #                         n.style.font.name = 'Tahoma'
# # # #                         n.style.font.bold = True
# # # #                         n.style.font.size = Pt(12)
# # # #                 if n.text == '10r':
# # # #                     for row in range(0, len(analysts)):
# # # #                         analyst_and_result = {
# # # #                             'analyst': analysts[row],
# # # #                             'result': results[row]
# # # #                         }
# # # #                         if analyst_and_result['analyst'] == 'Herps   IgG':
# # # #                             k = analyst_and_result['result']
# # # #                         n.text =f'  {5}'
# # # #                         n.style.font.name = 'Tahoma'
# # # #                         n.style.font.bold = True
# # # #                         n.style.font.size = Pt(12)
# # # #                 if n.text == '11r':
# # # #                     for row in range(0, len(analysts)):
# # # #                         analyst_and_result = {
# # # #                             'analyst': analysts[row],
# # # #                             'result': results[row]
# # # #                         }
# # # #                         if analyst_and_result['analyst'] == 'Herpes  IgM':
# # # #                             k = analyst_and_result['result']
# # # #                         n.text =f'  {5}'
# # # #                         n.style.font.name = 'Tahoma'
# # # #                         n.style.font.bold = True
# # # #                         n.style.font.size = Pt(12)
# # # #                 if n.text == 'Date:    /     / 20':
# # # #                     n.text = f'Date:   {day} /  {month} / {year}'
# # # #                     n.style.font.name = 'Tahoma'
# # # #                     n.style.font.bold = True
# # # #                     n.style.font.size = Pt(14)
# # # #                 print('1'+n.text+'2')
# # # # document.save('word/مختبر بغـداد.docx')
# # # # f.close()
# # # '''
# # # 1     أســم الـمــريــض    :2
# # # 1حـضـرة الـدكتـور    الـفاضـــل :                                                                                            الـمـحـتـرم2
# # # 1Toxoplasma IgG2
# # # 1Toxoplasma IgM2
# # # 1Cytomegalo Virus IgG2
# # # 1Cytomegalo Virus IgM2
# # # 1Rubella IgG2
# # # 1Rubella IgM2
# # # 1Anti - Phspholipin IgG2
# # # 1Anti - Phspholipin  IgM2
# # # 1Anti - Cardiolipin  IgG2
# # # 1Anti - Cardiolipin  IgM2
# # # 1Herps   IgG2
# # # 1Herpes  IgM2
# # #
# # #
# # # 1T32
# # # 1T42
# # # 1TSH2
# # # 1LH2
# # # 1FSH2
# # # 1Prolactin2
# # # 1Testosterone2
# # #
# # # '''
# f = open('word/GUE latest.docx', 'rb')
# f.read()
#
# document = Document(f)
# for i in document.tables:
#     for k in i.rows:
#         for j in k.cells:
#             for n in j.paragraphs:
#                 if n.text == 'أسـم المريض :       المحترم':
#                     n.text = 'أسـم المريض :5       المحترم'
#                     run = n.runs
#                     font = run[0].font
#                     font.name = 'Monotype Koufi'
#                     font.size = Pt(11)
#                     font.bold=True
#
#
#                 if n.text == 'حضرة الدكتور   :       المحترم':
#                     n.text = f'حضرة الدكتور   : {5}      المحترم'
#                     run = n.runs
#                     font = run[0].font
#                     font.name = 'Monotype Koufi'
#                     font.size = Pt(11)
#                     font.bold = True
#
#
#                 if n.text == 'Appearance :':
#                     n.text = f'Appearance : {"j"}'
#                     run = n.runs
#                     font = run[0].font
#                     font.name = 'Tahoma'
#                     font.size = Pt(14)
#
#
#                 if n.text == 'Reaction      :':
#                     n.text = f'Reaction      : {6}'
#                     run = n.runs
#                     font = run[0].font
#                     font.name = 'Tahoma'
#                     font.size = Pt(14)
#
#
#                 if n.text == 'Albumin       :':
#                     n.text = f'Albumin       : {6}'
#                     run = n.runs
#                     font = run[0].font
#                     font.name = 'Tahoma'
#                     font.size = Pt(14)
#
#
#                 if n.text == 'Sugar          :':
#                     n.text = f'Sugar          : {6}'
#                     run = n.runs
#                     font = run[0].font
#                     font.name = 'Tahoma'
#                     font.size = Pt(14)
#
#
#                 if n.text == 'RBCs         :':
#                     n.text = f'RBCs         : {6}'
#                     run = n.runs
#                     font = run[0].font
#                     font.name = 'Tahoma'
#                     font.size = Pt(14)
#
#
#                 if n.text == 'Pus cells    :':
#                     n.text = f'Pus cells    : {6}'
#                     run = n.runs
#                     font = run[0].font
#                     font.name = 'Tahoma'
#                     font.size = Pt(14)
#
#
#                 if n.text == 'Epith .cells :':
#                     n.text = f'Epith .cells : {6}'
#                     run = n.runs
#                     font = run[0].font
#                     font.name = 'Tahoma'
#                     font.size = Pt(14)
#
#
#                 if n.text == 'Crystals     :':
#                     n.text = f'Crystals     : {6}'
#                     run = n.runs
#                     font = run[0].font
#                     font.name = 'Tahoma'
#                     font.size = Pt(14)
#
#
#                 if n.text == 'Casts        :':
#                     n.text = f'Casts        : {6}'
#                     run = n.runs
#                     font = run[0].font
#                     font.name = 'Tahoma'
#                     font.size = Pt(14)
#
#
#                 if n.text == 'Other        :':
#                     n.text = f'Other        : {6}'
#                     run = n.runs
#                     font = run[0].font
#                     font.name = 'Tahoma'
#                     font.size = Pt(14)
#
#                 if n.text == 'Date:    /     / 20':
#                     n.text = f'Date:   {10} /  {3} / {2120}'
#                     run = n.runs
#                     font = run[0].font
#                     font.name = 'Tahoma'
#                     font.bold = True
#                     font.size = Pt(12)
#
#
# document.save('word/GUE latest.docx')
# f.close()
from win32com import client
import time
f = open('word/bio latest17.docx', 'rb')
f.read()

document = Document(f)
for i in document.tables:
    for k in i.rows:
        for j in k.cells:
            for n in j.paragraphs:
                if n.text == 'أسـم المريض :       المحترم':
                    n.text = f'أسـم المريض :{"kf"}       المحترم'
                    run=n.runs
                    font = run[0].font
                    font.name = 'Monotype Koufi'
                    font.bold = True
                    font.size = Pt(11)
                if n.text == 'حضرة الدكتور   :       المحترم':
                    n.text = f'حضرة الدكتور   : {"doctor"}      المحترم'
                    run = n.runs
                    font = run[0].font
                    font.name = 'Monotype Koufi'
                    font.bold = True
                    font.size = Pt(11)

                if n.text == 'Random  blood sugar :':
                            n.text = f'Random  blood sugar : {3}'
                            run = n.runs
                            font = run[0].font
                            font.bold=True
                            font.size = Pt(11)
                            font.name = 'Tahoma'
                if n.text == 'Blood Urea               :':

                            n.text = f'Blood Urea               : {3}'
                            run = n.runs
                            font = run[0].font
                            font.bold = True
                            font.size = Pt(11)
                            font.name = 'Tahoma'
                if n.text == 'S. Creatinin               :':
                            n.text = f'S. Creatinin               : {3}'
                            run = n.runs
                            font = run[0].font
                            font.bold = True
                            font.size = Pt(11)
                            font.name = 'Tahoma'
                if n.text == 'S. Uric acid                  :':
                            n.text = f'S. Uric acid                  : {3}'
                            run = n.runs
                            font = run[0].font
                            font.bold = True
                            font.size = Pt(11)
                            font.name = 'Tahoma'
                if n.text == 'S. Cholesterol            :':
                            n.text = f'S. Cholesterol            : {3}'
                            run = n.runs
                            font = run[0].font
                            font.bold = True
                            font.size = Pt(11)
                            font.name = 'Tahoma'
                if n.text == 'S. Triglycerid             :':
                            n.text = f'S. Triglycerid             : {3}'
                            run = n.runs
                            font = run[0].font
                            font.bold = True
                            font.size = Pt(11)
                            font.name = 'Tahoma'
                if n.text == 'Total serum Bilirubin:':
                            n.text = f'Total serum Bilirubin: {3}'
                            run = n.runs
                            font = run[0].font
                            font.bold = True
                            font.size = Pt(11)
                            font.name = 'Tahoma'
                if n.text == 'S.Calcium :':
                            n.text = f'S.Calcium : {3}'
                            run = n.runs
                            font = run[0].font
                            font.bold = True
                            font.size = Pt(11)
                            font.name = 'Tahoma'
                if n.text == 'Vitamin D              :':
                            n.text = f'Vitamin D              : {3}'
                            run = n.runs
                            font = run[0].font
                            font.bold = True
                            font.size = Pt(11)
                            font.name = 'Tahoma'
                if n.text == 'Date:    /     / 20':
                    n.text = f'Date:   {10} /  {2} / {8239}'
                    run = n.runs
                    font = run[0].font
                    font.bold = True
                    font.size = Pt(11)
                    font.name = 'Tahoma'
document.save('word/bio latest17.docx')
f.close()

word = client.Dispatch("Word.Application")

word.Documents.Open(r'F:\برنامج التحليلات\word\bio latest17.docx')
# word.Application.ActivePrinter = "PostScript"
word.ActiveDocument.PrintOut()
time.sleep(2)
word.ActiveDocument.Close()
# # f = open('word/GSE latest.docx', 'rb')
# # f.read()
# #
# # document = Document(f)
# # for i in document.tables:
# #     for k in i.rows:
# #         for j in k.cells:
# #             for n in j.paragraphs:
# #                 if n.text == 'أسـم المريض :       المحترم':
# #                     n.text = f'أسـم المريض :{"name"}       المحترم'
# #                     run = n.runs
# #                     font = run[0].font
# #                     font.name = 'Monotype Koufi'
# #                     font.bold = True
# #                     font.size = Pt(11)
# #                 if n.text == 'حضرة الدكتور   :       المحترم':
# #                     n.text = f'حضرة الدكتور   : {"doctor"}      المحترم'
# #                     run = n.runs
# #                     font = run[0].font
# #                     font.name = 'Monotype Koufi'
# #                     font.bold = True
# #                     font.size = Pt(11)
# #                 if n.text == 'Color:':
# #
# #                             n.text = f'Color: {"k"}'
# #                             run = n.runs
# #                             font = run[0].font
# #                             font.bold = True
# #                             font.size = Pt(14)
# #                             font.name = 'Tahoma'
# #                 if n.text == 'Consistency:':
# #
# #
# #                             n.text = f'Consistency: {"k"}'
# #                             run = n.runs
# #                             font = run[0].font
# #                             font.bold = True
# #                             font.size = Pt(14)
# #                             font.name = 'Tahoma'
# #                 if n.text == 'R.B.Cs:':
# #
# #
# #                             n.text = f'R.B.Cs:: {"k"}'
# #                             run = n.runs
# #                             font = run[0].font
# #                             font.bold = True
# #                             font.size = Pt(14)
# #                             font.name = 'Tahoma'
# #                 if n.text == 'Pus cells:':
# #
# #
# #                             n.text = f'Pus cells: {"k"}'
# #                             run = n.runs
# #                             font = run[0].font
# #                             font.bold = True
# #                             font.size = Pt(14)
# #                             font.name = 'Tahoma'
# #                 if n.text == 'E. Histolytica:':
# #
# #                             n.text = f'E. Histolytica: {"k"}'
# #                             run = n.runs
# #                             font = run[0].font
# #                             font.bold = True
# #                             font.size = Pt(14)
# #                             font.name = 'Tahoma'
# #                 if n.text == 'G. Lembilia:':
# #
# #                             n.text = f'G. Lembilia: {"k"}'
# #                             run = n.runs
# #                             font = run[0].font
# #                             font.bold = True
# #                             font.size = Pt(14)
# #                             font.name = 'Tahoma'
# #                 if n.text == 'Ova:':
# #
# #                             n.text = f'Ova: {"k"}'
# #                             run = n.runs
# #                             font = run[0].font
# #                             font.bold = True
# #                             font.size = Pt(14)
# #                             font.name = 'Tahoma'
# #                 if n.text == 'Other:':
# #
# #                             n.text = f'Other: {"k"}'
# #                             run = n.runs
# #                             font = run[0].font
# #                             font.bold = True
# #                             font.size = Pt(14)
# #                             font.name = 'Tahoma'
# #
# #
# #                 if n.text == 'Date:    /     / 20':
# #                     n.text = f'Date:   {"day"} /  {"month"} / {"year"}'
# #                     run = n.runs
# #                     font = run[0].font
# #                     font.bold = True
# #                     font.size = Pt(12)
# #                     font.name = 'Tahoma'
# # document.save('word/GSE latest.docx')
# # f.close()
# # f = open('word/هرمونات مشترك latest.docx', 'rb')
# # f.read()
# #
# # document = Document(f)
# # k='9'
# # for i in document.tables:
# #     for k in i.rows:
# #         for j in k.cells:
# #
# #             for n in j.paragraphs:
# #                 if n.text == '     أســم الـمــريــض    :':
# #                     n.text = f'     أســم الـمــريــض    :{"name"}'
# #                     run = n.runs
# #                     font = run[0].font
# #                     font.bold = True
# #                     font.size = Pt(12)
# #                     font.name = 'Monotype Koufi'
# #                 if n.text == 'حـضـرة الـدكتـور    الـفاضـــل :                                                                                            الـمـحـتـرم':
# #                     n.text = f'حـضـرة الـدكتـور    الـفاضـــل :{"doctor"}                                                                                            الـمـحـتـرم'
# #                     run = n.runs
# #                     font = run[0].font
# #                     font.bold = True
# #                     font.size = Pt(12)
# #                     font.name = 'Monotype Koufi'
# #                 if n.text == '0r':
# #                         n.text = f'  {"l"}'
# #                         run = n.runs
# #                         font = run[0].font
# #                         font.bold = True
# #                         font.size = Pt(12)
# #                         font.name = 'Tahoma'
# #                 if n.text == '1r':
# #
# #                         n.text =f'90{"l"}'
# #                         run = n.runs
# #                         font = run[0].font
# #                         font.bold = True
# #                         font.size = Pt(12)
# #                         font.name = 'Tahoma'
# #                 if n.text == '2r':
# #
# #                         n.text = f'  {"k"}'
# #                         run = n.runs
# #                         font = run[0].font
# #                         font.bold = True
# #                         font.size = Pt(12)
# #                         font.name = 'Tahoma'
# #                 if n.text == '3r':
# #
# #                         n.text = f'  {"k"}'
# #                         run = n.runs
# #                         font = run[0].font
# #                         font.bold = True
# #                         font.size = Pt(12)
# #                         font.name = 'Tahoma'
# #                 # if n.text == '4r':
# #                 #
# #                 #         n.text = f'  {k}'
# #                 #         run = n.runs
# #                 #         font = run[0].font
# #                 #         font.bold = True
# #                 #         font.size = Pt(12)
# #                 #         font.name = 'Tahoma'
# #                 if n.text == '5r':
# #
# #                         n.text = f'  {"k"}'
# #                         run = n.runs
# #                         font = run[0].font
# #                         font.bold = True
# #                         font.size = Pt(12)
# #                         font.name = 'Tahoma'
# #                 if n.text == '6r':
# #
# #                         n.text = f'  {"k"}'
# #                         run = n.runs
# #                         font = run[0].font
# #                         font.bold = True
# #                         font.size = Pt(12)
# #                         font.name = 'Tahoma'
# #                 if n.text == '7r':
# #
# #                         n.text = f'  {"k"}'
# #                         run = n.runs
# #                         font = run[0].font
# #                         font.bold = True
# #                         font.size = Pt(12)
# #                         font.name = 'Tahoma'
# #                 if n.text == '8r':
# #
# #                         n.text = f'  {"k"}'
# #                         run = n.runs
# #                         font = run[0].font
# #                         font.bold = True
# #                         font.size = Pt(12)
# #                         font.name = 'Tahoma'
# #                 if n.text == '9r':
# #
# #                         n.text = f'  {"k"}'
# #                         run = n.runs
# #                         font = run[0].font
# #                         font.bold = True
# #                         font.size = Pt(12)
# #                         font.name = 'Tahoma'
# #                 if n.text == '10r':
# #
# #                         n.text = f'  {"k"}'
# #                         run = n.runs
# #                         font = run[0].font
# #                         font.bold = True
# #                         font.size = Pt(12)
# #                         font.name = 'Tahoma'
# #                 if n.text == '11r':
# #
# #                         n.text = f'  {"k"}'
# #                         run = n.runs
# #                         font = run[0].font
# #                         font.bold = True
# #                         font.size = Pt(12)
# #                         font.name = 'Tahoma'
# #                 if n.text == 'r0r':
# #
# #                         n.text = f'  {"k"}'
# #                         run = n.runs
# #                         font = run[0].font
# #                         font.bold = True
# #                         font.size = Pt(12)
# #                         font.name = 'Tahoma'
# #
# #                 if n.text == 'r1r':
# #
# #                         n.text = f'  {"k"}'
# #                         run = n.runs
# #                         font = run[0].font
# #                         font.bold = True
# #                         font.size = Pt(12)
# #                         font.name = 'Tahoma'
# #
# #                 if n.text == 'r2r':
# #
# #                         n.text = f'  {"k"}'
# #                         run = n.runs
# #                         font = run[0].font
# #                         font.bold = True
# #                         font.size = Pt(12)
# #                         font.name = 'Tahoma'
# #
# #                 if n.text == 'r3r':
# #                         n.text = f'  {"k"}'
# #                         run = n.runs
# #                         font = run[0].font
# #                         font.bold = True
# #                         font.size = Pt(12)
# #                         font.name = 'Tahoma'
# #
# #                 # if n.text == 'r4r':
# #                 #
# #                 #         n.text = f'  {k}'
# #                 #         run = n.runs
# #                 #         font = run[0].font
# #                 #         font.bold = True
# #                 #         font.size = Pt(12)
# #                 #         font.name = 'Tahoma'
# #
# #                 if n.text == 'r5r':
# #
# #                         n.text = f'  {"k"}'
# #                         run = n.runs
# #                         font = run[0].font
# #                         font.bold = True
# #                         font.size = Pt(12)
# #                         font.name = 'Tahoma'
# #                 if n.text == 'r6r':
# #
# #                         n.text = f'  {"l"}'
# #                         run = n.runs
# #                         font = run[0].font
# #                         font.bold = True
# #                         font.size = Pt(12)
# #                         font.name = 'Tahoma'
# #                 if n.text == 'Date:    /     / 20':
# #                     n.text = f'Date:   {"day"} /  {"month"} / {"year"}'
# #                     run = n.runs
# #                     font = run[0].font
# #                     font.bold = True
# #                     font.size = Pt(14)
# #                     font.name = 'Tahoma'
# #                 if n.text =='4r' or n.text=='r4r':
# #                     n.text=''
# # document.save('word/هرمونات مشترك latest.docx')
# # f.close()
'''
                if if_print:
                    print('its start')
                    if os.path.exists(r'F:\برنامج التحليلات\bio latest17.docx'):
                        os.remove(r'F:\برنامج التحليلات\bio latest17.docx')
                if if_print:
                    if os.path.exists(r'F:\برنامج التحليلات\GSE latest.docx'):
                      os.remove(r'F:\برنامج التحليلات\GSE latest.docx')
                if if_print:
                    if os.path.exists(r'F:\برنامج التحليلات\GUE latest.docx'):
                        os.remove(r'F:\برنامج التحليلات\GUE latest.docx')
                if if_print:
                    if os.path.exists(r'F:\برنامج التحليلات\hematology latest.docx'):
                        os.remove(r'F:\برنامج التحليلات\hematology latest.docx')
                if if_print:
                    if os.path.exists(r'F:\برنامج التحليلات\هرمونات مشترك latest.docx'):
                        os.remove(r'F:\برنامج التحليلات\هرمونات مشترك latest.docx')
'''