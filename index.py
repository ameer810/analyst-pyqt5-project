import datetime
import os
import sys
import time
import MySQLdb
# import pyautogui
from PyQt5 import QtGui
from PyQt5.QtCore import *
from PyQt5.QtGui import *
from PyQt5.QtWidgets import *
from PyQt5.uic import loadUiType

from PyQt5.uic.properties import QtCore
from docx import *
# from shutil import copyfile
from docx.shared import Pt
# import smtplib
# from email.mime.text import MIMEText
# from email.mime.multipart import MIMEMultipart
from win32com import client

# from mymain import Ui_MainWindow as main_wind
FORM_CLASS, _ = loadUiType("design.ui")
show_all_sales_in_clients_page = False
user_id = 4
client_id_glob = 0
chick_if_add_new = False
if_print = False
analysts_name_glo = []
clients_name_glo = []
clients_name_glo_clients_page = []
two_files=False

# GSE_row1=[]
class mainapp(QMainWindow, FORM_CLASS):
    def __init__(self, parent=None):
        super(mainapp, self).__init__(parent)
        QMainWindow.__init__(self)
        self.setupUi(self)
        self.tabWidget.tabBar().setVisible(False)
        self.DB_Connect()
        self.Delete_Files()
        self.handel_buttons()
        self.Show_All_The_Sales()
        self.Show_all_analysts_in_combo()
        # self.Show_all_buys()
        # self.History()
        # self.groupBox.setEnabled(False)
        self.Show_paths()
        self.add_Analyst_to_list()
        self.Auto_complete_combo()
        self.add_client_to_list()
        self.Auto_complete_combo2()
        self.add_client_to_list4()
        self.Auto_complete_combo4()

    def DB_Connect(self):
        self.db = MySQLdb.connect(host='localhost', user='root', password='12345', db='tahlel', charset="utf8",
                                  use_unicode=True, port=3306)
        self.cur = self.db.cursor()

    def handel_buttons(self):
        self.pushButton_15.clicked.connect(self.Light_Blue_Theme)
        self.pushButton_9.clicked.connect(self.Dark_Orange_Theme)
        self.pushButton_13.clicked.connect(self.Dark_Blue_Theme)
        self.pushButton_14.clicked.connect(self.Dark_Theme)
        self.pushButton_11.clicked.connect(self.Dark_Gray_Theme)
        self.pushButton.clicked.connect(self.Open_Sales_Page)
        # self.pushButton_6.clicked.connect(self.Open_Login_Page)
        self.pushButton_4.clicked.connect(self.Open_Settings_Page)
        self.pushButton_12.clicked.connect(self.Open_Print_Page)
        # self.pushButton_3.clicked.connect(self.Open_History_Page)
        self.pushButton_5.clicked.connect(self.Open_clients_Page)
        # self.pushButton_2.clicked.connect(self.Open_Analyse_Page)
        # self.pushButton_8.clicked.connect(self.Open_ResetPassword_Page)
        self.pushButton_17.clicked.connect(self.Sales_Page)
        self.pushButton_30.clicked.connect(self.get_client_id)
        self.pushButton_31.clicked.connect(self.Chick_analyst_category)
        self.pushButton_32.clicked.connect(self.get_total_price)
        # self.pushButton_33.clicked.connect(self.Show_analyst_in_Edit_Or_Delete)
        self.pushButton_29.clicked.connect(self.Clients_Page)
        # self.pushButton_16.clicked.connect(self.Add_Buys)
        # self.pushButton_20.clicked.connect(self.Show_All_The_Analysts)
        # self.pushButton_27.clicked.connect(self.Edit_Analyst)
        # self.pushButton_7.clicked.connect(self.Log_In_Chieck)
        # self.pushButton_28.clicked.connect(self.Delete_Analyst)
        # self.pushButton_22.clicked.connect(self.Delete_All_History_Data)
        self.pushButton_18.clicked.connect(self.Print_Sale_Data)
        self.pushButton_19.clicked.connect(self.Search_In_All_Sales)
        # self.pushButton_21.clicked.connect(self.Search_In_History)
        self.pushButton_35.clicked.connect(self.clear_data_in_sales)
        # self.pushButton_10.clicked.connect(self.Reset_password)
        # self.pushButton_26.clicked.connect(self.Add_Analyst)
        self.pushButton_34.clicked.connect(self.Preview)
        self.pushButton_23.clicked.connect(self.Print_empty_papers)
        self.pushButton_24.clicked.connect(self.Add_Path)
        self.pushButton_36.clicked.connect(self.Show_All_Clients)
        self.my_def()

    def Show_all_clients_without_search(self):
        self.cur.execute(''' SELECT client_name FROM addclient ORDER BY -date ''')
        analyst_data = self.cur.fetchall()
        all_names = {}
        all_data = []
        id = 0

        for item in analyst_data:
            if item[0] not in all_names:
                all_names[item[0]] = 0

        for s in all_names.keys():
            self.cur.execute(''' SELECT id FROM addclient WHERE client_name=%s ORDER BY id ''', (s,))
            analyst_data2 = self.cur.fetchall()
            all_names[s] = analyst_data2[0][0]
            id = analyst_data2[0][0]
            self.cur.execute(''' SELECT * FROM addclient WHERE id=%s ORDER BY -date ''', (id,))
            latest_data = self.cur.fetchall()
            data = [{'name': s, 'value': latest_data}]
            all_data.append(data)
        self.tableWidget_2.setRowCount(0)
        self.tableWidget_2.insertRow(0)
        l_data = ()
        for ih in all_data:
            hs_data = ih[0]['value']
            l_data += tuple(hs_data)
        for row, form in enumerate(l_data):
            for col, item in enumerate(form):
                if col == 3:
                    self.tableWidget_2.setItem(row, col, QTableWidgetItem(str(l_data[row][4])))
                else:
                    print('d')
                    self.tableWidget_2.setItem(row, col, QTableWidgetItem(str(item)))
                col += 1

            row_pos2 = self.tableWidget_2.rowCount()
            self.tableWidget_2.insertRow(row_pos2)

    def my_def(self):
        print('sdge')
        self.cur.execute(''' SELECT client_name FROM addclient ORDER BY -date ''')
        data = self.cur.fetchall()
        all_names = {}
        all_data = []
        id = 0

        for item in data:
            if item[0] not in all_names:
                all_names[item[0]] = 0

        for s in all_names.keys():
            self.cur.execute(''' SELECT id FROM addclient WHERE client_name=%s ORDER BY id ''', (s,))
            analyst_data2 = self.cur.fetchall()
            all_names[s] = analyst_data2[0][0]
            id = analyst_data2[0][0]
            self.cur.execute(''' SELECT * FROM addclient WHERE id=%s ORDER BY -date ''', (id,))
            latest_data = self.cur.fetchall()
            data = [{'name': s, 'value': latest_data}]
            all_data.append(data)
        # self.tableWidget_2.setRowCount(0)
        # self.tableWidget_2.insertRow(0)
        l_data = ()
        for ih in all_data:
            hs_data = ih[0]['value']
            l_data += tuple(hs_data)
        # print(l_data, 'k')
        for row, form in enumerate(l_data):
            for col, item in enumerate(form):
                if col == 3:
                    self.tableWidget_2.setItem(row, col, QTableWidgetItem(str(l_data[row][4])))
                else:
                    self.tableWidget_2.setItem(row, col, QTableWidgetItem(str(item)))
                col += 1
            row_pos = self.tableWidget_2.rowCount()
            self.tableWidget_2.insertRow(row_pos)

    def add_Analyst_to_list(self):
        self.cur.execute(''' SELECT name FROM addanalyst ''')
        data = self.cur.fetchall()
        for i in data:
            if i[0] not in analysts_name_glo:
                analysts_name_glo.append(i[0])

    def Auto_Complete(self, model):
        model.setStringList(analysts_name_glo)

    def Auto_complete_combo(self):
        combo = self.comboBox_16
        completer = QCompleter()
        combo.setCompleter(completer)
        model = QStringListModel()
        completer.setModel(model)
        self.Auto_Complete(model)

    def add_client_to_list(self):
        self.cur.execute(''' SELECT client_name FROM addclient WHERE DATE(date)=%s ''', (datetime.date.today(),))
        data = self.cur.fetchall()
        for i in data:
            if i[0] not in clients_name_glo:
                clients_name_glo.append(i[0])

    def Auto_Complete2(self, model):
        model.setStringList(clients_name_glo)

    def Auto_complete_combo2(self):
        combo = self.lineEdit_25, self.lineEdit_27
        completer = QCompleter()
        for i in combo:
            i.setCompleter(completer)
            model = QStringListModel()
            completer.setModel(model)
            self.Auto_Complete2(model)

    def add_client_to_list4(self):
        self.cur.execute(''' SELECT client_name FROM addclient''')
        data = self.cur.fetchall()
        for i in data:
            if i[0] not in clients_name_glo_clients_page:
                clients_name_glo_clients_page.append(i[0])

    def Auto_Complete4(self, model):
        model.setStringList(clients_name_glo_clients_page)

    def Auto_complete_combo4(self):
        combo = self.lineEdit_27
        completer = QCompleter()
        combo.setCompleter(completer)
        model = QStringListModel()
        completer.setModel(model)
        self.Auto_Complete4(model)

    def Delete_Files(self):
        self.cur.execute(''' SELECT * FROM paths WHERE id=1 ''')
        mydata = self.cur.fetchone()
        word_files = mydata[1]
        save_word_files = mydata[2]
        try:
            if os.path.exists(r'%s\bio and hemo2.docx' % save_word_files):
                os.remove(r'%s\bio and hemo2.docx' % save_word_files)

            if os.path.exists(r'%s\GSE latest.docx' % save_word_files):
                os.remove(r'%s\GSE latest.docx' % save_word_files)

            if os.path.exists(r'%s\SFA latest.docx' % save_word_files):
                os.remove(r'%s\SFA latest.docx' % save_word_files)

            if os.path.exists(r'%s\GUE latest.docx' % save_word_files):
                os.remove(r'%s\GUE latest.docx' % save_word_files)

            if os.path.exists(r'%s\bio and hemo.docx' % save_word_files):
                os.remove(r'%s\bio and hemo.docx' % save_word_files)

            if os.path.exists(r'%s\هرمونات مشترك latest.docx' % save_word_files):
                os.remove(r'%s\هرمونات مشترك latest.docx' % save_word_files)
        except Exception as e:
            word = client.Dispatch("Word.Application")
            word.ActiveDocument.Close()
            self.Delete_Files()
            print(e)

    def Show_paths(self):
        self.cur.execute(''' SELECT * FROM paths WHERE id=1 ''')
        mydata = self.cur.fetchone()
        word_files = mydata[1]
        save_word_files = mydata[2]
        self.lineEdit_14.setText(word_files)
        self.lineEdit_15.setText(save_word_files)

    def Add_Path(self):
        files_path = self.lineEdit_14.text()
        save_files_path = self.lineEdit_15.text()
        self.cur.execute(''' UPDATE paths SET file_path=%s,save_file_path=%s WHERE id=1''',
                         (files_path, save_files_path))
        self.db.commit()
        QMessageBox.information(self, '', 'تم تطبيق المعلومات بنجاح')
        self.Show_paths()

    def Print_empty_papers(self):
        self.cur.execute(''' SELECT * FROM paths WHERE id=1 ''')
        mydata = self.cur.fetchone()
        word_files = mydata[1]
        save_word_files = mydata[2]
        paper_type = self.comboBox_27.currentIndex()
        copyes_num = self.spinBox_4.value()
        word = client.Dispatch("Word.Application")

        category = ''
        file = ''
        if paper_type == 0:
            category = 'bio'
            file = 'bio org.docx'
        if paper_type == 1:
            category = 'GSE'
            file = 'GSE org.docx'
        if paper_type == 2:
            category = 'GUE'
            file = 'GUE org.docx'
        if paper_type == 3:
            category = 'hematology'
            file = 'hematology org.docx'
        if paper_type == 5:
            category = 'SFA'
            file = 'SFA org.docx'
        if paper_type == 4:
            category = 'هرمونات مشترك'
            file = 'هرمونات مشترك الاصلي.docx'
        word.Documents.Open(r'%s\%s' % (word_files, file))
        for ihjk in range(1, copyes_num + 1):
            word.ActiveDocument.PrintOut(Background=False)
        QMessageBox.information(self, 'info', 'تمت الطباعة بنجاح')

        # time.sleep(2)
        # pyautogui.press('enter')
        # time.sleep(1.5)
        warning = QMessageBox.warning(self, '',
                                      "هل قمت بطباعة الملفات؟\n لا تضغط نعم الا لو قمت بطباعته",
                                      QMessageBox.Yes | QMessageBox.No)
        if warning == QMessageBox.Yes:
            move = 1
            try:
                word.ActiveDocument.Close()
            except:
                pass
        # word.ActiveDocument.Close()

    def Preview(self):
        global if_print
        self.Print_Sale_Data('T')
        self.cur.execute(''' SELECT * FROM paths WHERE id=1 ''')
        mydata = self.cur.fetchone()
        word_files = mydata[1]
        save_word_files = mydata[2]
        warning = QMessageBox.warning(self, '',
                                      "هل قمت بطباعة جميع الملفات التي تمت معاينتها؟\n لا تضغط نعم الا لو قمت بطباعتها",
                                      QMessageBox.Yes | QMessageBox.No)
        if warning == QMessageBox.Yes:
            try:
                if_print = True
                if os.path.exists(r'%s\bio latest17.docx' % save_word_files):
                    os.remove(r'%s\bio latest17.docx' % save_word_files)

                if os.path.exists(r'%s\GSE latest.docx' % save_word_files):
                    os.remove(r'%s\GSE latest.docx' % save_word_files)

                if os.path.exists(r'%s\SFA latest.docx' % save_word_files):
                    os.remove(r'%s\SFA latest.docx' % save_word_files)

                if os.path.exists(r'%s\GUE latest.docx' % save_word_files):
                    os.remove(r'%s\GUE latest.docx' % save_word_files)

                if os.path.exists(r'%s\hematology latest.docx' % save_word_files):
                    os.remove(r'%s\hematology latest.docx' % save_word_files)

                if os.path.exists(r'%s\هرمونات مشترك latest.docx' % save_word_files):
                    os.remove(r'%s\هرمونات مشترك latest.docx' % save_word_files)
            except:
                word = client.Dispatch("Word.Application")
                word.ActiveDocument.Close()
                self.Delete_Files()

    # def Reset_password(self):
    #     user_name = self.lineEdit_7.text()
    #     self.cur.execute(''' SELECT * FROM adduser ''')
    #     data = self.cur.fetchall()
    #     ruser_name = ''
    #     a = 0
    #     for row in data:
    #         if row[1] == user_name:
    #             ruser_name = row[1]
    #         else:
    #             a = 5
    #     if a == 5:
    #         QMessageBox.information(self, 'info', 'اسم المستخدم الذي ادخلته غير صحيح')
    #     self.cur.execute(''' SELECT user_email,userpassword FROM adduser WHERE user_name=%s ''', (ruser_name,))
    #     email_data = self.cur.fetchone()
    # email = "ameersaad810@gmail.com" # the email where you sent the email
    # password = "aahmpredtiddvxlo"
    # send_to_email = email_data[0] # for whom
    # subject = "n"
    # message = f'Hello.\n your password is {email_data[1]} '
    # msg = MIMEMultipart()
    # msg["From"] = email
    # msg["To"] = send_to_email
    # msg["Subject"] = subject
    #
    # msg.attach(MIMEText(message, 'plain'))
    #
    # server = smtplib.SMTP("smtp.gmail.com", 587)
    # server.starttls()
    # server.login(email, password)
    # text = msg.as_string()
    # server.sendmail(email, send_to_email, text)
    # server.quit()
    # print('ok')

    def clear_data_in_sales(self):
        self.tableWidget_5.setRowCount(0)
        self.tableWidget_5.insertRow(0)
        self.lineEdit_20.setText('')
        self.lineEdit_21.setText('')
        self.spinBox_7.setValue(20)
        self.doubleSpinBox_7.setValue(0)
        self.comboBox_14.setCurrentIndex(0)
        self.comboBox_16.setCurrentIndex(0)
        self.comboBox_17.setCurrentIndex(0)
        self.comboBox_15.setEnabled(True)
        self.comboBox_15.setCurrentIndex(0)
        self.textEdit.setPlainText('')
        self.lineEdit_20.setEnabled(True)
        self.spinBox_7.setEnabled(True)
        self.comboBox_14.setEnabled(True)
        self.textEdit.setEnabled(True)
        self.spinBox.setValue(0)

    def Show_All_Clients(self):
        global show_all_sales_in_clients_page
        client_name = self.lineEdit_27.text()
        if client_name != '0':
            self.cur.execute(''' SELECT * FROM addclient WHERE client_name=%s ORDER BY id''', (client_name,))
        else:
            self.cur.execute(''' SELECT * FROM addclient ORDER BY -date''')
        analyst_data = self.cur.fetchall()
        self.tableWidget_2.setRowCount(0)
        self.tableWidget_2.insertRow(0)
        try:
            if client_name != '0' and client_name != '':
                for i, k in enumerate(analyst_data[0]):
                    if i == 3:
                        self.tableWidget_2.setItem(0, i, QTableWidgetItem(str(analyst_data[0][4])))
                    else:
                        self.tableWidget_2.setItem(0, i, QTableWidgetItem(str(k)))
            else:
                show_all_sales_in_clients_page = True
                self.Show_all_clients_without_search()
        except:
            pass
            # self.tableWidget_2.setRowCount(0)
            # self.tableWidget_2.insertRow(0)

    def Search_In_History(self):
        actionsd = self.comboBox_24.currentIndex()
        tabley = self.comboBox_20.currentIndex()
        if actionsd != 0 and tabley == 0:
            try:
                self.cur.execute(f'SELECT action,tabled,id,dates FROM his WHERE action = {actionsd}')
            except Exception as e:
                print(e)
        elif tabley != 0 and actionsd == 0:
            try:
                self.cur.execute(f'SELECT action,tabled,id,dates FROM his WHERE tabled={tabley}')
            except Exception as e:
                print(e)
        elif actionsd != 0 and tabley != 0:
            try:
                self.cur.execute(
                    f' SELECT action,tabled,id,dates FROM his WHERE action={actionsd} AND tabled={tabley} ')

            except Exception as e:
                print(e)
        else:
            try:
                self.cur.execute(
                    f' SELECT action,tabled,id,dates FROM his')

            except Exception as e:
                print(e)
        data = self.cur.fetchall()
        self.tableWidget_8.setRowCount(0)
        self.tableWidget_8.insertRow(0)
        for row, form in enumerate(data):
            for col, item in enumerate(form):
                if col == 0:
                    sql = '''SELECT uid FROM his WHERE id =%s'''
                    self.cur.execute(sql, [data[row][2]])
                    datas = self.cur.fetchone()
                    sql = '''SELECT user_name FROM adduser WHERE id =%s'''
                    self.cur.execute(sql, [datas[0]])
                    user_name = self.cur.fetchone()
                    self.tableWidget_8.setItem(row, col, QTableWidgetItem(str(user_name[0])))
                if col == 1:
                    action = ''
                    if data[row][0] == 1:
                        action = 'تسجيل الدخول'
                    if data[0][0] == 2:
                        action = 'تسجيل الخروج'
                    if data[row][0] == 3:
                        action = 'اضافة'
                    if data[row][0] == 4:
                        action = 'تعديل'
                    if data[row][0] == 5:
                        action = 'حذف'
                    if data[row][0] == 6:
                        action = 'بحث'
                    if data[row][0] == 7:
                        action = 'طباعة'
                    self.tableWidget_8.setItem(row, col, QTableWidgetItem(str(action)))
                if col == 2:
                    tables = ''
                    if data[row][1] == 1:
                        tables = 'مبيع يومي'
                    if data[row][1] == 2:
                        tables = 'تحليل'
                    if data[row][1] == 3:
                        tables = 'مشتريات'
                    if data[row][1] == 4:
                        tables = 'مراجعين'
                    if data[row][1] == 5:
                        tables = 'مستخدم'
                    self.tableWidget_8.setItem(row, col, QTableWidgetItem(str(tables)))
                if col == 3:
                    self.tableWidget_8.setItem(row, col, QTableWidgetItem(str(data[row][3])))

                col += 1
            row_pos = self.tableWidget_8.rowCount()
            self.tableWidget_8.insertRow(row_pos)

    def Search_In_All_Sales(self):
        client_name = self.lineEdit_25.text()
        if client_name != '0':
            self.cur.execute(''' SELECT * FROM addclient WHERE client_name=%s AND DATE(date)=%s ORDER BY id''',
                             (client_name, datetime.date.today(),))
        else:
            self.cur.execute(''' SELECT * FROM addclient WHERE DATE(date)=%s ORDER BY -date''',
                             (datetime.date.today(),))
        analyst_data = self.cur.fetchall()
        self.tableWidget_6.setRowCount(0)
        self.tableWidget_6.insertRow(0)
        try:
            if client_name != '0' and client_name != '':
                for i, k in enumerate(analyst_data[0]):
                    if i == 3:
                        self.tableWidget_6.setItem(0, i, QTableWidgetItem(str(analyst_data[0][4])))
                    else:
                        self.tableWidget_6.setItem(0, i, QTableWidgetItem(str(k)))
            else:
                self.Show_All_The_Sales()
        except:
            self.tableWidget_6.setRowCount(0)
            self.tableWidget_6.insertRow(0)
        # self.Add_Data_To_history(6, 1)

    def Print_Sale_Data(self, prev):
        genuses = self.comboBox_14.currentIndex()
        all_analyst = []
        all_result = []
        date = datetime.datetime.now()
        day = date.year
        month = date.month
        year = date.day
        word_type = []
        for row in range(0, self.tableWidget_5.rowCount() - 1):
            try:
                real_name = self.tableWidget_5.item(row, 0).text()
                analyst = self.tableWidget_5.item(row, 1).text()
            except:
                pass
            self.cur.execute(''' SELECT sub_category FROM addanalyst WHERE name=%s ''', (analyst,))
            if analyst not in all_analyst:
                all_analyst.append(str(analyst))
            data = self.cur.fetchone()
            if data != None:
                for item in data:
                    word_type.append(data[0])
            try:
                result = self.tableWidget_5.cellWidget(row, 2).currentText()
            except:
                try:
                    result = self.tableWidget_5.item(row, 2).text()
                except:
                    pass
            all_result.append(str(result))
            try:
                real_doctor = self.tableWidget_5.item(row, 3).text()
            except:
                pass
        # try:
        self.Bio_Word(real_name, real_doctor, all_analyst, all_result, year, month, day, word_type, prev, genuses)

        # except Exception as e:
        #     print(e)
        #     QMessageBox.information(self, 'خطأ', 'هنالك خطأ يرجى مراجعة العملية')

        if prev != 'T':
            self.Delete_Files()
        for rowj in range(0, self.tableWidget_5.rowCount() - 1):
            try:
                r2_doctor = self.tableWidget_5.item(rowj, 3).text()
                r2_analyst_name = self.tableWidget_5.item(rowj, 1).text()
                r2_client_name = self.lineEdit_20.text()
            except:
                pass
            try:
                r2_result = self.tableWidget_5.cellWidget(rowj, 2).currentText()
            except:
                try:
                    r2_result = self.tableWidget_5.item(rowj, 2).text()
                except:
                    pass
            try:
                self.cur.execute(
                    ''' UPDATE addnewitem SET doctor_name=%s ,analyst_name=%s ,analyst_result=%s WHERE client_name=%s AND analyst_name=%s''',
                    (r2_doctor, r2_analyst_name, r2_result, r2_client_name, r2_analyst_name))
                self.db.commit()
            except:
                print('exept')

    def get_total_price(self):
        total_price = 0
        GSE = ['Color:GSE', 'Consistency', 'Pus cells:GSE', 'R.B.Cs:GSE', 'E. Histolytica', 'Ova', 'Other:GSE',
               'Bacteria:GSE', 'Monillia:GSE', 'Fatty drop:GSE']
        GUE = ['Appearance', 'Reaction:GUE', 'Albumin', 'Sugar', 'RBCs:GUE', 'Pus cells:GUE', 'Epith .cells',
               'Crystals', 'Casts', 'Other:GUE', 'Bacteria:GUE', 'Monillia:GUE', 'Mucuse:GUE']
        SFA = ['Volume', 'Reaction:SFA', 'Colour:SFA', 'Liquefaction', 'Motility:Active', 'Motility:Sluggish',
               'Motility:Dead', 'Morphology:Normal', 'Morphology:Abnormal', 'Morphology:Pus cells', 'Other:SFA',
               'Count']
        My_num = ['Appearance', 'Reaction:GUE', 'Albumin', 'Sugar', 'RBCs:GUE', 'Pus cells:GUE', 'Epith .cells',
                  'Crystals',
                  'Casts', 'Other:GUE', 'Volume', 'Reaction:SFA', 'Colour:SFA', 'Liquefaction', 'Motility:Active',
                  'Motility:Sluggish', 'Motility:Dead', 'Morphology:Normal', 'Morphology:Abnormal',
                  'Morphology:Pus cells', 'Other:SFA', 'Color:GSE', 'Consistency', 'E. Histolytica', 'G. Lembilia',
                  'Ova', 'Pus cells:GSE', 'R.B.Cs:GSE', 'Bacteria:GSE', 'Bacteria:GUE', 'Monillia:GSE', 'Monillia:GUE',
                  'Fatty drop:GSE', 'Fatty drop:GUE', 'Mucuse:GUE']
        c_GSE = False
        c_GUE = False
        c_SFA = False
        rMy_num = 0
        for row in range(0, self.tableWidget_5.rowCount() - 1):
            try:
                analyst_name = self.tableWidget_5.item(row, 1).text()
            except:
                analyst_name = ''
            for kg in GSE:
                if kg == analyst_name:
                    c_GSE = True
            for kg1 in GUE:
                if kg1 == analyst_name:
                    c_GUE = True
            for kg2 in SFA:
                if kg2 == analyst_name:
                    c_SFA = True
            try:
                a = self.tableWidget_5.item(row, 4).text()
            except:
                a = 0
            try:
                total_price += int(a)
            except ValueError:
                a = 0
        if c_SFA == True:
            total_price += 3
        if c_GSE == True:
            total_price += 3
        if c_GUE == True:
            total_price += 3
        self.lineEdit_24.setText(str(total_price))

    def Sales_Page(self):
        self.lineEdit_20.setEnabled(False)
        global client_id_glob
        global chick_if_add_new
        global clients_name_glo
        My_num = ['Appearance', 'Reaction:GUE', 'Albumin', 'Sugar', 'RBCs:GUE', 'Pus cells:GUE', 'Epith .cells',
                  'Crystals',
                  'Casts', 'Other:GUE', 'Volume', 'Reaction:SFA', 'Colour:SFA', 'Liquefaction', 'Motility:Active',
                  'Motility:Sluggish', 'Motility:Dead', 'Morphology:Normal', 'Morphology:Abnormal',
                  'Morphology:Pus cells', 'Other:SFA', 'Color:GSE', 'Consistency', 'E. Histolytica', 'G. Lembilia',
                  'Ova', 'Pus cells:GSE', 'R.B.Cs:GSE', 'Bacteria:GSE', 'Bacteria:GUE', 'Monillia:GSE', 'Monillia:GUE',
                  'Fatty drop:GSE', 'Fatty drop:GUE', 'Mucuse:GUE']
        analyst_name = self.comboBox_16.currentText()
        client_name = self.lineEdit_20.text()
        client_age = self.spinBox_7.value()
        client_genus = self.comboBox_14.currentText()
        client_doctor = self.comboBox_15.currentText()
        client_genus = self.comboBox_14.currentText()
        analyst_or_clients_notes = self.textEdit.toPlainText()
        analyst_lineEdit_result = self.lineEdit_21.text()
        analyst_combo_result = self.comboBox_17.currentText()
        analyst_number_result = self.doubleSpinBox_7.value()
        client_id = self.spinBox.value()

        self.cur.execute('''SELECT price,category,sub_category FROM addanalyst WHERE name = %s''', (analyst_name,))
        analyst_price = self.cur.fetchone()

        latest_result = 1
        total_price = 0
        if analyst_price != None:
            if analyst_price[1] == 'عدد':
                if analyst_price[2] == 'SFA':
                    latest_result = ''
                else:
                    latest_result = analyst_number_result
            if analyst_price[1] == 'خيارات':
                latest_result = analyst_combo_result
            if analyst_price[1] == 'حقل كتابة':
                latest_result = analyst_lineEdit_result
            if analyst_price[1] == 'خيارات مع تعديل':
                latest_result = analyst_combo_result
            total_price = int(analyst_price[0])
        self.cur.execute('''
                   INSERT INTO addclient (client_name,client_age,client_genus,client_doctor,date)
                   VALUES (%s,%s,%s,%s,%s)
                ''', (
            str(client_name), str(client_age), str(client_genus), str(client_doctor), str(datetime.datetime.now())))
        self.db.commit()
        self.cur.execute('''SELECT id FROM addclient WHERE client_name = %s''', (client_name,))
        real_client_id = self.cur.fetchone()
        self.cur.execute('''
           INSERT INTO addnewitem (client_name,client_id,client_age,genus,doctor_name,notes,analyst_name,analyst_result,price,total_price,date)
           VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
        ''', (
            client_name, real_client_id, client_age, client_genus, client_doctor, analyst_or_clients_notes,
            analyst_name,
            latest_result,
            total_price, total_price, datetime.datetime.now()))
        self.db.commit()
        clients_name_glo.append(str(client_name))
        for rowj in range(0, self.tableWidget_5.rowCount() - 1):
            try:
                r2_doctor = self.tableWidget_5.item(rowj, 3).text()
                r2_analyst_name = self.tableWidget_5.item(rowj, 1).text()
                r2_client_name = self.lineEdit_20.text()
            except:
                r2_client_name = ''
                r2_analyst_name = ''
                r2_doctor = ''
            try:
                r2_result = self.tableWidget_5.cellWidget(rowj, 2).currentText()
            except:
                try:
                    r2_result = self.tableWidget_5.item(rowj, 2).text()
                except:
                    r2_result = ''
            try:
                self.cur.execute(
                    ''' UPDATE addnewitem SET doctor_name=%s ,analyst_name=%s ,analyst_result=%s WHERE client_name=%s AND analyst_name=%s''',
                    (r2_doctor, r2_analyst_name, r2_result, r2_client_name, r2_analyst_name))
                self.db.commit()
            except:
                print('except')
        self.cur.execute(
            ''' SELECT client_name,analyst_name,analyst_result,doctor_name,total_price FROM addnewitem WHERE client_name = %s AND DATE(date)=%s''',
            (self.lineEdit_20.text(), datetime.date.today(),))
        analyst_data = self.cur.fetchall()

        self.tableWidget_5.setRowCount(0)
        self.tableWidget_5.insertRow(0)
        for row, form in enumerate(analyst_data):
            for col, item in enumerate(form):
                # if col==4:
                #     print(total_price,'here')
                # self.tableWidget_5.setItem(row, col, QTableWidgetItem(str(total_price)))
                # else:
                self.tableWidget_5.setItem(row, col, QTableWidgetItem(str(item)))
                col += 1
            row_pos = self.tableWidget_5.rowCount()
            self.tableWidget_5.insertRow(row_pos)
        chick_if_add_new = True
        try:
            axd1 = self.tableWidget_5.findItems('Fatty drop:GSE', Qt.MatchContains)
            axd2 = self.tableWidget_5.findItems('Mucuse:GUE', Qt.MatchContains)
            axd3 = self.tableWidget_5.findItems('Morphology:Pus cells', Qt.MatchContains)
            self.tableWidget_5.insertRow(axd1[0].row() + 1)
            self.tableWidget_5.insertRow(axd1[0].row() + 2)
            self.tableWidget_5.insertRow(axd2[0].row() + 1)
            self.tableWidget_5.insertRow(axd2[0].row() + 2)
            self.tableWidget_5.insertRow(axd3[0].row() + 1)
            self.tableWidget_5.insertRow(axd3[0].row() + 2)
        except:
            pass
        for rowd in range(0, self.tableWidget_5.rowCount() - 1):
            all_name_items = []
            the_name = ''
            try:
                name = self.tableWidget_5.item(rowd, 0).text()
            except:
                pass
            for i in name:
                all_name_items.append(i)
                try:
                    if int(i):
                        all_name_items.remove(i)
                except:
                    pass
            for jikl in all_name_items:
                the_name += jikl
            self.tableWidget_5.setItem(rowd, 0, QTableWidgetItem(the_name))
            try:
                rs_name = self.tableWidget_5.item(rowd, 1).text()
            except:
                rs_name = ''
            mycobmbo = QComboBox(self)

            if rs_name == 'Color:GSE' or rs_name == 'Colour:SFA':
                mycobmbo = QComboBox(self)
                mycobmbo.addItems(['yallow', 'brown', 'green', 'milk'])
            if rs_name == 'Consistency':
                mycobmbo = QComboBox(self)
                mycobmbo.addItems(['Solid', 'Liquid', 'Semi solid', 'Semi liquid', 'Mucoid'])
            if rs_name == 'R.B.Cs:GSE' or rs_name == 'Pus cells:GSE' or rs_name == 'RBCs:GUE' or rs_name == 'Pus cells:GUE':
                mycobmbo = QComboBox(self)
                mycobmbo.addItems(['1 - 2', '1 - 3', '2 - 3', '2 - 4', '0 - 1', '0 - 2', '3 - 5', '4 - 6', '5 - 6',
                                   '5 - 7', '6 - 8', '6 - 7', '+', '++', '+++', '++++', 'Full Field'])
            if rs_name == 'E. Histolytica' or rs_name == 'G. Lembilia':
                mycobmbo = QComboBox(self)
                mycobmbo.addItems(['Cyst', 'Trophozoite'])
            if rs_name == 'Ova':
                mycobmbo = QComboBox(self)
                mycobmbo.addItems(['Nill'])
            if rs_name == 'Appearance':
                mycobmbo = QComboBox(self)
                mycobmbo.addItems(['Turbid', 'Clear'])
            if rs_name == 'Reaction:GUE' or rs_name == 'Reaction:SFA':
                mycobmbo = QComboBox(self)
                mycobmbo.addItems(['Acidic', 'Alkaline'])
            if rs_name == 'Albumin' or rs_name == 'Sugar':
                mycobmbo = QComboBox(self)
                mycobmbo.addItems(['Nil', '+', '++', '+++', 'Trace'])
            if rs_name == 'Epith .cells':
                mycobmbo = QComboBox(self)
                mycobmbo.addItems(['Few', '+', '++', '+++', '++++'])
            if rs_name == 'Crystals':
                mycobmbo = QComboBox(self)
                mycobmbo.addItems(['Am.Urate', 'Am.Phosphatase', 'Uric Acid', 'Ca.Oxalate'])
            if rs_name == 'Casts':
                mycobmbo = QComboBox(self)
                mycobmbo.addItems(['Granular cast +', 'Granular cast ++', 'Granular cast +++'])

            if rs_name == 'Bacteria:GSE' or rs_name == 'Monillia:GSE' or rs_name == 'Fatty drop:GSE' or rs_name == 'Monillia:GUE' or rs_name == 'Fatty drop:GUE' or rs_name == 'Bacteria:GUE' or rs_name == 'Mucuse:GUE':
                mycobmbo = QComboBox(self)
                mycobmbo.addItems(['Few', '+', '++', '+++', '++++'])
            if rs_name == 'Motility:Active' or rs_name == 'Motility:Sluggish' or rs_name == 'Motility:Dead' or rs_name == 'Morphology:Normal' or rs_name == 'Morphology:Abnormal' or rs_name == 'Morphology:Pus cells':
                mycobmbo = QComboBox(self)
                mycobmbo.addItems(
                    ['5', '10', '15', '20', '25', '30', '35', '40', '45', '50', '55', '60', '65', '70', '75', '80',
                     '85'])
            if rs_name == 'Volume':
                mycobmbo = QComboBox(self)
                mycobmbo.addItems(['0.5', '0.6', '0.7'])
            if rs_name == 'Liquefaction':
                mycobmbo = QComboBox(self)
                mycobmbo.addItems(['5', '10', '15', '20', '25', '30', '35', '40', '45'])
            # if rs_name=='':
            #     mycobmbo = QComboBox(self)
            #     mycobmbo.addItems([])
            # if rs_name=='':
            #     mycobmbo = QComboBox(self)
            #     mycobmbo.addItems([])
            # if rs_name=='':
            #     mycobmbo = QComboBox(self)
            #     mycobmbo.addItems([])
            # if rs_name=='':
            #     mycobmbo = QComboBox(self)
            #     mycobmbo.addItems([])
            # if rs_name=='':
            #     mycobmbo = QComboBox(self)
            #     mycobmbo.addItems([])
            # if rs_name=='':
            #     mycobmbo = QComboBox(self)
            #     mycobmbo.addItems([])
            # if rs_name=='':
            #     mycobmbo = QComboBox(self)
            #     mycobmbo.addItems([])
            # if rs_name=='':
            #     mycobmbo = QComboBox(self)
            #     mycobmbo.addItems([])
            # if rs_name=='':
            #     mycobmbo = QComboBox(self)
            #     mycobmbo.addItems([])
            # if rs_name=='':
            #     mycobmbo = QComboBox(self)
            #     mycobmbo.addItems([])
            # if rs_name=='':
            #     mycobmbo = QComboBox(self)
            #     mycobmbo.addItems([])
            # if rs_name=='':
            #     mycobmbo = QComboBox(self)
            #     mycobmbo.addItems([])
            for jjk in My_num:
                if rs_name == jjk:
                    # print('mystatrt')
                    r2_analyst_name = self.tableWidget_5.item(rowd, 1).text()
                    r2_client_name = self.lineEdit_20.text()
                    self.cur.execute(
                        ''' SELECT  analyst_result  FROM addnewitem WHERE client_name=%s AND analyst_name=%s ''',
                        (r2_client_name, r2_analyst_name))
                    myrs = self.cur.fetchall()

                    if myrs != '' or myrs != None:
                        index = mycobmbo.findText(myrs[0][0], Qt.MatchFixedString)
                        mycobmbo.setCurrentIndex(index)
                    self.tableWidget_5.setItem(rowd, 2, QTableWidgetItem(str('')))
                    self.tableWidget_5.setCellWidget(rowd, 2, mycobmbo)
                    if mycobmbo.currentText() == '' or mycobmbo.currentText() == ' ':
                        mycobmbo.setCurrentIndex(0)
        self.tableWidget_5.scrollToBottom()
        self.get_total_price()
        self.Show_All_The_Sales()
        self.Show_all_clients_without_search()
        # self.Show_All_Clients()
        # mycobmbo = QComboBox(self)

        # self.Show_All_one_client_analyst()
        # self.Add_Data_To_history(3, 1)
        # self.History()
        # client_name = self.lineEdit_20.setText('')
        # client_age = self.spinBox_7.setValue(0)
        # client_genus = self.comboBox_14.setCurrentIndex(0)
        # client_doctor = self.comboBox_15.setCurrentIndex(0)
        # analyst_name = self.comboBox_16.setCurrentIndex(0)
        # analyst_or_clients_notes = self.textEdit.setPlainText('')
        # analyst_lineEdit_result = self.lineEdit_21.setText('')
        # analyst_combo_result = self.comboBox_17.setCurrentIndex(0)
        # analyst_number_result = self.doubleSpinBox_7.setValue(0)
        # client_id = self.spinBox.setValue(0)
        # self.Show_All_one_client_analyst()
        for myitem in range(0, self.tableWidget_5.rowCount() - 1):
            self.tableWidget_5.item(myitem, 2).setBackground(QtGui.QColor(255, 156, 153))

    def Show_All_one_client_analyst(self):
        global client_id_glob
        global chick_if_add_new
        My_num = ['Appearance', 'Reaction:GUE', 'Albumin', 'Sugar', 'RBCs:GUE', 'Pus cells:GUE', 'Epith .cells',
                  'Crystals',
                  'Casts', 'Other:GUE', 'Volume', 'Reaction:SFA', 'Colour:SFA', 'Liquefaction', 'Motility:Active',
                  'Motility:Sluggish', 'Motility:Dead', 'Morphology:Normal', 'Morphology:Abnormal',
                  'Morphology:Pus cells', 'Other:SFA', 'Color:GSE', 'Consistency', 'E. Histolytica', 'G. Lembilia',
                  'Ova', 'Pus cells:GSE', 'R.B.Cs:GSE', 'Bacteria:GSE', 'Bacteria:GUE', 'Monillia:GSE', 'Monillia:GUE',
                  'Fatty drop:GSE', 'Fatty drop:GUE', 'Mucuse:GUE']
        client_name = self.lineEdit_20.text()
        # self.cur.execute('''
        #      SELECT client_name,analyst_name,analyst_result,doctor_name FROM addnewitem WHERE client_name = %s AND DATE(date)=%s
        # ''', (client_name,))
        # analyst_data = self.cur.fetchall()
        self.cur.execute('''
                     SELECT client_name,analyst_name,analyst_result,doctor_name,total_price,client_id,client_age,genus,notes FROM addnewitem WHERE client_id = %s
                ''', (self.spinBox.value(),))
        analyst_data = self.cur.fetchall()
        # need doctor name enabled i will make it with current text
        if self.spinBox.value() == 0 and chick_if_add_new == False:
            self.tableWidget_5.setRowCount(0)
            self.tableWidget_5.insertRow(0)
            self.lineEdit_20.setText('')
            self.lineEdit_21.setText('')
            self.spinBox_7.setValue(0)
            self.doubleSpinBox_7.setValue(0)
            self.comboBox_14.setCurrentIndex(0)
            self.comboBox_17.setCurrentIndex(0)
            self.textEdit.setPlainText('')
            self.lineEdit_20.setEnabled(True)
            self.spinBox_7.setEnabled(True)
            self.comboBox_14.setEnabled(True)
            self.textEdit.setEnabled(True)
        if self.spinBox.value() != 0:
            try:
                self.comboBox_15.setCurrentText(str(analyst_data[0][3]))
                self.comboBox_15.setEnabled(False)
                self.comboBox_14.setEnabled(False)
                self.lineEdit_20.setText(analyst_data[0][0])
                self.lineEdit_20.setEnabled(False)
                self.spinBox_7.setValue(int(analyst_data[0][6]))
                self.spinBox_7.setEnabled(False)
                if analyst_data[0][6] == 'ذكر':
                    self.comboBox_14.setCurrentIndex(1)
                    self.comboBox_14.setEnabled(False)
                if analyst_data[0][6] == 'انثى':
                    self.comboBox_14.setCurrentIndex(0)
                    self.comboBox_14.setEnabled(False)
                self.textEdit.setPlainText(str(analyst_data[0][8]))
                self.textEdit.setEnabled(False)
                self.cur.execute(
                    '''SELECT client_name,analyst_name,analyst_result,doctor_name,total_price FROM addnewitem WHERE client_name = %s AND DATE(date)=%s''',
                    (self.lineEdit_20.text(), datetime.date.today(),))
                analyst_data = self.cur.fetchall()

                self.tableWidget_5.setRowCount(0)
                self.tableWidget_5.insertRow(0)

                for row, form in enumerate(analyst_data):
                    for col, item in enumerate(form):
                        self.tableWidget_5.setItem(row, col, QTableWidgetItem(str(item)))
                        col += 1
                    row_pos = self.tableWidget_5.rowCount()
                    self.tableWidget_5.insertRow(row_pos)
                chick_if_add_new = False
                self.Show_All_The_Sales()

                try:
                    axd1 = self.tableWidget_5.findItems('Fatty drop:GSE', Qt.MatchContains)
                    axd2 = self.tableWidget_5.findItems('Mucuse:GUE', Qt.MatchContains)
                    axd3 = self.tableWidget_5.findItems('Morphology:Pus cells', Qt.MatchContains)
                    self.tableWidget_5.insertRow(axd1[0].row() + 1)
                    self.tableWidget_5.insertRow(axd1[0].row() + 2)
                    self.tableWidget_5.insertRow(axd2[0].row() + 1)
                    self.tableWidget_5.insertRow(axd2[0].row() + 2)
                    self.tableWidget_5.insertRow(axd3[0].row() + 1)
                    self.tableWidget_5.insertRow(axd3[0].row() + 2)
                except:
                    print('ecxehue')
                    pass
                for rowd in range(0, self.tableWidget_5.rowCount() - 1):
                    all_name_items = []
                    the_name = ''
                    try:
                        name = self.tableWidget_5.item(rowd, 0).text()
                    except:
                        name = ''
                    for i in name:
                        all_name_items.append(i)
                        try:
                            if int(i):
                                all_name_items.remove(i)
                        except:
                            pass
                    for jikl in all_name_items:
                        the_name += jikl
                    self.tableWidget_5.setItem(rowd, 0, QTableWidgetItem(the_name))
                    try:
                        rs_name = self.tableWidget_5.item(rowd, 1).text()
                    except:
                        rs_name = ''
                    mycobmbo = QComboBox(self)

                    if rs_name == 'Color:GSE' or rs_name == 'Colour:SFA':
                        mycobmbo = QComboBox(self)
                        mycobmbo.addItems(['yallow', 'brown', 'green', 'milk'])
                    if rs_name == 'Consistency':
                        mycobmbo = QComboBox(self)
                        mycobmbo.addItems(['Solid', 'Liquid', 'Semi solid', 'Semi liquid', 'Mucoid'])
                    if rs_name == 'R.B.Cs:GSE' or rs_name == 'Pus cells:GSE' or rs_name == 'RBCs:GUE' or rs_name == 'Pus cells:GUE':
                        mycobmbo = QComboBox(self)
                        mycobmbo.addItems(
                            ['1 - 2', '1 - 3', '2 - 3', '2 - 4', '0 - 1', '0 - 2', '3 - 5', '4 - 6', '5 - 6',
                             '5 - 7', '6 - 8', '6 - 7', '+', '++', '+++', '++++', 'Full Field'])
                    if rs_name == 'E. Histolytica' or rs_name == 'G. Lembilia':
                        mycobmbo = QComboBox(self)
                        mycobmbo.addItems(['Cyst', 'Trophozoite'])
                    if rs_name == 'Ova':
                        mycobmbo = QComboBox(self)
                        mycobmbo.addItems(['Nill'])
                    if rs_name == 'Appearance':
                        mycobmbo = QComboBox(self)
                        mycobmbo.addItems(['Turbid', 'Clear'])
                    if rs_name == 'Reaction:GUE' or rs_name == 'Reaction:SFA':
                        mycobmbo = QComboBox(self)
                        mycobmbo.addItems(['Acidic', 'Alkaline'])
                    if rs_name == 'Albumin' or rs_name == 'Sugar':
                        mycobmbo = QComboBox(self)
                        mycobmbo.addItems(['Nil', '+', '++', '+++', 'Trace'])
                    if rs_name == 'Epith .cells':
                        mycobmbo = QComboBox(self)
                        mycobmbo.addItems(['Few', '+', '++', '+++', '++++'])
                    if rs_name == 'Crystals':
                        mycobmbo = QComboBox(self)
                        mycobmbo.addItems(['Am.Urate', 'Am.Phosphatase', 'Uric Acid', 'Ca.Oxalate'])
                    if rs_name == 'Casts':
                        mycobmbo = QComboBox(self)
                        mycobmbo.addItems(['Granular cast +', 'Granular cast ++', 'Granular cast +++'])

                    if rs_name == 'Bacteria:GSE' or rs_name == 'Monillia:GSE' or rs_name == 'Fatty drop:GSE' or rs_name == 'Monillia:GUE' or rs_name == 'Fatty drop:GUE' or rs_name == 'Bacteria:GUE' or rs_name == 'Mucuse:GUE':
                        mycobmbo = QComboBox(self)
                        mycobmbo.addItems(['Few', '+', '++', '+++', '++++'])
                    if rs_name == 'Motility:Active' or rs_name == 'Motility:Sluggish' or rs_name == 'Motility:Dead' or rs_name == 'Morphology:Normal' or rs_name == 'Morphology:Abnormal' or rs_name == 'Morphology:Pus cells':
                        mycobmbo = QComboBox(self)
                        mycobmbo.addItems(
                            ['5', '10', '15', '20', '25', '30', '35', '40', '45', '50', '55', '60', '65', '70', '75',
                             '80',
                             '85'])
                    if rs_name == 'Volume':
                        mycobmbo = QComboBox(self)
                        mycobmbo.addItems(['0.5', '0.6', '0.7'])
                    if rs_name == 'Liquefaction':
                        mycobmbo = QComboBox(self)
                        mycobmbo.addItems(['5', '10', '15', '20', '25', '30', '35', '40', '45'])
                    # if rs_name=='':
                    #     mycobmbo = QComboBox(self)
                    #     mycobmbo.addItems([])
                    # if rs_name=='':
                    #     mycobmbo = QComboBox(self)
                    #     mycobmbo.addItems([])
                    # if rs_name=='':
                    #     mycobmbo = QComboBox(self)
                    #     mycobmbo.addItems([])
                    # if rs_name=='':
                    #     mycobmbo = QComboBox(self)
                    #     mycobmbo.addItems([])
                    # if rs_name=='':
                    #     mycobmbo = QComboBox(self)
                    #     mycobmbo.addItems([])
                    # if rs_name=='':
                    #     mycobmbo = QComboBox(self)
                    #     mycobmbo.addItems([])
                    # if rs_name=='':
                    #     mycobmbo = QComboBox(self)
                    #     mycobmbo.addItems([])
                    # if rs_name=='':
                    #     mycobmbo = QComboBox(self)
                    #     mycobmbo.addItems([])
                    # if rs_name=='':
                    #     mycobmbo = QComboBox(self)
                    #     mycobmbo.addItems([])
                    # if rs_name=='':
                    #     mycobmbo = QComboBox(self)
                    #     mycobmbo.addItems([])
                    # if rs_name=='':
                    #     mycobmbo = QComboBox(self)
                    #     mycobmbo.addItems([])
                    # if rs_name=='':
                    #     mycobmbo = QComboBox(self)
                    #     mycobmbo.addItems([])
                    for jjk in My_num:
                        try:
                            if rs_name == jjk:
                                # print('mystatrt')
                                r2_analyst_name = self.tableWidget_5.item(rowd, 1).text()
                                r2_client_name = self.lineEdit_20.text()
                                self.cur.execute(
                                    ''' SELECT  analyst_result  FROM addnewitem WHERE client_name=%s AND analyst_name=%s ''',
                                    (r2_client_name, r2_analyst_name))
                                myrs = self.cur.fetchall()

                                if myrs != '' or myrs != None:
                                    index = mycobmbo.findText(myrs[0][0], Qt.MatchFixedString)
                                    mycobmbo.setCurrentIndex(index)
                                self.tableWidget_5.setItem(rowd, 2, QTableWidgetItem(str('')))
                                self.tableWidget_5.setCellWidget(rowd, 2, mycobmbo)
                                if mycobmbo.currentText() == '' or mycobmbo.currentText() == ' ':
                                    mycobmbo.setCurrentIndex(0)
                        except:
                            pass

            except IndexError:
                QMessageBox.information(self, 'Error',
                                        'الرقم الذي ادخلته غير صحيح يرجى ادخال رقم صحيح او مراجعة صفحة "كل المبيعات" للتأكد من الرقم')
            self.get_total_price()
            for myitem in range(0, self.tableWidget_5.rowCount() - 1):
                self.tableWidget_5.item(myitem, 2).setBackground(QtGui.QColor(255, 156, 153))

    def Chick_analyst_category(self):
        pationt_name = self.lineEdit_20.text()
        doctor_name = self.comboBox_15.currentText()

        analyst_name = self.comboBox_16.currentText()
        self.cur.execute('''SELECT category FROM addanalyst WHERE name = %s''', (analyst_name,))
        analyst_category = self.cur.fetchone()
        try:
            if analyst_category[0] == 'عدد':
                self.comboBox_17.hide()
                self.lineEdit_21.hide()
                self.doubleSpinBox_7.show()
            if analyst_category[0] == 'خيارات':
                self.lineEdit_21.hide()
                self.doubleSpinBox_7.hide()
                self.comboBox_17.show()
                self.comboBox_17.setEditable(False)
            if analyst_category[0] == 'حقل كتابة':
                self.lineEdit_21.show()
                self.doubleSpinBox_7.hide()
                self.comboBox_17.hide()
            if analyst_category[0] == 'خيارات مع تعديل':
                self.lineEdit_21.hide()
                self.doubleSpinBox_7.hide()
                # self.comboBox_17.clear()
                self.comboBox_17.show()
                self.comboBox_17.setEditable(True)
        except:
            if analyst_name != 'Full GSE' and analyst_name != 'Full GUE' and analyst_name != 'Full SFA':
                QMessageBox.information(self, 'تحذير', "يرجى اختيار تحليل صحيح")
        row_count = self.tableWidget_5.rowCount()
        if analyst_name == 'Full GSE':
            # QStatusBar.showMessage(self,'يرجى الانتظار سيتم تنفيذ طلبك خلال ثواني')
            self.comboBox_16.setCurrentText('Color:GSE')
            self.Sales_Page()
            self.comboBox_16.setCurrentText('Consistency')
            self.Sales_Page()
            self.comboBox_16.setCurrentText('R.B.Cs:GSE')
            self.Sales_Page()
            self.comboBox_16.setCurrentText('Pus cells:GSE')
            self.Sales_Page()
            self.comboBox_16.setCurrentText('E. Histolytica')
            self.Sales_Page()
            self.comboBox_16.setCurrentText('G. Lembilia')
            self.Sales_Page()
            self.comboBox_16.setCurrentText('Ova')
            self.Sales_Page()
            self.comboBox_16.setCurrentText('Bacteria:GSE')
            self.Sales_Page()
            self.comboBox_16.setCurrentText('Monillia:GSE')
            self.Sales_Page()
            self.comboBox_16.setCurrentText('Fatty drop:GSE')
            self.Sales_Page()
            self.comboBox_16.setCurrentIndex(0)
        if analyst_name == 'Full GUE':
            # QStatusBar.showMessage(self,'يرجى الانتظار سيتم تنفيذ طلبك خلال ثواني')
            self.comboBox_16.setCurrentText('Appearance')
            self.Sales_Page()
            self.comboBox_16.setCurrentText('Reaction:GUE')
            self.Sales_Page()
            self.comboBox_16.setCurrentText('Albumin')
            self.Sales_Page()
            self.comboBox_16.setCurrentText('Sugar')
            self.Sales_Page()
            self.comboBox_16.setCurrentText('RBCs:GUE')
            self.Sales_Page()
            self.comboBox_16.setCurrentText('Pus cells:GUE')
            self.Sales_Page()
            self.comboBox_16.setCurrentText('Epith .cells')
            self.Sales_Page()
            self.comboBox_16.setCurrentText('Crystals')
            self.Sales_Page()
            self.comboBox_16.setCurrentText('Casts')
            self.Sales_Page()
            self.comboBox_16.setCurrentText('Bacteria:GUE')
            self.Sales_Page()
            self.comboBox_16.setCurrentText('Monillia:GUE')
            self.Sales_Page()
            self.comboBox_16.setCurrentText('Mucuse:GUE')
            self.Sales_Page()
        if analyst_name == 'Full SFA':
            # QStatusBar.showMessage(self,'يرجى الانتظار سيتم تنفيذ طلبك خلال ثواني')
            self.comboBox_16.setCurrentText('Volume')
            self.Sales_Page()
            self.comboBox_16.setCurrentText('Reaction:SFA')
            self.Sales_Page()
            self.comboBox_16.setCurrentText('Colour:SFA')
            self.Sales_Page()
            self.comboBox_16.setCurrentText('Liquefaction')
            self.Sales_Page()
            self.comboBox_16.setCurrentText('Count')
            self.Sales_Page()
            self.comboBox_16.setCurrentText('Motility:Active')
            self.Sales_Page()
            self.comboBox_16.setCurrentText('Motility:Sluggish')
            self.Sales_Page()
            self.comboBox_16.setCurrentText('Motility:Dead')
            self.Sales_Page()
            self.comboBox_16.setCurrentText('Morphology:Normal')
            self.Sales_Page()
            self.comboBox_16.setCurrentText('Morphology:Abnormal')
            self.Sales_Page()
            self.comboBox_16.setCurrentText('Morphology:Pus cells')
            self.Sales_Page()

        colors = ['yallow', 'brown', 'green', 'milk']
        RBCs_Pus_cells_GUE_GSE = ['1 - 2', '1 - 3', '2 - 3', '2 - 4', '0 - 1', '0 - 2', '3 - 5', '4 - 6', '5 - 6',
                                  '5 - 7', '6 - 8', '6 - 7', '+', '++', '+++', '++++', 'Full Field']
        Consistency = ['Solid', 'Liquid', 'Semi solid', 'Semi liquid', 'Mucoid']
        G_Lembilia_E_Histolytica = ['Cyst', 'Trophozoite']
        Epith_cells = ['Few', '+', '++', '+++', '++++']
        Ova = ['Nil']
        Appearance = ['Turbid', 'Clear']
        Reaction = ['Acidic', 'Alkaline']
        Sugar = ['Nil', '+', '++', '+++', 'Trace']
        Crystals = ['Am.Urate', 'Am.Phosphatase', 'Uric Acid', 'Ca.Oxalate']
        Casts = ['Granular cast +', 'Granular cast ++', 'Granular cast +++']
        Blood_Group = ['A (+ve)', 'B (+ve)', 'AB (+ve)', 'O (+ve)', 'O (-ve)', 'A (-ve)', 'B (-ve)', 'AB (-ve)']
        Pregnancy_test_in_urine_serum = ['Positive (+ve)', 'Negative (-ve)', 'Weak Positive']
        Pregnancy_test_in_urine_serum2 = ['Positive (+ve)', 'Negative (-ve)']
        Blood_Group = ['A(+ve)', 'B(+ve)', 'AB(+ve)', 'O(+ve)', 'O(-ve)', 'A(-ve)', 'B(-ve)', 'AB(-ve)']
        test = ['Toxoplasma IgG', 'Toxoplasma IgM', 'Cytomegalo Virus IgG', 'Cytomegalo Virus IgM', 'Rubella IgG',
                'Rubella IgM', 'Anti - Phspholipin IgG', 'Anti - Phspholipin  IgM', 'Anti - Cardiolipin  IgG',
                'Anti - Cardiolipin  IgM', 'Herps   IgG', 'Herpes  IgM']
        test_choices = ['0.5 Negative', '0.6 Negative', '0.7 Negative', '0.8 Negative', '1.1 Positive', '1.1 Positive',
                        '1.2 Positive', '1.3 Positive', '1.4 Positive', '1.5 Positive']
        Hb = [4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18]
        motility = [5, 10, 15, 20, 25, 30, 35, 40, 45, 50, 55, 60, 65, 70, 75, 80, 85]
        self.comboBox_17.clear()
        if analyst_name == 'Colour:SFA' or analyst_name == 'Color:GSE':
            for item in colors:
                self.comboBox_17.addItem(str(item))
        if analyst_name == 'Blood Group':
            for itemg in Blood_Group:
                self.comboBox_17.addItem(str(itemg))
        if analyst_name == 'Motility:Active' or analyst_name == 'Motility:Sluggish' or analyst_name == 'Motility:Dead' or analyst_name == 'Morphology:Normal' or analyst_name == 'Morphology:Abnormal' or analyst_name == 'Morphology:Pus cells':
            for item101 in motility:
                self.comboBox_17.addItem(str(item101))
        if analyst_name == 'R.B.Cs:GSE' or analyst_name == 'Pus cells:GSE' or analyst_name == 'RBCs:GUE' or analyst_name == 'Pus cells:GUE' or analyst_name == 'Morphology:Pus cells':
            for item0 in RBCs_Pus_cells_GUE_GSE:
                self.comboBox_17.addItem(str(item0))
        if analyst_name == 'Consistency':
            for item1 in Consistency:
                self.comboBox_17.addItem(str(item1))
        if analyst_name == 'E. Histolytica' or analyst_name == 'G. Lembilia':
            for item2 in G_Lembilia_E_Histolytica:
                self.comboBox_17.addItem(str(item2))
        if analyst_name == 'Ova':
            for item4 in Ova:
                self.comboBox_17.addItem(str(item4))
        if analyst_name == 'Appearance':
            for item5 in Appearance:
                self.comboBox_17.addItem(str(item5))
        if analyst_name == 'Reaction:SFA' or analyst_name == 'Reaction:GUE':
            for item6 in Reaction:
                self.comboBox_17.addItem(str(item6))
        if analyst_name == 'Sugar' or analyst_name == 'Albumin':
            for item7 in Sugar:
                self.comboBox_17.addItem(str(item7))
        if analyst_name == 'Epith .cells' or analyst_name == 'Bacteria:GSE' or analyst_name == 'Bacteria:GUE' or analyst_name == 'Monillia:GSE' or analyst_name == 'Monillia:GUE' or analyst_name == 'Fatty drop:GSE' or analyst_name == 'Fatty drop:GUE' or analyst_name == 'Mucuse:GUE':
            for item8 in Epith_cells:
                self.comboBox_17.addItem(str(item8))
        if analyst_name == 'Crystals':
            for item9 in Crystals:
                self.comboBox_17.addItem(str(item9))
        if analyst_name == 'Casts':
            for item10 in Casts:
                self.comboBox_17.addItem(str(item10))
        if analyst_name == 'Hb':
            for item100 in Hb:
                self.comboBox_17.setEditable(True)
                self.comboBox_17.addItem(str(item100))
        if analyst_name == 'Pregnancy test  in serum' or analyst_name == 'Pregnancy test  in urine' or analyst_name == 'Salmonella typhi  IgG' or analyst_name == 'Salmonella typhi  IgM' or analyst_name == 'Rose-Bengal test' or analyst_name == 'Rh' or analyst_name == 'HBS Ag' or analyst_name == 'HCV Ab' or analyst_name == 'HIV':
            for item11 in Pregnancy_test_in_urine_serum2:
                self.comboBox_17.addItem(str(item11))
        if analyst_name == 'HBS Ag' or analyst_name == 'HCV Ab' or analyst_name == 'HIV':
            for item11 in Pregnancy_test_in_urine_serum:
                self.comboBox_17.addItem(str(item11))
        for fitem in test:
            if analyst_name == fitem:
                for item12 in test_choices:
                    self.comboBox_17.addItem(str(item12))

    def get_client_id(self):
        global client_id_glob
        client_id_glob = self.spinBox.value()
        self.Show_All_one_client_analyst()
        # self.Add_Data_To_history(6, 1)
        # self.History()

    def Show_All_The_Sales(self):
        global show_all_sales_in_clients_page
        self.cur.execute(''' SELECT client_name FROM addclient WHERE DATE(date)=%s ORDER BY -date ''',
                         (datetime.date.today(),))
        analyst_data = self.cur.fetchall()
        all_names = {}
        all_data = []
        id = 0

        for item in analyst_data:
            if item[0] not in all_names:
                all_names[item[0]] = 0

        for s in all_names.keys():
            self.cur.execute(''' SELECT id FROM addclient WHERE client_name=%s ORDER BY id ''', (s,))
            analyst_data2 = self.cur.fetchall()
            all_names[s] = analyst_data2[0][0]
            id = analyst_data2[0][0]
            self.cur.execute(''' SELECT * FROM addclient WHERE id=%s ORDER BY -date ''', (id,))
            latest_data = self.cur.fetchall()
            data = [{'name': s, 'value': latest_data}]
            all_data.append(data)
        self.tableWidget_6.setRowCount(0)
        self.tableWidget_6.insertRow(0)
        # if show_all_sales_in_clients_page:
        #     self.tableWidget_2.setRowCount(0)
        #     self.tableWidget_2.insertRow(0)
        l_data = ()
        for ih in all_data:
            hs_data = ih[0]['value']
            l_data += tuple(hs_data)
        for row, form in enumerate(l_data):
            for col, item in enumerate(form):
                if col == 3:
                    self.tableWidget_6.setItem(row, col, QTableWidgetItem(str(l_data[row][4])))
                    # if show_all_sales_in_clients_page:
                    #
                    #     self.tableWidget_2.setItem(row, col, QTableWidgetItem(str(l_data[row][4])))
                else:
                    print('d')
                    self.tableWidget_6.setItem(row, col, QTableWidgetItem(str(item)))
                    # if show_all_sales_in_clients_page:
                    #     self.tableWidget_2.setItem(row, col, QTableWidgetItem(str(item)))
                col += 1
            row_pos = self.tableWidget_6.rowCount()
            self.tableWidget_6.insertRow(row_pos)
            # if show_all_sales_in_clients_page:
            #     row_pos2 = self.tableWidget_2.rowCount()
            #     self.tableWidget_2.insertRow(row_pos2)
        # print(all_data)

    # def Show_All_The_Analysts(self):
    #     search_type = self.comboBox_19.currentText()
    #     search_words = self.lineEdit_26.text()
    #
    #     if self.comboBox_19.currentIndex() == 1:
    #         self.cur.execute(
    #             ''' SELECT name,category,default_result1,default_result2,price,sub_category FROM addanalyst WHERE name=%s ''',
    #             (search_words,))
    #     if self.comboBox_19.currentIndex() == 2:
    #         self.cur.execute(
    #             ''' SELECT name,category,default_result1,default_result2,price,sub_category FROM addanalyst WHERE price=%s ''',
    #             (search_words,))
    #     if self.comboBox_19.currentIndex() == 0:
    #         self.cur.execute(
    #             ''' SELECT name,category,default_result1,default_result2,price,sub_category FROM addanalyst''')
    #
    #     # combobox_value = ''
    #     # sql = ''
    #     # if search_type == 'اسم التحليل المطابق':
    #     #     combobox_value = 'name'
    #     # if search_type == 'السعر المطابق':
    #     #     combobox_value = 'price'
    #     # if search_type != 'اسم التحليل المطابق' and search_type != 'السعر المطابق':
    #     #     sql = ''' SELECT name,category,default_result1,default_result2,price,sub_category FROM addanalyst'''
    #     # else:
    #     #     sql = f' SELECT name,category,default_result1,default_result2,price,sub_category FROM addanalyst WHERE {combobox_value}=%s'
    #     # if type(search_words) == int:
    #     #     self.cur.execute(sql, search_words)
    #     # if type(search_words) == str:
    #     #     self.cur.execute(sql, search_words)
    #     analyst_data = self.cur.fetchall()
    #     self.tableWidget_7.setRowCount(0)
    #     self.tableWidget_7.insertRow(0)
    #     for row, form in enumerate(analyst_data):
    #         for col, item in enumerate(form):
    #             if col == 1:
    #                 category = ''
    #                 if analyst_data[row][1] == 'خيارات':
    #                     category = 'خيارات'
    #                 if analyst_data[row][1] == 'عدد':
    #                     category = 'عدد'
    #                 if analyst_data[row][1] == 'خيارات مع تعديل':
    #                     category = 'خيارات مع تعديل'
    #                 if analyst_data[row][1] == 'حقل كتابة':
    #                     category = 'حقل كتابة'
    #
    #                 self.tableWidget_7.setItem(row, col, QTableWidgetItem(str(category)))
    #             if col == 5:
    #                 self.tableWidget_7.setItem(row, col, QTableWidgetItem(str(analyst_data[row][5])))
    #
    #             self.tableWidget_7.setItem(row, col, QTableWidgetItem(str(item)))
    #             col += 1
    #         row_pos = self.tableWidget_7.rowCount()
    #         self.tableWidget_7.insertRow(row_pos)
    #     self.Add_Data_To_history(6, 2)
    #     self.History()

    # def Add_Analyst(self):
    #     analyst_name = self.lineEdit_28.text()
    #     analyst_result_category = self.comboBox_22.currentText()
    #     analyst_result1 = self.doubleSpinBox_3.value()
    #     analyst_result2 = self.doubleSpinBox_2.value()
    #     analyst_price = self.doubleSpinBox.value()
    #     sub_category = self.comboBox_23.currentText()
    #     date = datetime.datetime.now()
    #     self.cur.execute(
    #         ''' INSERT INTO addanalyst (name,default_result1,default_result2,price,category,sub_category,date) VALUES(%s,%s,%s,%s,%s,%s,%s) ''',
    #         (
    #             analyst_name, analyst_result1, analyst_result2, analyst_price, analyst_result_category, sub_category,
    #             date))
    #     self.db.commit()
    #     QMessageBox.information(self, 'info', 'تم اضافة التحليل بنجاح')
    #     analyst_name = self.lineEdit_28.setText('')
    #     analyst_result_category = self.comboBox_22.setCurrentIndex(0)
    #     analyst_result1 = self.doubleSpinBox_3.setValue(0)
    #     analyst_result2 = self.doubleSpinBox_2.setValue(0)
    #     analyst_price = self.doubleSpinBox.setValue(0)
    #     sub_category = self.comboBox_23.setCurrentIndex(0)
    #     self.Show_all_analysts_in_combo()
    #     self.Add_Data_To_history(3, 2)
    #     self.History()
    #
    # def Show_analyst_in_Edit_Or_Delete(self):
    #     analyst_current_name = self.comboBox_21.currentText()
    #     self.cur.execute(
    #         ''' SELECT name,default_result1,default_result2,price,category,sub_category,date FROM addanalyst WHERE name=%s ''',
    #         (analyst_current_name,))
    #     data = self.cur.fetchall()
    #     num = 0
    #     if data[0][4] == 'عدد':
    #         num = 4
    #     if data[0][4] == 'خيارات':
    #         num = 2
    #     if data[0][4] == 'حقل كتابة':
    #         num = 1
    #     if data[0][4] == 'خيارات مع تعديل':
    #         num = 3
    #     tp = 0
    #     if data[0][5] == 'bio':
    #         tp = 1
    #     if data[0][5] == 'GSE':
    #         tp = 2
    #     if data[0][5] == 'GUE':
    #         tp = 3
    #     if data[0][5] == 'hematology':
    #         tp = 4
    #     if data[0][5] == 'هرمونات مشترك':
    #         tp = 5
    #     sub_category = self.comboBox_26.setCurrentIndex(tp)
    #     analyst_name = self.lineEdit_29.setText(str(data[0][0]))
    #     analyst_result_category = self.comboBox_25.setCurrentIndex(num)
    #     analyst_result1 = self.doubleSpinBox_4.setValue(data[0][1])
    #     analyst_result2 = self.doubleSpinBox_5.setValue(data[0][2])
    #     analyst_price = self.doubleSpinBox_6.setValue(data[0][3])
    #     self.Add_Data_To_history(6, 2)
    #     self.History()
    #
    # def Edit_Analyst(self):
    #     analyst_current_name = self.comboBox_21.currentText()
    #     analyst_name = self.lineEdit_29.text()
    #     analyst_result_category = self.comboBox_23.currentText()
    #     analyst_result1 = self.doubleSpinBox_4.value()
    #     analyst_result2 = self.doubleSpinBox_5.value()
    #     analyst_price = self.doubleSpinBox_6.value()
    #     sub_category = self.comboBox_26.currentText()
    #     date = datetime.datetime.now()
    #     self.cur.execute(
    #         ''' UPDATE  addanalyst SET name=%s,default_result1=%s,default_result2=%s,price=%s,category=%s,sub_category=%s,date=%s WHERE name=%s ''',
    #         (
    #             analyst_name, analyst_result1, analyst_result2, analyst_price, analyst_result_category, sub_category,
    #             date, analyst_name))
    #     self.db.commit()
    #     QMessageBox.information(self, 'info', 'تم تعديل التحليل بنجاح')
    #     analyst_current_name = self.comboBox_21.setCurrentIndex(0)
    #     analyst_name = self.lineEdit_29.setText('')
    #     analyst_result_category = self.comboBox_23.setCurrentIndex(0)
    #     analyst_result1 = self.doubleSpinBox_4.setValue(0)
    #     analyst_result2 = self.doubleSpinBox_5.setValue(0)
    #     analyst_price = self.doubleSpinBox_6.setValue(0)
    #     sub_category = self.comboBox_26.setCurrentIndex(0)
    #     self.Add_Data_To_history(4, 2)
    #     self.History()
    #     self.Show_all_analysts_in_combo()
    #
    # def Delete_Analyst(self):
    #
    #     warning = QMessageBox.warning(self, 'احذر', "هل انت متأكد من انك تريد مسح التحليل",
    #                                   QMessageBox.Yes | QMessageBox.No)
    #     if warning == QMessageBox.Yes:
    #         analyst_current_name = self.comboBox_21.currentText()
    #         sql = ''' DELETE FROM addanalyst WHERE name=%s '''
    #         self.cur.execute(sql, [(analyst_current_name)])
    #         self.db.commit()
    #         QMessageBox.information(self, 'info', 'تم حذف التحليل بنجاح')
    #
    #         self.Show_all_analysts_in_combo()
    #     self.Add_Data_To_history(5, 2)
    #     self.History()
    #     self.Show_all_analysts_in_combo()

    def Show_all_analysts_in_combo(self):
        self.cur.execute(
            ''' SELECT name FROM addanalyst WHERE sub_category=%s OR sub_category=%s OR sub_category=%s ''',
            ('bio', 'هرمونات مشترك', 'hematology'))
        data = self.cur.fetchall()
        self.comboBox_21.clear()
        self.comboBox_16.clear()
        self.comboBox_21.addItem('----------------')
        self.comboBox_16.addItem('----------------')
        self.comboBox_16.addItem('Full GSE')
        self.comboBox_16.addItem('Full GUE')
        self.comboBox_16.addItem('Full SFA')
        for item in data:
            self.comboBox_21.addItem(str(item[0]))
            self.comboBox_16.addItem(str(item[0]))

    def Clients_Page(self):
        self.tableWidget_4.setRowCount(0)
        self.tableWidget_4.insertRow(0)
        self.tableWidget_9.setRowCount(0)
        self.tableWidget_9.insertRow(0)
        id = self.spinBox_2.value()
        self.cur.execute(''' SELECT client_name,client_age,client_genus,client_doctor FROM addclient WHERE id=%s''',
                         (str(id),))
        client_data = self.cur.fetchall()
        self.cur.execute(
            ''' SELECT price,analyst_name,analyst_result,client_name FROM addnewitem WHERE client_id=%s''',
            (str(id),))
        client_analyst_data = self.cur.fetchall()
        num = 0
        all_client_analyst = []
        total = 0
        rMy_num = 0
        GSE = ['Color:GSE', 'Consistency', 'Pus cells:GSE', 'R.B.Cs:GSE', 'E. Histolytica', 'Ova', 'Other:GSE',
               'Bacteria:GSE', 'Monillia:GSE', 'Fatty drop:GSE']
        GUE = ['Appearance', 'Reaction:GUE', 'Albumin', 'Sugar', 'RBCs:GUE', 'Pus cells:GUE', 'Epith .cells',
               'Crystals', 'Casts', 'Other:GUE', 'Bacteria:GUE', 'Monillia:GUE', 'Mucuse:GUE']
        SFA = ['Volume', 'Reaction:SFA', 'Colour:SFA', 'Liquefaction', 'Motility:Active', 'Motility:Sluggish',
               'Motility:Dead', 'Morphology:Normal', 'Morphology:Abnormal', 'Morphology:Pus cells', 'Other:SFA',
               'Count']

        c_GSE = False
        c_GUE = False
        c_SFA = False
        My_num = ['Appearance', 'Reaction:GUE', 'Albumin', 'Sugar', 'RBCs:GUE', 'Pus cells:GUE', 'Epith .cells',
                  'Crystals',
                  'Casts', 'Other:GUE', 'Volume', 'Reaction:SFA', 'Colour:SFA', 'Liquefaction', 'Motility:Active',
                  'Motility:Sluggish', 'Motility:Dead', 'Morphology:Normal', 'Morphology:Abnormal',
                  'Morphology:Pus cells', 'Other:SFA', 'Color:GSE', 'Consistency', 'E. Histolytica', 'G. Lembilia',
                  'Ova', 'Pus cells:GSE', 'R.B.Cs:GSE', 'Bacteria:GSE', 'Bacteria:GUE', 'Monillia:GSE', 'Monillia:GUE',
                  'Fatty drop:GSE', 'Fatty drop:GUE', 'Mucuse:GUE']

        for i in client_analyst_data:
            num += 1
        for price in range(0, num):
            if client_analyst_data[price][1] not in GSE and client_analyst_data[price][1] not in GUE and \
                    client_analyst_data[price][1] not in SFA:
                total += client_analyst_data[price][0]
        for k in range(0, num):
            for kg in GSE:
                if kg == client_analyst_data[k][1]:
                    c_GSE = True
            for kg1 in GUE:
                if kg1 == client_analyst_data[k][1]:
                    c_GUE = True
            for kg2 in SFA:
                if kg2 == client_analyst_data[k][1]:
                    c_SFA = True
        for j in range(0, num):
            all_client_analyst.append(str(client_analyst_data[j][1]))
        if c_SFA == True:
            total += 3
        if c_GSE == True:
            total += 3
        if c_GUE == True:
            total += 3
        self.tableWidget_4.setRowCount(0)
        self.tableWidget_4.insertRow(0)
        for row, form in enumerate(client_data):

            for col, item in enumerate(form):
                self.tableWidget_4.setItem(row, 4, QTableWidgetItem(str(num)))
                self.tableWidget_4.setItem(row, 5, QTableWidgetItem(str(','.join(all_client_analyst))))
                self.tableWidget_4.setItem(row, 6, QTableWidgetItem(str(total)))
                self.tableWidget_4.setItem(row, col, QTableWidgetItem(str(item)))
                col += 1

            row_pos = self.tableWidget_4.rowCount()
            self.tableWidget_4.insertRow(row_pos)
        self.tableWidget_9.setRowCount(0)
        self.tableWidget_9.insertRow(0)
        for row, form in enumerate(client_analyst_data):

            for col, item in enumerate(form):
                if col == 0:
                    self.tableWidget_9.setItem(row, col, QTableWidgetItem(str(client_analyst_data[0][3])))
                if col == 1:
                    self.tableWidget_9.setItem(row, col, QTableWidgetItem(str(client_analyst_data[row][1])))
                if col == 2:
                    self.tableWidget_9.setItem(row, col, QTableWidgetItem(str(client_analyst_data[row][2])))
                col += 1

            row_pos = self.tableWidget_9.rowCount()
            self.tableWidget_9.insertRow(row_pos)
        # self.Add_Data_To_history(6, 5)
        # self.History()

    # def Add_Buys(self):
    #     item_name = self.lineEdit_13.text()
    #     item_type = self.comboBox_6.currentText()
    #     quantity = self.spinBox_3.value()
    #     quantity_avaliable = quantity
    #     signal_item_price = self.lineEdit_12.text()
    #     total_item_price = self.lineEdit_11.text()
    #     date_create_before = self.dateEdit.date()
    #     date_create = str(date_create_before.toPyDate())
    #     self.cur.execute(
    #         ''' INSERT INTO addbuys (item_name,signal_item_price,total_price,buys_type,quantity,date) VALUES (%s,%s,%s,%s,%s,%s)'''
    #         , (item_name, signal_item_price, total_item_price, item_type, quantity, date_create))
    #     self.db.commit()
    #     self.Show_all_buys()
    # self.Add_Data_To_history(3, 4)
    # self.History()

    # def Show_all_buys(self):
    #     self.cur.execute(''' SELECT item_name,buys_type,quantity,signal_item_price,total_price,date FROM addbuys ''')
    #     data = self.cur.fetchall()
    #     self.tableWidget_3.setRowCount(0)
    #     self.tableWidget_3.insertRow(0)
    #     for row, form in enumerate(data):
    #         for col, item in enumerate(form):
    #             self.tableWidget_3.setItem(row, col, QTableWidgetItem(str(item)))
    #             col += 1
    #         row_pos = self.tableWidget_3.rowCount()
    #         self.tableWidget_3.insertRow(row_pos)

    # def Add_Data_To_history(self, action, table):
    #     global user_id
    #     self.cur.execute(
    #         ''' INSERT INTO his VALUES (DEFAULT,%s,%s,%s,%s,%s)''',
    #         (user_id, action, table, datetime.datetime.now(), 1))
    #     self.db.commit()

    # def History(self):
    #     self.cur.execute('''SELECT * FROM his ORDER BY -dates''')
    #     analyst_data = self.cur.fetchall()
    #     self.tableWidget_8.setRowCount(0)
    #     self.tableWidget_8.insertRow(0)
    #     for row, form in enumerate(analyst_data):
    #         for col, item in enumerate(form):
    #             if col == 0:
    #                 sql = '''SELECT uid FROM his WHERE id =%s'''
    #                 self.cur.execute(sql, [item])
    #                 data = self.cur.fetchone()
    #                 sql = '''SELECT user_name FROM adduser WHERE id =%s'''
    #                 self.cur.execute(sql, [data[0]])
    #                 user_name = self.cur.fetchone()
    #                 self.tableWidget_8.setItem(row, col, QTableWidgetItem(str(user_name[0])))
    #             if col == 1:
    #                 action = ''
    #                 if analyst_data[row][2] == 1:
    #                     action = 'تسجيل الدخول'
    #                 if analyst_data[0][2] == 2:
    #                     action = 'تسجيل الخروج'
    #                 if analyst_data[row][2] == 3:
    #                     action = 'اضافة'
    #                 if analyst_data[row][2] == 4:
    #                     action = 'تعديل'
    #                 if analyst_data[row][2] == 5:
    #                     action = 'حذف'
    #                 if analyst_data[row][2] == 6:
    #                     action = 'بحث'
    #                 if analyst_data[row][2] == 7:
    #                     action = 'طباعة'
    #                 self.tableWidget_8.setItem(row, col, QTableWidgetItem(str(action)))
    #             if col == 2:
    #                 tables = ''
    #                 if analyst_data[row][3] == 1:
    #                     tables = 'مبيع يومي'
    #                 if analyst_data[row][3] == 2:
    #                     tables = 'تحليل'
    #                 if analyst_data[row][3] == 3:
    #                     tables = 'مشتريات'
    #                 if analyst_data[row][3] == 4:
    #                     tables = 'مراجعين'
    #                 if analyst_data[row][3] == 5:
    #                     tables = 'مستخدم'
    #                 self.tableWidget_8.setItem(row, col, QTableWidgetItem(str(tables)))
    #             if col == 3:
    #                 self.tableWidget_8.setItem(row, col, QTableWidgetItem(str(analyst_data[row][4])))
    #
    #             col += 1
    #         row_pos = self.tableWidget_8.rowCount()
    #         self.tableWidget_8.insertRow(row_pos)

    # def Show_avaliable_quantity(self):
    #     self.cur.execute(''' SELECT quantity FROM addbuys WHERE ''')
    #     quantity_avaliable = 0
    #     if quantity_avaliable < 1:
    #         date_quantity_ended = datetime.datetime.now()
    #         self.cur.execute(''' UPDATE addbuys SET quantity_avaliable=%s,date_ended=%s''',
    #                          (quantity_avaliable, date_quantity_ended))
    # def Log_In_Chieck(self):
    #     global user_id
    #     user_name = self.lineEdit.text()
    #     user_password = self.lineEdit_2.text()
    #     self.cur.execute(''' SELECT id,user_password,user_name FROM adduser WHERE user_name=%s ''', (user_name,))
    #     # self.cur.execute(''' SELECT * FROM adduser ''')
    #     data = self.cur.fetchone()
    #     if data != None:
    #         if data[2] == user_name and data[1] == user_password:
    #             user_id = data[0]
    #             self.groupBox.setEnabled(True)
    #             self.tabWidget.setCurrentIndex(3)
    #             # self.Add_Data_To_history(1, 5)
    #             # self.History()
    #         else:
    #             warning = QMessageBox.warning(self, '',
    #                                           "كلمة المرور او اسم المستخدم غير صحيحة هل تريد استعادة كلمة المرور؟",
    #                                           QMessageBox.Yes | QMessageBox.No)
    #             if warning == QMessageBox.Yes:
    #                 self.Open_ResetPassword_Page()
    #     else:
    #         warning = QMessageBox.warning(self, '',
    #                                       "كلمة المرور او اسم المستخدم غير صحيحة هل تريد استعادة كلمة المرور؟",
    #                                       QMessageBox.Yes | QMessageBox.No)
    #         if warning == QMessageBox.Yes:
    #             self.Open_ResetPassword_Page()

    # def Delete_All_History_Data(self):
    #     sql = ''' DELETE FROM his'''
    #     self.cur.execute(sql)
    #     self.db.commit()
    #     QMessageBox.information(self, 'info', 'تم حذف محتويات السجل بنجاح')
    #     self.tableWidget_8.setRowCount(0)
    #     self.tableWidget_8.insertRow(0)

    def Open_Sales_Page(self):
        self.tabWidget.setCurrentIndex(3)

    # def Open_Login_Page(self):
    #     self.tabWidget.setCurrentIndex(0)

    def Open_Settings_Page(self):
        self.tabWidget.setCurrentIndex(6)

    def Open_Print_Page(self):
        self.tabWidget.setCurrentIndex(7)

    # def Open_History_Page(self):
    #     self.tabWidget.setCurrentIndex(5)

    def Open_clients_Page(self):
        self.tabWidget.setCurrentIndex(2)

    # def Open_Analyse_Page(self):
    #     self.tabWidget.setCurrentIndex(4)

    # def Open_ResetPassword_Page(self):
    #     self.tabWidget.setCurrentIndex(1)

    def Light_Blue_Theme(self):
        style = open('thems/light_blue.css', 'r')
        style = style.read()
        self.setStyleSheet(style)

    def Dark_Blue_Theme(self):
        style = open('thems/darkblue.css', 'r')
        style = style.read()
        self.setStyleSheet(style)

    def Dark_Gray_Theme(self):
        style = open('thems/darkgray.css', 'r')
        style = style.read()
        self.setStyleSheet(style)

    def Dark_Orange_Theme(self):
        style = open('thems/darkorange.css', 'r')
        style = style.read()
        self.setStyleSheet(style)

    def Dark_Theme(self):
        style = open('thems/qdark.css', 'r')
        style = style.read()
        self.setStyleSheet(style)

    def Bio_Word(self, name, doctor, analysts, results, year, month, day, word_types, prev, genus):
        global if_print
        global two_files
        bio_analysts = []
        self.cur.execute(''' SELECT * FROM paths WHERE id=1 ''')
        mydata = self.cur.fetchone()
        word_files = mydata[1]
        save_word_files = mydata[2]
        self.cur.execute(
            ''' SELECT name,sub_category FROM addanalyst WHERE sub_category='bio' OR sub_category='hematology' ''')
        the_data = self.cur.fetchall()
        bio_hemo_all_analysts = []
        for tahlel_item in the_data:
            bio_hemo_all_analysts.append(tahlel_item[0])
            if tahlel_item[1] == 'bio':
                bio_analysts.append(tahlel_item[0])
        all_analyst = []
        all_result = []
        list_of_hemo_and_bio_analysts2 = []
        list_of_hemo_and_bio_results2 = []

        for index, itr in enumerate(all_analyst):
            if itr in bio_hemo_all_analysts and itr not in list_of_hemo_and_bio_analysts2:
                list_of_hemo_and_bio_analysts2.append(itr)
                list_of_hemo_and_bio_results2.append(all_result[index])
        number_of_docx_rows = 23
        if len(bio_hemo_all_analysts) < number_of_docx_rows:
            f = open('test-mydocx.docx', 'rb')
            f.read()
            document = Document(f)
        else:
            two_files=True
            print('two files open')
            f = open('test-mydocx.docx', 'rb')
            f.read()
            document = Document(f)
            f2 = open('test-mydocx2.docx', 'rb')
            document2 = Document(f2)
        row_num = False
        for word_type in word_types:
            for i in document.tables:
                for k in i.rows:
                    for j in k.cells:
                        for n in j.paragraphs:
                            for row in range(0, len(list_of_hemo_and_bio_analysts2)):
                                if 'bio' in word_types and 'hemo' in word_types:
                                    if n.text == str(row + 1):
                                        print(all_analyst[row], 'not')
                                        n.text = str(list_of_hemo_and_bio_analysts2[row]) + '  :'
                                        run = n.runs
                                        font = run[0].font
                                        font.bold = True
                                        font.size = Pt(11)
                                        font.name = 'Tahoma'
                                        n2 = n.add_run(str(list_of_hemo_and_bio_results2[row]))
                                        run = n2
                                        font = run.font
                                        font.bold = False
                                        font.size = Pt(11)
                                        font.name = 'Tahoma'
                                    if n.text == str(row + 1) + 'unit':
                                        tahlel = str(list_of_hemo_and_bio_analysts2[row])
                                        if tahlel in bio_analysts:
                                            n.text = 'mg/dl'
                                        else:
                                            if tahlel == 'Hb':
                                                n.text = 'gm/dl'
                                            if tahlel == 'PCV':
                                                n.text = '%'
                                            if tahlel == 'WBCs':
                                                n.text = 'cell/cumm'
                                            if tahlel == 'E.S.R':
                                                n.text = 'mm/1hr.'
                                            if tahlel == 'R.B.Sugar':
                                                n.text = 'mg/dl'
                                        run1 = n
                                        font1 = run1.runs[0].font
                                        font1.bold = True
                                        font1.size = Pt(12)
                                        font1.name = 'Times New Roman'
                                    # if n.text == str(row + 1) + 'defult':
                                    #     # n.text = analyst_and_result['normal']
                                    #     n.text = 'mydef'
                                    #     run1 = n
                                    #     font1 = run1.runs[0].font
                                    #     font1.bold = True
                                    #     font1.size = Pt(12)
                                    #     font1.name = 'Times New Roman'
                                    if row > 23:
                                        row_num = True

                                else:
                                    print('hahahah you will dont showing hahahaha')
                                    if 'bio' in word_type or 'hemo' in word_type:
                                        if n.text == str(row + 1):
                                            print(all_analyst[row], 'not')
                                            n.text = str(list_of_hemo_and_bio_analysts2[row]) + '  :'
                                            run = n.runs
                                            font = run[0].font
                                            font.bold = True
                                            font.size = Pt(11)
                                            font.name = 'Tahoma'
                                            n2 = n.add_run(str(list_of_hemo_and_bio_results2[row]))
                                            run = n2
                                            font = run.font
                                            font.bold = False
                                            font.size = Pt(11)
                                            font.name = 'Tahoma'
                                        if n.text == str(row + 1) + 'unit':
                                            tahlel = str(list_of_hemo_and_bio_analysts2[row])
                                            if tahlel in bio_analysts:
                                                n.text = 'mg/dl'
                                            else:
                                                if tahlel == 'Hb':
                                                    n.text = 'gm/dl'
                                                if tahlel == 'PCV':
                                                    n.text = '%'
                                                if tahlel == 'WBCs':
                                                    n.text = 'cell/cumm'
                                                if tahlel == 'E.S.R':
                                                    n.text = 'mm/1hr.'
                                                if tahlel == 'R.B.Sugar':
                                                    n.text = 'mg/dl'
                                            run1 = n
                                            font1 = run1.runs[0].font
                                            font1.bold = True
                                            font1.size = Pt(12)
                                            font1.name = 'Times New Roman'
                                        if n.text == str(row + 1) + 'defult':
                                            # n.text = analyst_and_result['normal']
                                            n.text = 'mydef'
                                            run1 = n
                                            font1 = run1.runs[0].font
                                            font1.bold = True
                                            font1.size = Pt(12)
                                            font1.name = 'Times New Roman'
                            for mynum in range(1, 25):
                                if n.text == str(mynum) or n.text == str(mynum) + 'unit' or n.text == str(
                                        mynum) + 'defult' or n.text == '':
                                    n.text = ''
                                    n.clear()
            if row_num:
                try:
                    for i2 in document2.tables:
                        for k2 in i2.rows:
                            for j2 in k2.cells:
                                for n3 in j2.paragraphs:
                                    for row in range(23, len(list_of_hemo_and_bio_analysts2)):
                                        if n3.text == str(row - 22):
                                            # print(all_analyst[row],'not')
                                            n3.text = str(list_of_hemo_and_bio_analysts2[row]) + '  :'
                                            run = n3.runs
                                            font = run[0].font
                                            font.bold = True
                                            font.size = Pt(11)
                                            font.name = 'Tahoma'
                                            n2 = n3.add_run(str(list_of_hemo_and_bio_results2[row]))
                                            run = n2
                                            font = run.font
                                            font.bold = False
                                            font.size = Pt(11)
                                            font.name = 'Tahoma'
                                        if n3.text == str(row - 22) + 'unit':
                                            tahlel = str(list_of_hemo_and_bio_analysts2[row])
                                            if tahlel in bio_analysts:
                                                n.text = 'mg/dl'
                                            else:
                                                if tahlel == 'Hb':
                                                    n.text = 'gm/dl'
                                                if tahlel == 'PCV':
                                                    n.text = '%'
                                                if tahlel == 'WBCs':
                                                    n.text = 'cell/cumm'
                                                if tahlel == 'E.S.R':
                                                    n.text = 'mm/1hr.'
                                                if tahlel == 'R.B.Sugar':
                                                    n.text = 'mg/dl'
                                            run1 = n3
                                            font1 = run1.runs[0].font
                                            font1.bold = True
                                            font1.size = Pt(12)
                                            font1.name = 'Times New Roman'
                                        if n3.text == str(row - 22) + 'defult':
                                            # n.text = analyst_and_result['normal']
                                            n3.text = 'mydef'
                                            run1 = n3
                                            font1 = run1.runs[0].font
                                            font1.bold = True
                                            font1.size = Pt(12)
                                            font1.name = 'Times New Roman'
                                    for mynum in range(1, 20):
                                        if n3.text == str(mynum) or n3.text == str(mynum) + 'unit' or n3.text == str(
                                                mynum) + 'defult' or n3.text == '':
                                            n3.text = ''
                                            n3.clear()
                except Exception as e:
                    print(e)
            row_num = False
            document.save('bio and hemo.docx')
            f.close()
            if two_files:
                document2.save('bio and hemo2.docx')
                f2.close()
            if word_type == 'GSE':
                f = open(r'%s\GSE latest.docx' % word_files, 'rb')
                f.read()

                document = Document(f)
                for i in document.tables:
                    for k in i.rows:
                        for j in k.cells:
                            for n in j.paragraphs:
                                if n.text == 'أسـم المريض :':
                                    if genus == 1:
                                        n.text = f'أسـم المريض                                {name}'
                                    else:
                                        n.text = f'أسـم المريضة                                {name}'
                                    run = n.runs
                                    font = run[0].font
                                    font.name = 'Monotype Koufi'
                                    font.bold = True
                                    font.size = Pt(11)
                                if n.text == 'حضرة الدكتورة   :':
                                    if doctor == 'عدوية شمس سعيد':
                                        n.text = f'حضرة الدكتورة                                {doctor}'
                                    else:
                                        n.text = f'حضرة الدكتور                                {doctor}'
                                    run = n.runs
                                    font = run[0].font
                                    font.name = 'Monotype Koufi'
                                    font.bold = True
                                    font.size = Pt(11)
                                if n.text == 'المحترم':
                                    if genus == 1:
                                        n.text = 'المحترم'
                                    else:
                                        n.text = 'المحترمة'
                                    run = n.runs

                                    font = run[0].font
                                    font.name = 'Monotype Koufi'
                                    font.bold = True
                                    font.size = Pt(11)
                                if n.text == 'المحترمة2':
                                    if doctor == 'عدوية شمس سعيد':
                                        n.text = f'المحترمة'
                                    else:
                                        n.text = f'المحترم'
                                    run = n.runs
                                    font = run[0].font
                                    font.name = 'Monotype Koufi'
                                    font.bold = True
                                    font.size = Pt(11)

                                    if n.text == 'Color:':
                                        for row in range(0, len(analysts)):
                                            if analyst_and_result['analyst'] == 'Color:GSE':
                                                k = analyst_and_result['result']
                                                n2 = n.add_run(f'  {k}')
                                                run = n2
                                                font = run.font
                                                font.bold = False
                                                font.size = Pt(14)
                                                font.name = 'Tahoma'

                                            if analyst_and_result['analyst'] == 'Consistency':
                                                if n.text == 'Consistency:':
                                                    k = analyst_and_result['result']
                                                    n2 = n.add_run(f'  {k}')
                                                    run = n2
                                                    font = run.font
                                                    font.bold = False
                                                    font.size = Pt(14)
                                                    font.name = 'Tahoma'

                                            if analyst_and_result['analyst'] == 'R.B.Cs:GSE':
                                                if n.text == 'R.B.Cs:':
                                                    k = analyst_and_result['result']
                                                    n2 = n.add_run(f'  {k}')
                                                    run = n2
                                                    font = run.font
                                                    font.bold = False
                                                    font.size = Pt(14)
                                                    font.name = 'Tahoma'

                                            if analyst_and_result['analyst'] == 'Pus cells:GSE':
                                                if n.text == 'Pus cells:':
                                                    k = analyst_and_result['result']
                                                    n2 = n.add_run(f'  {k}')
                                                    run = n2
                                                    font = run.font
                                                    font.bold = False
                                                    font.size = Pt(14)
                                                    font.name = 'Tahoma'

                                            if analyst_and_result['analyst'] == 'E. Histolytica':
                                                if n.text == 'E. Histolytica:':
                                                    k = analyst_and_result['result']
                                                    n2 = n.add_run(f'  {k}')
                                                    run = n2
                                                    font = run.font
                                                    font.bold = False
                                                    font.size = Pt(14)
                                                    font.name = 'Tahoma'

                                            if analyst_and_result['analyst'] == 'G. Lembilia':
                                                if n.text == 'G. Lembilia:':
                                                    k = analyst_and_result['result']
                                                    n2 = n.add_run(f'  {k}')
                                                    run = n2
                                                    font = run.font
                                                    font.bold = False
                                                    font.size = Pt(14)
                                                    font.name = 'Tahoma'

                                            if analyst_and_result['analyst'] == 'Ova':
                                                if n.text == 'Ova:':
                                                    k = analyst_and_result['result']
                                                    n2 = n.add_run(f'  {k}')
                                                    run = n2
                                                    font = run.font
                                                    font.bold = False
                                                    font.size = Pt(14)
                                                    font.name = 'Tahoma'

                                            if analyst_and_result['analyst'] == 'Bacteria:GSE':
                                                if n.text == 'Bacteria:':
                                                    k = analyst_and_result['result']
                                                    n2 = n.add_run(f'  {k}')
                                                    run = n2
                                                    font = run.font
                                                    font.bold = False
                                                    font.name = 'Tahoma'
                                                    font.size = Pt(14)
                                            if analyst_and_result['analyst'] == 'Monillia:GSE':
                                                if n.text == 'Monillia:':
                                                    k = analyst_and_result['result']
                                                    n2 = n.add_run(f'  {k}')
                                                    run = n2
                                                    font = run.font
                                                    font.bold = False
                                                    font.name = 'Tahoma'
                                                    font.size = Pt(14)

                                            if analyst_and_result['analyst'] == 'Fatty drop:GSE':
                                                if n.text == 'Fatty drop:':
                                                    k = analyst_and_result['result']
                                                    n2 = n.add_run(f'  {k}')
                                                    run = n2
                                                    font = run.font
                                                    font.bold = False
                                                    font.name = 'Tahoma'
                                                    font.size = Pt(14)

                                            if analyst_and_result['analyst'] == 'Other:GSE':
                                                if n.text == 'Other:':
                                                    k = analyst_and_result['result']
                                                    n2 = n.add_run(f' {k}')
                                                    run = n2
                                                    font = run.font
                                                    font.bold = False
                                                    font.size = Pt(14)
                                                    font.name = 'Tahoma'

                                if n.text == 'Date:    /     / 20':
                                    n.text = f'Date:   {year} /  {month} / {day}'
                                    run = n.runs
                                    font = run[0].font
                                    font.bold = True
                                    font.size = Pt(12)
                                    font.name = 'Tahoma'
                document.save(r'%s\GSE latest.docx' % save_word_files)
                f.close()

            if word_type == 'GUE':

                f = open(r'%s\GUE latest.docx' % word_files, 'rb')
                f.read()

                document = Document(f)
                for i in document.tables:
                    for k in i.rows:
                        for j in k.cells:
                            for n in j.paragraphs:
                                if n.text == 'أسـم المريض :':
                                    if genus == 1:
                                        n.text = f'أسـم المريض                                {name}'
                                    else:
                                        n.text = f'أسـم المريضة                                {name}'
                                    run = n.runs
                                    font = run[0].font
                                    font.name = 'Monotype Koufi'
                                    font.size = Pt(11)
                                if n.text == 'حضرة الدكتورة   :':
                                    if doctor == 'عدوية شمس سعيد':
                                        n.text = f'حضرة الدكتورة                                {doctor}'
                                    else:
                                        n.text = f'حضرة الدكتور                                {doctor}'
                                    run = n.runs
                                    font = run[0].font
                                    font.name = 'Monotype Koufi'
                                    font.size = Pt(11)
                                if n.text == 'المحترم':
                                    if genus == 1:
                                        n.text = 'المحترم'
                                    else:
                                        n.text = 'المحترمة'
                                    run = n.runs

                                    font = run[0].font
                                    font.name = 'Monotype Koufi'
                                    font.bold = True
                                    font.size = Pt(11)
                                if n.text == 'المحترمة2':
                                    if doctor == 'عدوية شمس سعيد':
                                        n.text = f'المحترمة'
                                    else:
                                        n.text = f'المحترم'
                                    run = n.runs
                                    font = run[0].font
                                    font.name = 'Monotype Koufi'
                                    font.bold = True
                                    font.size = Pt(11)

                                for row in range(0, len(analysts)):
                                    analyst_and_result = {
                                        'analyst': analysts[row],
                                        'result': results[row]
                                    }

                                    if analyst_and_result['analyst'] == 'Appearance':
                                        if n.text == 'Appearance :':
                                            k = analyst_and_result['result']
                                            n2 = n.add_run(f'  {k}')
                                            run = n2
                                            font = run.font
                                            font.name = 'Tahoma'
                                            font.bold = False
                                            font.size = Pt(14)

                                    if analyst_and_result['analyst'] == 'Reaction:GUE':
                                        if n.text == 'Reaction      :':
                                            k = analyst_and_result['result']
                                            n2 = n.add_run(f'  {k}')
                                            run = n2
                                            font = run.font
                                            font.name = 'Tahoma'
                                            font.bold = False
                                            font.size = Pt(14)

                                    if analyst_and_result['analyst'] == 'Albumin':
                                        if n.text == 'Albumin       :':
                                            k = analyst_and_result['result']
                                            n2 = n.add_run(f'  {k}')
                                            run = n2
                                            font = run.font
                                            font.name = 'Tahoma'
                                            font.bold = False
                                            font.size = Pt(14)

                                    if analyst_and_result['analyst'] == 'Sugar':
                                        if n.text == 'Sugar          :':
                                            k = analyst_and_result['result']
                                            n2 = n.add_run(f'  {k}')
                                            run = n2
                                            font = run.font
                                            font.name = 'Tahoma'
                                            font.bold = False
                                            font.size = Pt(14)

                                    if analyst_and_result['analyst'] == 'RBCs:GUE':
                                        if n.text == 'RBCs         :':
                                            k = analyst_and_result['result']
                                            n2 = n.add_run(f'  {k}')
                                            run = n2
                                            font = run.font
                                            font.name = 'Tahoma'
                                            font.bold = False
                                            font.size = Pt(14)

                                    if analyst_and_result['analyst'] == 'Pus cells:GUE':
                                        if n.text == 'Pus cells    :':
                                            k = analyst_and_result['result']
                                            n2 = n.add_run(f'  {k}')
                                            run = n2
                                            font = run.font
                                            font.name = 'Tahoma'
                                            font.bold = False
                                            font.size = Pt(14)

                                    if analyst_and_result['analyst'] == 'Epith .cells':
                                        if n.text == 'Epith .cells :':
                                            k = analyst_and_result['result']
                                            n2 = n.add_run(f'  {k}')
                                            run = n2
                                            font = run.font
                                            font.name = 'Tahoma'
                                            font.bold = False
                                            font.size = Pt(14)

                                    if analyst_and_result['analyst'] == 'Crystals':
                                        if n.text == 'Crystals     :':
                                            k = analyst_and_result['result']
                                            n2 = n.add_run(f'  {k}')
                                            run = n2
                                            font = run.font
                                            font.name = 'Tahoma'
                                            font.bold = False
                                            font.size = Pt(14)

                                    if analyst_and_result['analyst'] == 'Casts':
                                        if n.text == 'Casts        :':
                                            k = analyst_and_result['result']
                                            n2 = n.add_run(f'  {k}')
                                            run = n2
                                            font = run.font
                                            font.name = 'Tahoma'
                                            font.bold = False
                                            font.size = Pt(14)

                                    if analyst_and_result['analyst'] == 'Bacteria:GUE':
                                        if n.text == 'Bacteria    :':
                                            k = analyst_and_result['result']
                                            n2 = n.add_run(f'  {k}')
                                            run = n2
                                            font = run.font
                                            font.name = 'Tahoma'
                                            font.bold = False
                                            font.size = Pt(14)

                                    if analyst_and_result['analyst'] == 'Monillia:GUE':
                                        if n.text == 'Monillia     :':
                                            k = analyst_and_result['result']
                                            n2 = n.add_run(f'  {k}')
                                            run = n2
                                            font = run.font
                                            font.bold = False
                                            font.name = 'Tahoma'
                                            font.size = Pt(14)

                                    if analyst_and_result['analyst'] == 'Mucuse:GUE':
                                        if n.text == 'Mucuse        :':
                                            k = analyst_and_result['result']
                                            n2 = n.add_run(f'  {k}')
                                            run = n2
                                            font = run.font
                                            font.name = 'Tahoma'
                                            font.bold = False
                                            font.size = Pt(14)

                                    if analyst_and_result['analyst'] == 'Other:GUE':
                                        if n.text == 'Other        :':
                                            k = analyst_and_result['result']
                                            n2 = n.add_run(f'  {k}')
                                            run = n2
                                            font = run.font
                                            font.name = 'Tahoma'
                                            font.size = Pt(14)
                                if n.text == 'Date:    /     / 20':
                                    n.text = f'Date:   {year} /  {month} / {day}'
                                    run = n.runs
                                    font = run[0].font
                                    font.name = 'Tahoma'
                                    font.bold = True
                                    font.size = Pt(12)
                document.save(r'%s\GUE latest.docx' % save_word_files)
                f.close()
            if word_type == 'SFA':
                f = open(r'%s\SFA latest.docx' % word_files, 'rb')
                f.read()
                document = Document(f)
                for i in document.tables:
                    for k in i.rows:
                        for j in k.cells:

                            for n in j.paragraphs:
                                if n.text == 'أسـم المريض :':
                                    if genus == 1:
                                        n.text = f'أسـم المريض                                {name}'
                                    else:
                                        n.text = f'أسـم المريضة                                {name}'
                                    run = n.runs
                                    font = run[0].font
                                    font.name = 'Monotype Koufi'
                                    font.bold = True
                                    font.size = Pt(11)
                                if n.text == 'حضرة الدكتورة   :':
                                    if doctor == 'عدوية شمس سعيد':
                                        n.text = f'حضرة الدكتورة                                {doctor}'
                                    else:
                                        n.text = f'حضرة الدكتور                                {doctor}'
                                    run = n.runs
                                    font = run[0].font
                                    font.name = 'Monotype Koufi'
                                    font.bold = True
                                    font.size = Pt(11)
                                if n.text == 'المحترم':
                                    if genus == 1:
                                        n.text = 'المحترم'
                                    else:
                                        n.text = 'المحترمة'
                                    run = n.runs

                                    font = run[0].font
                                    font.name = 'Monotype Koufi'
                                    font.bold = True
                                    font.size = Pt(11)
                                if n.text == 'المحترمة2':
                                    if doctor == 'عدوية شمس سعيد':
                                        n.text = f'المحترمة'
                                    else:
                                        n.text = f'المحترم'
                                    run = n.runs
                                    font = run[0].font
                                    font.name = 'Monotype Koufi'
                                    font.bold = True
                                    font.size = Pt(11)
                                for row in range(0, len(analysts)):
                                    analyst_and_result = {
                                        'analyst': analysts[row],
                                        'result': results[row]
                                    }

                                    if analyst_and_result['analyst'] == 'Volume':
                                        if n.text == 'Volume       :':
                                            k = analyst_and_result['result']
                                            n2 = n.add_run(f'  {k}')
                                            run = n2
                                            font = run.font
                                            font.name = 'Tahoma'
                                            font.bold = False
                                            font.size = Pt(14)

                                    if analyst_and_result['analyst'] == 'Reaction:SFA':
                                        if n.text == 'Reaction      :':
                                            k = analyst_and_result['result']
                                            n2 = n.add_run(f'  {k}')
                                            run = n2
                                            font = run.font
                                            font.name = 'Tahoma'
                                            font.bold = False
                                            font.size = Pt(14)

                                    if analyst_and_result['analyst'] == 'Colour:SFA':
                                        if n.text == 'Colour        :':
                                            k = analyst_and_result['result']
                                            n2 = n.add_run(f'  {k}')
                                            run = n2
                                            font = run.font
                                            font.name = 'Tahoma'
                                            font.bold = False
                                            font.size = Pt(14)

                                    if analyst_and_result['analyst'] == 'Liquefaction':
                                        if n.text == 'Liquefaction :':
                                            k = analyst_and_result['result']
                                            n2 = n.add_run(f'  {k}')
                                            run = n2
                                            font = run.font
                                            font.name = 'Tahoma'
                                            font.bold = False
                                            font.size = Pt(14)

                                    if analyst_and_result['analyst'] == 'Count':
                                        if n.text == 'Count          :':
                                            k = analyst_and_result['result']
                                            n2 = n.add_run(f'  {k}')
                                            run = n2
                                            font = run.font
                                            font.name = 'Tahoma'
                                            font.bold = False
                                            font.size = Pt(14)

                                    if analyst_and_result['analyst'] == 'Motility:Active':
                                        if n.text == 'Active          :':
                                            k = analyst_and_result['result']
                                            n2 = n.add_run(f'  {k}')
                                            run = n2
                                            font = run.font
                                            font.name = 'Tahoma'
                                            font.bold = False
                                            font.size = Pt(14)

                                    if analyst_and_result['analyst'] == 'Motility:Sluggish':
                                        if n.text == 'Sluggish       :':
                                            k = analyst_and_result['result']
                                            n2 = n.add_run(f'  {k}')
                                            run = n2
                                            font = run.font
                                            font.name = 'Tahoma'
                                            font.bold = False
                                            font.size = Pt(14)

                                    if analyst_and_result['analyst'] == 'Motility:Dead':
                                        if n.text == 'Dead           :':
                                            k = analyst_and_result['result']
                                            n2 = n.add_run(f'  {k}')
                                            run = n2
                                            font = run.font
                                            font.name = 'Tahoma'
                                            font.bold = False
                                            font.size = Pt(14)

                                    if analyst_and_result['analyst'] == 'Morphology:Normal':
                                        if n.text == 'Normal        :':
                                            k = analyst_and_result['result']
                                            n2 = n.add_run(f'  {k}')
                                            run = n2
                                            font = run.font
                                            font.name = 'Tahoma'
                                            font.bold = False
                                            font.size = Pt(14)

                                    if analyst_and_result['analyst'] == 'Morphology:Abnormal':
                                        if n.text == 'Abnormal     :':
                                            k = analyst_and_result['result']
                                            n2 = n.add_run(f'  {k}')
                                            run = n2
                                            font = run.font
                                            font.name = 'Tahoma'
                                            font.bold = False
                                            font.size = Pt(14)

                                    if analyst_and_result['analyst'] == 'Morphology:Pus cells':
                                        if n.text == 'Pus cells       :':
                                            k = analyst_and_result['result']
                                            n2 = n.add_run(f'  {k}')
                                            run = n2
                                            font = run.font
                                            font.name = 'Tahoma'
                                            font.bold = False
                                            font.size = Pt(14)

                                    if analyst_and_result['analyst'] == 'Other:SFA':
                                        if n.text == 'Other           :':
                                            k = analyst_and_result['result']
                                            n2 = n.add_run(f'  {k}')
                                            run = n2
                                            font = run.font
                                            font.name = 'Tahoma'
                                            font.bold = False
                                            font.size = Pt(14)
                                if n.text == 'Date:    /     / 20':
                                    n.text = f'Date:   {year} /  {month} / {day}'
                                    run = n.runs
                                    font = run[0].font
                                    font.name = 'Tahoma'
                                    font.bold = True
                                    font.size = Pt(12)
                                # print('1'+n.text+'2')
                # for kd in document.paragraphs:
                #     print('1'+kd.text+'2'+'h')
                document.save(r'%s\SFA latest.docx' % save_word_files)
                f.close()
            if word_type == 'هرمونات مشترك':
                f = open(r'%s\هرمونات مشترك latest.docx' % word_files, 'rb')
                f.read()

                document = Document(f)

                for i in document.tables:
                    for k in i.rows:
                        for j in k.cells:

                            for n in j.paragraphs:
                                if n.text == '     أســم الـمــريــض    :':
                                    if genus == 1:
                                        n.text = f'أســم الـمــريــض                                {name}'
                                    else:
                                        n.text = f'أســم الـمــريــضة                                {name}'
                                    run = n.runs
                                    font = run[0].font
                                    font.bold = True
                                    font.size = Pt(14)
                                    font.name = 'Monotype Koufi'
                                if n.text == 'حـضـرة الـدكتـورة    الـفاضـــلة :                                                                                            ':
                                    if doctor != 'عدوية شمس سعيد':
                                        n.text = f'حـضـرة الـدكتـورة    الـفاضـــل                                {doctor}'
                                    else:
                                        n.text = f'حـضـرة الـدكتـورة    الـفاضـــلة                                {doctor}'
                                    run = n.runs
                                    font = run[0].font
                                    font.bold = True
                                    font.size = Pt(14)
                                    font.name = 'Monotype Koufi'
                                if n.text == 'الـمـحـتـرم':
                                    if genus == 1:
                                        n.text = 'المحترم'
                                    else:
                                        n.text = 'المحترمة'
                                    run = n.runs

                                    font = run[0].font
                                    font.name = 'Monotype Koufi'
                                    font.bold = True
                                    font.size = Pt(14)
                                if n.text == 'الـمـحـتـرمة2':
                                    if doctor == 'عدوية شمس سعيد':
                                        n.text = f'المحترمة'
                                    else:
                                        n.text = f'المحترم'
                                    run = n.runs
                                    font = run[0].font
                                    font.name = 'Monotype Koufi'
                                    font.bold = True
                                    font.size = Pt(14)
                                for row in range(0, len(analysts)):
                                    analyst_and_result = {
                                        'analyst': analysts[row],
                                        'result': results[row]
                                    }

                                    if analyst_and_result['analyst'] == 'Toxoplasma IgG':
                                        if n.text == '0r':
                                            k = analyst_and_result['result']

                                            n.text = f'  {k}'
                                            run = n.runs
                                            font = run[0].font
                                            font.bold = True
                                            font.size = Pt(10)
                                            font.name = 'Tahoma'

                                    if analyst_and_result['analyst'] == 'Toxoplasma IgM':
                                        if n.text == '1r':
                                            k = analyst_and_result['result']

                                            # print(k, ng, 'here ohf')
                                            n.text = f'  {k}'
                                            run = n.runs
                                            font = run[0].font
                                            font.bold = True
                                            font.size = Pt(10)
                                            font.name = 'Tahoma'

                                    if analyst_and_result['analyst'] == 'Cytomegalo Virus IgG':
                                        if n.text == '2r':
                                            k = analyst_and_result['result']

                                            n.text = f'  {k}'
                                            run = n.runs
                                            font = run[0].font
                                            font.bold = True
                                            font.size = Pt(10)
                                            font.name = 'Tahoma'

                                    if analyst_and_result['analyst'] == 'Cytomegalo Virus IgM':
                                        if n.text == '3r':
                                            k = analyst_and_result['result']

                                            n.text = f'  {k}'
                                            run = n.runs
                                            font = run[0].font
                                            font.bold = True
                                            font.size = Pt(10)
                                            font.name = 'Tahoma'

                                    if analyst_and_result['analyst'] == 'Rubella IgG':
                                        if n.text == '4r':
                                            k = analyst_and_result['result']

                                            n.text = f'  {k}'
                                            run = n.runs
                                            font = run[0].font
                                            font.bold = True
                                            font.size = Pt(10)
                                            font.name = 'Tahoma'

                                    if analyst_and_result['analyst'] == 'Rubella IgM':
                                        if n.text == '5r':
                                            k = analyst_and_result['result']
                                            n.text = f'  {k}'
                                            run = n.runs
                                            font = run[0].font
                                            font.bold = True
                                            font.size = Pt(10)
                                            font.name = 'Tahoma'

                                    if analyst_and_result['analyst'] == 'Anti - Phspholipin IgG':
                                        if n.text == '6r':
                                            k = analyst_and_result['result']
                                            n.text = f'  {k}'
                                            run = n.runs
                                            font = run[0].font
                                            font.bold = True
                                            font.size = Pt(10)
                                            font.name = 'Tahoma'

                                    if analyst_and_result['analyst'] == 'Anti - Phspholipin  IgM':
                                        if n.text == '7r':
                                            k = analyst_and_result['result']
                                            n.text = f'  {k}'
                                            run = n.runs
                                            font = run[0].font
                                            font.bold = True
                                            font.size = Pt(10)
                                            font.name = 'Tahoma'

                                    if analyst_and_result['analyst'] == 'Anti - Cardiolipin  IgG':
                                        if n.text == '8r':
                                            k = analyst_and_result['result']
                                            n.text = f'  {k}'
                                            run = n.runs
                                            font = run[0].font
                                            font.bold = True
                                            font.size = Pt(10)
                                            font.name = 'Tahoma'

                                    if analyst_and_result['analyst'] == 'Anti - Cardiolipin  IgM':
                                        if n.text == '9r':
                                            k = analyst_and_result['result']
                                            n.text = f'  {k}'
                                            run = n.runs
                                            font = run[0].font
                                            font.bold = True
                                            font.size = Pt(10)
                                            font.name = 'Tahoma'

                                    if analyst_and_result['analyst'] == 'Herps   IgG':
                                        if n.text == '10r':
                                            k = analyst_and_result['result']
                                            n.text = f'  {k}'
                                            run = n.runs
                                            font = run[0].font
                                            font.bold = True
                                            font.size = Pt(10)
                                            font.name = 'Tahoma'

                                    if analyst_and_result['analyst'] == 'Herpes  IgM':
                                        if n.text == '11r':
                                            k = analyst_and_result['result']
                                            n.text = f'  {k}'
                                            run = n.runs
                                            font = run[0].font
                                            font.bold = True
                                            font.size = Pt(10)
                                            font.name = 'Tahoma'

                                    if analyst_and_result['analyst'] == 'T3':
                                        if n.text == 'r0r':
                                            k = analyst_and_result['result']
                                            n.text = f'  {k}'
                                            run = n.runs
                                            font = run[0].font
                                            font.bold = True
                                            font.size = Pt(10)
                                            font.name = 'Tahoma'

                                    if analyst_and_result['analyst'] == 'T4':
                                        k = analyst_and_result['result']
                                        n.text = f'  {k}'
                                        run = n.runs
                                        font = run[0].font
                                        font.bold = True
                                        font.size = Pt(10)
                                        font.name = 'Tahoma'

                                    if analyst_and_result['analyst'] == 'TSH':
                                        if n.text == 'r2r':
                                            k = analyst_and_result['result']
                                            n.text = f'  {k}'
                                            run = n.runs
                                            font = run[0].font
                                            font.bold = True
                                            font.size = Pt(10)
                                            font.name = 'Tahoma'

                                    if analyst_and_result['analyst'] == 'LH':
                                        if n.text == 'r3r':
                                            k = analyst_and_result['result']
                                            n.text = f'  {k}'
                                            run = n.runs
                                            font = run[0].font
                                            font.bold = True
                                            font.size = Pt(10)
                                            font.name = 'Tahoma'

                                    if analyst_and_result['analyst'] == 'FSH':
                                        if n.text == 'r4r':
                                            k = analyst_and_result['result']
                                            n.text = f'  {k}'
                                            run = n.runs
                                            font = run[0].font
                                            font.bold = True
                                            font.size = Pt(10)
                                            font.name = 'Tahoma'

                                    if analyst_and_result['analyst'] == 'Prolactin':
                                        if n.text == 'r5r':
                                            k = analyst_and_result['result']
                                            n.text = f'  {k}'
                                            run = n.runs
                                            font = run[0].font
                                            font.bold = True
                                            font.size = Pt(10)
                                            font.name = 'Tahoma'

                                    if analyst_and_result['analyst'] == 'Testosterone':
                                        if n.text == 'r6r':
                                            k = analyst_and_result['result']
                                            n.text = f'  {k}'
                                            run = n.runs
                                            font = run[0].font
                                            font.bold = True
                                            font.size = Pt(10)
                                            font.name = 'Tahoma'
                                if n.text == 'Date:    /     / 20':
                                    n.text = f'Date:{year}/{month}/{day}'
                                    run = n.runs
                                    font = run[0].font
                                    font.bold = True
                                    font.size = Pt(12)
                                    font.name = 'Tahoma'
                                if n.text == '0r':
                                    n.text = ''
                                if n.text == '1r':
                                    n.text = ''
                                if n.text == '2r':
                                    n.text = ''
                                if n.text == '3r':
                                    n.text = ''
                                if n.text == '4r':
                                    n.text = ''
                                if n.text == '5r':
                                    n.text = ''
                                if n.text == '6r':
                                    n.text = ''
                                if n.text == '7r':
                                    n.text = ''
                                if n.text == '8r':
                                    n.text = ''
                                if n.text == '9r':
                                    n.text = ''
                                if n.text == '10r':
                                    n.text = ''
                                if n.text == '11r':
                                    n.text = ''
                                if n.text == 'r0r':
                                    n.text = ''
                                if n.text == 'r1r':
                                    n.text = ''
                                if n.text == 'r2r':
                                    n.text = ''
                                if n.text == 'r3r':
                                    n.text = ''
                                if n.text == 'r4r':
                                    n.text = ''
                                if n.text == 'r5r':
                                    n.text = ''
                                if n.text == 'r6r':
                                    n.text = ''
                for element in document.paragraphs:
                    if element.text == 'Date:    /     / 20':
                        element.text = f'Date:{year}/{month}/{day}'
                        run = element.runs
                        font = run[0].font
                        font.bold = True
                        font.size = Pt(12)
                        font.name = 'Tahoma'
                document.save(r'%s\هرمونات مشترك latest.docx' % save_word_files)
                f.close()

        # try:
        if prev == 'T':
            word = client.Dispatch("Word.Application")
            if os.path.exists(r'%s\bio and hemo.docx' % save_word_files):
                word.Documents.Open(r'%s\bio and hemo.docx' % save_word_files)
            if two_files: # it will don't do any thing because if else: there was no file i know you will don't understand me and i dont understand what i say
                if os.path.exists(r'%s\bio and hemo2.docx' % save_word_files):
                    word.Documents.Open(r'%s\bio and hemo2.docx' % save_word_files)
            if os.path.exists(r'%s\GSE latest.docx' % save_word_files):
                word.Documents.Open(r'%s\GSE latest.docx' % save_word_files)
            if os.path.exists(r'%s\SFA latest.docx' % save_word_files):
                word.Documents.Open(r'%s\SFA latest.docx' % save_word_files)
            if os.path.exists(r'%s\GUE latest.docx' % save_word_files):
                word.Documents.Open(r'%s\GUE latest.docx' % save_word_files)
            if os.path.exists(r'%s\هرمونات مشترك latest.docx' % save_word_files):
                word.Documents.Open(r'%s\هرمونات مشترك latest.docx' % save_word_files)


        else:
            move = 1
            word = client.Dispatch("Word.Application")
            if move == 1:
                if os.path.exists(r'%s\bio and hemo.docx' % save_word_files):
                    word.Documents.Open(r'%s\bio and hemo.docx' % save_word_files)
                    # word.ActiveDocument()
                    # word.ActiveDocument.ActiveWindow.View()
                    word.ActiveDocument.PrintOut(Background=False)
                    # time.sleep(2)
                    # pyautogui.press('enter')
                    # time.sleep(1.5)
                    # word.ActiveDocument.Close()
                    move = 0
                    warning = QMessageBox.warning(self, '',
                                                  "هل قمت بطباعة الملف؟\n لا تضغط نعم الا لو قمت بطباعته",
                                                  QMessageBox.Yes | QMessageBox.No)
                    if warning == QMessageBox.Yes:
                        move = 1
                        try:
                            word.ActiveDocument.Close()
                        except:
                            pass
                    os.remove(r'%s\bio and hemo2.docx' % save_word_files)

            if move == 1:
                if os.path.exists(r'%s\GSE latest.docx' % save_word_files):
                    word.Documents.Open(r'%s\GSE latest.docx' % save_word_files)
                    word.ActiveDocument.PrintOut(Background=False)
                    # time.sleep(2)
                    # pyautogui.press('enter')
                    # time.sleep(1.5)
                    # word.ActiveDocument.Close()
                    move = 0
                    warning = QMessageBox.warning(self, '',
                                                  "هل قمت بطباعة الملف؟\n لا تضغط نعم الا لو قمت بطباعته",
                                                  QMessageBox.Yes | QMessageBox.No)
                    if warning == QMessageBox.Yes:
                        move = 1
                        try:
                            word.ActiveDocument.Close()
                        except:
                            pass
                    os.remove(r'%s\GSE latest.docx' % save_word_files)

            if move == 1:

                if os.path.exists(r'%s\SFA latest.docx' % save_word_files):
                    word.Documents.Open(r'%s\SFA latest.docx' % save_word_files)
                    word.ActiveDocument.PrintOut(Background=False)
                    # time.sleep(2)
                    # pyautogui.press('enter')
                    # time.sleep(1.5)
                    # word.ActiveDocument.Close()
                    move = 0
                    warning = QMessageBox.warning(self, '',
                                                  "هل قمت بطباعة الملف؟\n لا تضغط نعم الا لو قمت بطباعته",
                                                  QMessageBox.Yes | QMessageBox.No)
                    if warning == QMessageBox.Yes:
                        move = 1
                        try:
                            word.ActiveDocument.Close()
                        except:
                            pass
                    # word.ActiveDocument.Close()
                    os.remove(r'%s\SFA latest.docx' % save_word_files)

            if move == 1:

                if os.path.exists(r'%s\GUE latest.docx' % save_word_files):
                    word.Documents.Open(r'%s\GUE latest.docx' % save_word_files)
                    word.ActiveDocument.PrintOut(Background=False)
                    # time.sleep(2)
                    # pyautogui.press('enter')
                    # time.sleep(1.5)
                    # word.ActiveDocument.Close()
                    move = 0
                    warning = QMessageBox.warning(self, '',
                                                  "هل قمت بطباعة الملف؟\n لا تضغط نعم الا لو قمت بطباعته",
                                                  QMessageBox.Yes | QMessageBox.No)
                    if warning == QMessageBox.Yes:
                        move = 1
                        try:
                            word.ActiveDocument.Close()
                        except:
                            pass
                    # word.ActiveDocument.Close()
                    os.remove(r'%s\GUE latest.docx' % save_word_files)

            if move == 1:
                if two_files:
                    if os.path.exists(r'%s\bio and hemo2.docx' % save_word_files):
                        word.Documents.Open(r'%s\bio and hemo2.docx' % save_word_files)
                        word.ActiveDocument.PrintOut(Background=False)
                        # time.sleep(2)
                        # pyautogui.press('enter')
                        # time.sleep(1.5)
                        # word.ActiveDocument.Close()
                        move = 0
                        warning = QMessageBox.warning(self, '',
                                                      "هل قمت بطباعة الملف؟\n لا تضغط نعم الا لو قمت بطباعته",
                                                      QMessageBox.Yes | QMessageBox.No)
                        if warning == QMessageBox.Yes:
                            move = 1
                            try:
                                word.ActiveDocument.Close()
                            except:
                                pass
                        # word.ActiveDocument.Close()
                        os.remove(r'%s\bio and hemo2.docx' % save_word_files)

            if move == 1:

                if os.path.exists(r'%s\هرمونات مشترك latest.docx' % save_word_files):
                    word.Documents.Open(r'%s\هرمونات مشترك latest.docx' % save_word_files)
                    word.ActiveDocument.PrintOut(Background=False)
                    # time.sleep(2)
                    # pyautogui.press('enter')
                    # time.sleep(1.5)
                    # word.ActiveDocument.Close()
                    move = 0
                    warning = QMessageBox.warning(self, '',
                                                  "هل قمت بطباعة الملف؟\n لا تضغط نعم الا لو قمت بطباعته",
                                                  QMessageBox.Yes | QMessageBox.No)
                    if warning == QMessageBox.Yes:
                        move = 1
                        try:
                            word.ActiveDocument.Close()
                        except:
                            pass
                    # word.ActiveDocument.Close()
                    os.remove(r'%s\هرمونات مشترك latest.docx' % save_word_files)
            self.Delete_Files()


        # except Exception as e:
        #     print(e)
        two_files=False



def main():
    app = QApplication(sys.argv)
    app.processEvents()
    window = mainapp()
    window.show()
    app.exec_()


if __name__ == '__main__':
    main()
